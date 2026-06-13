"""Build a supervisor-facing project report document.

The report is a readable synthesis of the current repository state. It reads
only local deterministic artifacts and does not run collectors, agents, MCP,
LLMs, ML, database writes, or live endpoints.
"""
from __future__ import annotations

import argparse
import json
import os
import sqlite3
from dataclasses import dataclass
from datetime import UTC, datetime
from html import escape
from pathlib import Path
from typing import Any, Sequence

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[2]
RESULTS_DIR = REPO_ROOT / "data" / "results"
DOCS_PROJECT_DIR = REPO_ROOT / "docs" / "project"
DEFAULT_ASSET_DIR = DOCS_PROJECT_DIR / "dozentenbericht_assets"
DEFAULT_MD_OUTPUT = DOCS_PROJECT_DIR / "dozentenbericht_ba_thesis.md"
DEFAULT_HTML_OUTPUT = DOCS_PROJECT_DIR / "dozentenbericht_ba_thesis.html"
DEFAULT_DOCX_OUTPUT = DOCS_PROJECT_DIR / "dozentenbericht_ba_thesis.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
TABLE_WIDTH_DXA = 9360


@dataclass(frozen=True)
class FigureSpec:
    """A report figure with a caption and optional paragraph."""

    path: Path
    caption: str
    note: str


def build_report(
    *,
    markdown_output: Path = DEFAULT_MD_OUTPUT,
    html_output: Path = DEFAULT_HTML_OUTPUT,
    docx_output: Path = DEFAULT_DOCX_OUTPUT,
    asset_dir: Path = DEFAULT_ASSET_DIR,
) -> dict[str, Any]:
    """Generate the Markdown and DOCX supervisor report."""

    asset_dir.mkdir(parents=True, exist_ok=True)
    data = collect_report_data()
    overview_path = write_pipeline_overview(asset_dir / "project_pipeline_overview.png")
    data["figures"].insert(
        0,
        FigureSpec(
            path=overview_path,
            caption="Projektlogik: deterministische Analyse vor Interpretation.",
            note=(
                "Die Abbildung zeigt die einfache Lesart fuer die Praesentation: "
                "Daten werden validiert, in Python analysiert und erst danach "
                "als Bericht, Dashboard oder spaeter als bounded Interpretation verwendet."
            ),
        ),
    )
    markdown = render_markdown(data, markdown_output=markdown_output)
    markdown_output.parent.mkdir(parents=True, exist_ok=True)
    markdown_output.write_text(markdown, encoding="utf-8")
    html_output.parent.mkdir(parents=True, exist_ok=True)
    html_output.write_text(render_html(data, html_output=html_output), encoding="utf-8")
    write_docx(data, docx_output)
    return {
        "markdown_output": str(markdown_output),
        "html_output": str(html_output),
        "docx_output": str(docx_output),
        "overview_figure": str(overview_path),
        "figure_count": len(data["figures"]),
        "generated_at_utc": data["generated_at_utc"],
    }


def collect_report_data() -> dict[str, Any]:
    """Read compact evidence from the current repository artifacts."""

    h1_summary = _read_csv("data/results/thesis_h1_summary.csv")
    h1_quality_pairwise = _read_csv("data/results/h1_forecast_quality_pairwise.csv")
    h1_forecast_synthesis = _read_csv(
        "data/results/h1_forecast_quality_synthesis.csv"
    )
    h1_claim_audit_summary = _read_csv(
        "data/results/h1_claim_evidence_audit_summary.csv"
    )
    h1_poll_comparison_result = _read_csv(
        "data/results/h1_poll_comparison_result_summary.csv"
    )
    h1_poll_claim_readiness = _read_csv(
        "data/results/h1_poll_claim_readiness_summary.csv"
    )
    h1_poll_scope_frontier = _read_csv(
        "data/results/h1_poll_scope_frontier_summary.csv"
    )
    h1_poll_decision_matrix = _read_csv(
        "data/results/h1_poll_decision_matrix_summary.csv"
    )
    h1_robust_poll_scope_quality_pairwise = _read_csv(
        "data/results/h1_robust_poll_scope_quality_pairwise.csv"
    )
    h1_robust_poll_scope_quality_summary = _read_csv(
        "data/results/h1_robust_poll_scope_quality_summary.csv"
    )
    h1_robust_poll_scope_unit_quality = _read_csv(
        "data/results/h1_robust_poll_scope_unit_quality_summary.csv"
    )
    h1_poll_comparison_unit_robustness = _read_csv(
        "data/results/h1_poll_comparison_unit_robustness_summary.csv"
    )
    h1_direct_poll_loss_decomposition = _read_csv(
        "data/results/h1_direct_poll_loss_decomposition_summary.csv"
    )
    h1_direct_poll_state_cluster = _read_csv(
        "data/results/h1_direct_poll_state_cluster_diagnostic_summary.csv"
    )
    h1_direct_poll_outlier_robustness = _read_csv(
        "data/results/h1_direct_poll_outlier_robustness_summary.csv"
    )
    h1_calibration_summary = _read_csv(
        "data/results/h1_calibration_diagnostic_summary.csv"
    )
    h1_calibration_pairwise = _read_csv(
        "data/results/h1_calibration_diagnostic_pairwise.csv"
    )
    h1_final_snapshot = _read_csv("data/results/h1_final_snapshot_summary.csv")
    h1_state_poll_snapshot = _read_csv(
        "data/results/h1_state_poll_snapshot_summary.csv"
    )
    h1_popular_vote = _read_csv("data/results/h1_popular_vote_summary.csv")
    h1_margin_threshold_readiness = _read_csv(
        "data/results/h1_margin_threshold_readiness.csv"
    )
    h1_state_poll_panel = _read_csv("data/results/h1_state_poll_panel_summary.csv")
    h1_state_poll_panel_temporal = _read_csv(
        "data/results/h1_state_poll_panel_temporal_claim_audit.csv"
    )
    h1_state_poll_panel_horizon = _read_csv(
        "data/results/h1_state_poll_panel_horizon_claim_audit.csv"
    )
    h1_state_poll_panel_horizon_state = _read_csv(
        "data/results/h1_state_poll_panel_horizon_state_support_summary.csv"
    )
    h1_state_poll_panel_near_quality = _read_csv(
        "data/results/h1_state_poll_panel_near_window_quality_summary.csv"
    )
    h1_state_poll_sensitivity = _read_csv(
        "data/results/h1_state_poll_snapshot_sensitivity.csv"
    )
    h1_state_poll_coverage = _read_csv(
        "data/results/h1_state_poll_snapshot_coverage.csv"
    )
    h1_rieke_state_forecast = _read_csv(
        "data/results/h1_rieke_state_forecast_summary.csv"
    )
    h1_270towin_state_forecast = _read_csv(
        "data/results/h1_270towin_state_forecast_summary.csv"
    )
    h1_270towin_poll_average = _read_csv(
        "data/results/h1_270towin_poll_average_summary.csv"
    )
    h1_state_source_consensus = _read_csv(
        "data/results/h1_state_source_consensus_summary.csv"
    )
    h1_competitive_state = _read_csv(
        "data/results/h1_competitive_state_diagnostic_summary.csv"
    )
    h1_state_poll_panel_competitiveness = _read_csv(
        "data/results/h1_state_poll_panel_competitiveness_summary.csv"
    )
    h1_state_poll_panel_state_significance = _read_csv(
        "data/results/h1_state_poll_panel_state_significance_summary.csv"
    )
    h2_summary = _read_csv("data/results/thesis_h2_summary.csv")
    h3_summary = _read_csv("data/results/thesis_h3_summary.csv")
    monitor_summary = _read_csv("data/results/monitor_v2_bounded_summary.csv")
    monitor_anomaly_review_summary = _read_csv(
        "data/results/monitor_anomaly_review_summary.csv"
    )
    swiss_comparison = _read_csv("data/results/swiss_referendum_10mio_comparison.csv")
    swiss_latest_source = _read_csv(
        "data/results/swiss_referendum_10mio_latest_source_comparison.csv"
    )
    swiss_information = _read_csv(
        "data/results/swiss_referendum_10mio_information_response.csv"
    )
    event_seed = _read_csv("data/events_timeline_seed.csv")
    swiss_polls = _read_csv("data/swiss_referendum_10mio_polls.csv")
    literature = _read_csv("data/literature/literature_index.csv")
    thesis_metadata = _read_json("data/results/thesis_consolidation_metadata.json")
    thesis_captions = _read_csv("data/results/thesis_table_figure_captions.csv")
    thesis_next_work = _read_csv("data/results/thesis_next_work_plan.csv")
    thesis_project_highlevel = _read_csv("data/results/thesis_project_highlevel_view.csv")
    thesis_source_worksheet = _read_csv("data/results/thesis_source_review_worksheet.csv")
    thesis_execution_checklist = _read_csv("data/results/thesis_execution_checklist.csv")
    thesis_advisor_handoff = _read_csv("data/results/thesis_advisor_handoff_package.csv")
    thesis_submission_readiness = _read_csv(
        "data/results/thesis_submission_readiness_board.csv"
    )
    thesis_drafting_sequence = _read_csv("data/results/thesis_drafting_sequence.csv")
    thesis_h1_h2_h3_bounded_chapter_draft = _read_csv(
        "data/results/thesis_h1_h2_h3_bounded_chapter_draft.csv"
    )
    thesis_h1_h2_h3_source_gated_drafting = _read_csv(
        "data/results/thesis_h1_h2_h3_source_gated_thesis_drafting_pass.csv"
    )
    thesis_h1_h2_h3_worksheet_drafting_bridge = _read_csv(
        "data/results/thesis_h1_h2_h3_worksheet_drafting_bridge.csv"
    )

    return {
        "generated_at_utc": datetime.now(UTC).replace(microsecond=0).isoformat(),
        "project": {
            "working_title": (
                "Informationelle Effizienz dezentraler Prognosemaerkte "
                "am Beispiel Polymarket im Vergleich zu traditionellen "
                "Prognosequellen"
            ),
            "database": _database_summary(),
            "folder_inventory": _folder_inventory(),
            "test_summary": _status_test_summary(),
        },
        "h1": _h1_data(
            h1_summary,
            h1_quality_pairwise,
            h1_forecast_synthesis,
            h1_claim_audit_summary,
            h1_poll_comparison_result,
            h1_poll_claim_readiness,
            h1_poll_scope_frontier,
            h1_poll_decision_matrix,
            h1_robust_poll_scope_quality_pairwise,
            h1_robust_poll_scope_quality_summary,
            h1_robust_poll_scope_unit_quality,
            h1_poll_comparison_unit_robustness,
            h1_direct_poll_loss_decomposition,
            h1_direct_poll_state_cluster,
            h1_direct_poll_outlier_robustness,
            h1_calibration_summary,
            h1_calibration_pairwise,
            h1_final_snapshot,
            h1_state_poll_snapshot,
            h1_popular_vote,
            h1_margin_threshold_readiness,
            h1_state_poll_panel,
            h1_state_poll_panel_temporal,
            h1_state_poll_panel_horizon,
            h1_state_poll_panel_horizon_state,
            h1_state_poll_panel_near_quality,
            h1_state_poll_sensitivity,
            h1_state_poll_coverage,
            h1_rieke_state_forecast,
            h1_270towin_state_forecast,
            h1_270towin_poll_average,
            h1_state_source_consensus,
            h1_competitive_state,
            h1_state_poll_panel_competitiveness,
            h1_state_poll_panel_state_significance,
        ),
        "h2": _h2_data(h2_summary, event_seed),
        "h3": _h3_data(h3_summary),
        "monitor": _monitor_data(monitor_summary, monitor_anomaly_review_summary),
        "swiss": _swiss_data(
            swiss_comparison,
            swiss_latest_source,
            swiss_information,
            swiss_polls,
        ),
        "literature": _literature_data(literature),
        "source_review": _source_review_worksheet_data(thesis_source_worksheet),
        "thesis_highlevel": _thesis_highlevel_data(thesis_metadata, thesis_captions),
        "project_highlevel": _project_highlevel_report_data(thesis_project_highlevel),
        "next_work": _next_work_report_data(thesis_next_work),
        "execution_checklist": _execution_checklist_report_data(thesis_execution_checklist),
        "advisor_handoff": _advisor_handoff_report_data(thesis_advisor_handoff),
        "submission_readiness": _submission_readiness_report_data(
            thesis_submission_readiness
        ),
        "drafting_sequence": _drafting_sequence_report_data(thesis_drafting_sequence),
        "bounded_chapter_draft": _bounded_chapter_draft_report_data(
            thesis_h1_h2_h3_bounded_chapter_draft
        ),
        "source_gated_drafting": _source_gated_drafting_report_data(
            thesis_h1_h2_h3_source_gated_drafting
        ),
        "worksheet_drafting_bridge": _worksheet_drafting_bridge_report_data(
            thesis_h1_h2_h3_worksheet_drafting_bridge
        ),
        "source_counts": {
            "curated_events": len(event_seed),
            "literature_rows": len(literature),
            "swiss_poll_rows": len(swiss_polls),
        },
        "figures": _figure_specs(),
    }


def render_markdown(data: dict[str, Any], *, markdown_output: Path) -> str:
    """Render the report body as Markdown for transparent review."""

    h1 = data["h1"]
    h2 = data["h2"]
    h3 = data["h3"]
    monitor = data["monitor"]
    swiss = data["swiss"]
    literature = data["literature"]
    source_review = data["source_review"]
    highlevel = data["thesis_highlevel"]
    project_highlevel = data["project_highlevel"]
    next_work = data["next_work"]
    execution_checklist = data["execution_checklist"]
    advisor_handoff = data["advisor_handoff"]
    submission_readiness = data["submission_readiness"]
    drafting_sequence = data["drafting_sequence"]
    bounded_chapter_draft = data["bounded_chapter_draft"]
    source_gated_drafting = data["source_gated_drafting"]
    worksheet_drafting_bridge = data["worksheet_drafting_bridge"]
    worksheet_drafting_bridge = data["worksheet_drafting_bridge"]
    db = data["project"]["database"]
    folders = data["project"]["folder_inventory"]
    insight_rows = [
        "| Bereich | Erkenntnis | Evidenz | Interpretation | Grenze |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in _interpretation_rows(data):
        insight_rows.append(
            "| {bereich} | {erkenntnis} | {evidenz} | {interpretation} | {grenze} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    method_rows = [
        "| Entscheidung | Begruendung | Konsequenz |",
        "| --- | --- | --- |",
    ]
    for decision, reason, consequence in _method_decision_rows():
        method_rows.append(
            f"| {decision.replace('|', ',')} | {reason.replace('|', ',')} | {consequence.replace('|', ',')} |"
        )
    literature_rows = [
        "| Quelle | Rolle in der Arbeit | Beitrag zur Interpretation | Status |",
        "| --- | --- | --- | --- |",
    ]
    for source in literature["sources"]:
        citation = (
            f"{source['authors']} ({source['year']}): "
            f"{source['title']}"
        ).replace("|", ",")
        role = str(source["role"]).replace("|", ",")
        note = str(source["research_note"]).replace("|", ",")
        literature_rows.append(
            f"| `{source['source_id']}` - {citation} | {role} | {note} | {source['status']} |"
        )

    next_work_rows = [
        "| Prioritaet | Workstream | Naechste Aktion | Guardrail |",
        "| --- | --- | --- | --- |",
    ]
    for row in next_work["rows"]:
        next_work_rows.append(
            "| {priority_order} | {workstream} | {next_action} | {guardrail} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    execution_rows = [
        "| Task | Kapitel | Schreibaktion | Fertig wenn | Advisor-Fragen |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in execution_checklist["rows"]:
        execution_rows.append(
            "| {task_id} | {chapter_title} | {draft_action_de} | {done_when_de} | {advisor_question_ids} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    advisor_handoff_rows = [
        "| Reihenfolge | Datei | Verwendung | Entscheidung | Grenze |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in advisor_handoff["rows"]:
        advisor_handoff_rows.append(
            "| {package_order} | `{path}` | {handoff_use_de} | {advisor_decision_de} | {boundary_de} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    project_highlevel_rows = [
        "| Projektteil | Status | Entscheidung | Naechstes Gate |",
        "| --- | --- | --- | --- |",
    ]
    for row in project_highlevel["rows"]:
        project_highlevel_rows.append(
            "| {project_layer} | {status_de} | {current_decision_de} | {next_gate_de} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    readiness_rows = [
        "| Gate | Status | Naechste Aktion | Grenze | Thesis-Nutzung |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in submission_readiness["rows"]:
        readiness_rows.append(
            "| {gate_area} | {current_status} | {next_action_de} | {blocker_or_limit_de} | {thesis_use_de} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    drafting_rows = [
        "| Prioritaet | Thesis-Abschnitt | Erlaubnis | Schreibaktion | Nicht behaupten |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in drafting_sequence["rows"]:
        drafting_rows.append(
            "| {priority_order} | {thesis_section} | {draft_permission} | {writing_action_de} | {must_not_claim_de} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    bounded_chapter_rows = [
        "| Kapitel | Methoden | Interpretationen | Literatur/Artefakte | Tabelle/Figur | Gate |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    for row in bounded_chapter_draft["chapter_rows"]:
        bounded_chapter_rows.append(
            "| {thesis_area} | `{method_evidence_ids}` | `{interpretation_evidence_ids}` | {literature_artifact_summary_de} | {table_figure_de} | {source_review_gate_summary_de} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    source_gated_chapter_rows = [
        "| Kapitel | Schritte | Manual Source Review | Tabelle/Figur | Status |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in source_gated_drafting["chapter_rows"]:
        source_gated_chapter_rows.append(
            "| {thesis_area} | {step_count}: {step_summary_de} | {manual_review_summary_de} | {table_figure_de} | {status_de} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    source_gated_step_rows = [
        "| Ordnung | Kapitel | Schreibschritt | Writer Action | Finalgate |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in source_gated_drafting["step_rows"]:
        source_gated_step_rows.append(
            "| {draft_sequence_order} | {thesis_area} | {draft_section_de} | {writer_action_de} | {final_gate_short_de} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )
    worksheet_bridge_rows = [
        "| Kapitel | Worksheets | Drafting | Tabelle/Figur | Gate |",
        "| --- | --- | --- | --- | --- |",
    ]
    for row in worksheet_drafting_bridge["chapter_rows"]:
        worksheet_bridge_rows.append(
            "| {thesis_area} | {worksheet_summary_de} | {drafting_summary_de} | {table_figure_de} | {gate_short_de} |".format(
                **{key: str(value).replace("|", ",") for key, value in row.items()}
            )
        )

    lines = [
        "# Dozentenbericht zur Bachelorarbeit",
        "",
        f"Erstellt: {data['generated_at_utc']}",
        "",
        f"**Arbeitstitel:** {data['project']['working_title']}.",
        "",
        "## Kurzfazit",
        "",
        (
            "Das Projekt hat eine deterministische Analysegrundlage aufgebaut. "
            "Statistische Kennzahlen werden in Python berechnet; LLMs oder "
            "Agenten interpretieren keine Rohdaten und berechnen keine Metriken."
        ),
        "",
        "- H1 ist ein Forecast-Qualitaetsvergleich.",
        "- H2 ist eine taegliche Event-Window-Analyse.",
        "- H3 ist eine Wallet-Tier-Timing-Diagnostik.",
        "- Der Monitor ist ein read-only Forschungsprototyp mit deterministischer Anomaly-Review-Queue.",
        "- Der Schweizer 10-Millionen-Referendumsvergleich laeuft als separater Datensammlungs-Track.",
        "",
        "## Aufbau wie in der Bachelorarbeit",
        "",
        "Dieser Dozentenbericht ist als Zwischenstand im Stil einer Bachelorarbeit aufgebaut:",
        "",
        "- Einleitung und Forschungsfrage: Warum Prediction Markets als Informationsmaerkte relevant sind.",
        "- Theorie und Literatur: Effizienz, Prediction Markets, Polling-Vergleiche, Wallet- und Mikrostrukturgrenzen.",
        "- Methodik: deterministische Python-Pipeline, validierte Artefakte, keine LLM-Metriken.",
        "- Empirie: H1 Forecast-Qualitaet, H2 Event-Window-Reaktion, H3 Wallet-Tier-Timing.",
        "- Erweiterung: read-only Monitor und Schweizer Referendumsvergleich als laufender Track.",
        "- Diskussion: Grenzen, belastbare Formulierungen und naechste Arbeitsschritte.",
        "",
        "## Highlevel-Projektstand",
        "",
        (
            "Der Review-Access bleibt pausiert. Der aktuelle Fortschritt liegt "
            "in der Thesis-Konsolidierung: Methoden, Interpretationen, Quellen, "
            "Tabellen und Figuren sind auf deterministische Artefakte gemappt."
        ),
        "",
        f"- Aktive Phase: {highlevel['active_phase']}.",
        (
            f"- Thesis-Paket: {highlevel['core_tables']} Kern-Tabellen und "
            f"{highlevel['core_figures']} Kern-Figuren; insgesamt "
            f"{highlevel['caption_rows']} Caption-Zeilen."
        ),
        (
            f"- Evidenzkarte: {highlevel['evidence_rows']} Evidence-Zeilen; "
            f"{highlevel['core_result_rows']} zentrale Resultatzeilen; "
            f"{highlevel['chapter_rows']} Kapitelplan-Zeilen."
        ),
        (
            f"- Citation-Gate: {highlevel['citation_packets']} Review-Pakete, "
            f"davon {highlevel['full_review_packets']} mit Full-Source-Review "
            "vor finaler Zitation."
        ),
        "- Agenten bleiben nur dokumentierter Ausblick; keine Runtime-Agenten, kein MCP, keine Modell-Router.",
        "",
        "| Ebene | Stand | Konsequenz fuer die Thesis |",
        "| --- | --- | --- |",
        *[
            f"| {row[0]} | {row[1]} | {row[2]} |"
            for row in highlevel["rows"]
        ],
        "",
        "## Dozentenpaket und Uebergabereihenfolge",
        "",
        (
            f"Das Advisor-Handoff-Paket ordnet {advisor_handoff['row_count']} "
            "Dateien fuer die naechste Betreuung. Zuerst kommt der "
            f"`{advisor_handoff['first_deliverable']}`, danach folgen "
            "Word-Bericht, Absprache-Checklist, Submission Readiness Board, "
            "Drafting Sequence, Feedback-Log und die Arbeitsdateien fuer "
            "Kapitel und Source Review."
        ),
        "",
        *advisor_handoff_rows,
        "",
        "## Projektmatrix fuer die naechste Abstimmung",
        "",
        (
            f"Die Projektmatrix fasst {project_highlevel['row_count']} Ebenen "
            "als Status-, Entscheidungs- und Gate-Sicht zusammen. Sie zeigt "
            "explizit, dass Review-Access pausiert bleibt und Agenten nur "
            "dokumentierter Ausblick sind."
        ),
        "",
        *project_highlevel_rows,
        "",
        "## Submission Readiness und finale Gates",
        "",
        (
            f"Das Submission Readiness Board trennt {submission_readiness['row_count']} "
            f"Gates in Draft-Arbeit, finale Blocker und Future Work. "
            f"Draft-ready Gates: {submission_readiness['draft_ready_count']}; "
            f"final blockierte Gates: {submission_readiness['final_blocked_count']}. "
            f"Source Review steht auf `{submission_readiness['source_review_status']}`, "
            f"Swiss steht auf `{submission_readiness['swiss_status']}`, Agenten "
            f"stehen auf `{submission_readiness['agent_status']}`."
        ),
        "",
        *readiness_rows,
        "",
        "## Schreibsequenz fuer den naechsten Entwurf",
        "",
        (
            f"Die Drafting Sequence ordnet {drafting_sequence['row_count']} "
            f"Schritte vom Quellenreview bis zur finalen QA. Erste Sequenz ist "
            f"`{drafting_sequence['first_step']}`, letzte Sequenz ist "
            f"`{drafting_sequence['final_step']}`. Bounded write-now: "
            f"{drafting_sequence['bounded_write_now_count']}; final blockiert: "
            f"{drafting_sequence['final_blocked_count']}; Future-work-only: "
            f"{drafting_sequence['future_work_only_count']}."
        ),
        "",
        *drafting_rows,
        "",
        "## Bounded H1-H2-H3 Kapitelentwurf",
        "",
        (
            f"Der neue H1-H2-H3 Bounded Chapter Draft liefert "
            f"{bounded_chapter_draft['row_count']} geordnete Prosa-Bausteine: "
            f"{bounded_chapter_draft['rows_per_chapter']} je H1, H2 und H3. "
            f"Bounded-draft-ready: {bounded_chapter_draft['bounded_ready_count']}; "
            f"final-submission-ready: {bounded_chapter_draft['final_ready_count']}."
        ),
        (
            "Der Dozent sieht damit direkt, dass jede empirische Methode und "
            "jede Interpretation eine Evidence-ID, Literatur-IDs, "
            "deterministische Artefakte, ein kuratiertes Tabellen-/Figurenpaar, "
            "Limitationen und ein Source-Review-Gate hat."
        ),
        "",
        *bounded_chapter_rows,
        "",
        (
            "Diese Sektion ist eine Schreibvorlage, kein finaler Zitations- "
            "oder Abgabeclaim. Keine neuen Kennzahlen, keine Rohartefakt-Dumps, "
            "keine Quellenstatus-Hochstufung und keine Runtime-Agenten."
        ),
        "",
        "## Source-Gated H1-H2-H3 Drafting Sequence",
        "",
        (
            f"Der source-gated Thesis-Drafting-Pass bringt den H1-H2-H3-Kern "
            f"in {source_gated_drafting['row_count']} paragraphenweise "
            f"Schreibschritte: {source_gated_drafting['rows_per_chapter']} je "
            "Kapitel. Bounded-draft-ready: "
            f"{source_gated_drafting['bounded_ready_count']}; "
            f"final-submission-ready: {source_gated_drafting['final_ready_count']}. "
            f"Manual Source Review: {source_gated_drafting['manual_rows_linked']} "
            f"Rows verlinkt, {source_gated_drafting['manual_pending_rows']} pending, "
            f"{source_gated_drafting['manual_final_ready_rows']} final-ready."
        ),
        (
            "Fuer den Dozenten ist das die konkrete Schreibreihenfolge nach dem "
            "Bounded Chapter Draft: Methode/Resultat setzen, Interpretation und "
            "Limitation setzen, Tabelle/Figur einbauen, Manual Source Review "
            "ausfuehren, Finalgate und Future-Agent-Grenze sichtbar lassen."
        ),
        source_gated_drafting["review_control_de"],
        "",
        *source_gated_chapter_rows,
        "",
        *source_gated_step_rows,
        "",
        (
            "Auch diese Sequenz ist kein finaler Zitations- oder Abgabeclaim: "
            "Page-/Section-Note, Claim-Support, Blocked-Wording und Citation-Use "
            "bleiben manuelle Gates; Agenten bleiben documentation-only Future Work."
        ),
        "",
        "## Worksheet-to-Drafting Bridge fuer H1-H2-H3",
        "",
        (
            f"Die Worksheet-Drafting-Bridge verbindet {worksheet_drafting_bridge['worksheet_rows']} "
            "manuelle H1-H2-H3 Worksheet-Zeilen mit "
            f"{worksheet_drafting_bridge['drafting_steps']} source-gated Schreibschritten. "
            f"Method rows: {worksheet_drafting_bridge['method_rows']}; "
            f"Interpretation rows: {worksheet_drafting_bridge['interpretation_rows']}; "
            f"Source/artifact gaps: {worksheet_drafting_bridge['source_artifact_gap_rows']}; "
            f"final-release rows: {worksheet_drafting_bridge['final_release_ready_rows']}."
        ),
        worksheet_drafting_bridge["source_artifact_rule_de"],
        (
            "Damit sieht der Dozent in einer kleinen Bruecke, welche Quellen- "
            "und Artefaktpflicht vor jedem Absatz gilt und warum der Haupttext "
            "bei T2/F1, T3/F2 und T4/F3 bleibt."
        ),
        "",
        *worksheet_bridge_rows,
        "",
        worksheet_drafting_bridge["future_agent_boundary_de"],
        "",
        "## Naechste Arbeitsschritte",
        "",
        (
            f"Der Next-Work-Plan ordnet {next_work['row_count']} Workstreams. "
            f"Erste Prioritaet ist `{next_work['first_workstream']}`, letzte "
            f"QA-Prioritaet ist `{next_work['final_workstream']}`."
        ),
        "",
        *next_work_rows,
        "",
        "## Kapitelweise Umsetzungscheckliste",
        "",
        (
            f"Die Execution-Checkliste uebersetzt die Highlevel-View in "
            f"{execution_checklist['row_count']} Kapitelaufgaben. Erste Aufgabe "
            f"ist `{execution_checklist['first_task']}`, letzte Aufgabe ist "
            f"`{execution_checklist['final_task']}`."
        ),
        "",
        *execution_rows,
        "",
        (
            "Die Liste ist kein neues empirisches Ergebnis. Sie zeigt nur, "
            "welche Kapitel mit welchen Inputs, Done-Kriterien und "
            "Advisor-Fragen abgearbeitet werden sollen."
        ),
        "",
        "## Forschungsfrage und Hypothesen",
        "",
        (
            "Die Leitfrage lautet, inwiefern Polymarket-Preise Informationen "
            "waehrend politischer Ereignisse abbilden, schneller oder anders "
            "als traditionelle Prognosequellen reagieren und ob aggregierte "
            "Wallet-Aktivitaet als frueher Timing-Indikator sichtbar wird."
        ),
        "",
        "- H1: Polymarket wird als Probability-Forecast gegen traditionelle Forecast- oder Poll-derived Vergleichsquellen getestet.",
        "- H2: Vorab kuratierte Ereignisse werden in taeglichen Event-Windows ausgewertet.",
        "- H3: Wallet-Aktivitaet wird ueber verteilungsbasierte Tiers als Timing-Diagnostik analysiert.",
        "",
        "## Wissenschaftlicher Quellenrahmen",
        "",
        (
            f"Der lokale Literaturindex umfasst {literature['source_count']} "
            f"Quellen; fuer diesen Bericht werden "
            f"{literature['selected_source_count']} wissenschaftlich relevante "
            f"Kernquellen als Rahmen verwendet. Statusverteilung: "
            f"{literature['status_counts_text']}."
        ),
        (
            f"Das neue Source-Review-Worksheet enthaelt "
            f"{source_review['worksheet_rows']} manuelle Review-Zeilen, davon "
            f"{source_review['priority_1_rows']} Priority-1-Methodenquellen "
            f"und {source_review['blocked_rows']} blockierte oder Future-Work-Quelle. "
            "Alle Reviewer-Entscheide bleiben pending."
        ),
        "",
        *literature_rows,
        "",
        literature["citation_boundary"],
        "",
        "## Methodisches Design und Begruendung",
        "",
        (
            "Die Arbeit operationalisiert informationelle Effizienz nicht als "
            "direkt beobachtbare Eigenschaft, sondern ueber drei Proxies. "
            "Forecast-Qualitaet wird mit Brier-Verlusten und Vergleichstests "
            "gemessen; Ereignisreaktionen werden nur fuer vorab kuratierte "
            "Events ausgewertet; Wallet-Signale bleiben aggregierte "
            "Timing-Diagnostik. Damit sind die Resultate reproduzierbar und "
            "methodisch begrenzt."
        ),
        "",
        "- Alle Kennzahlen stammen aus Python-Artefakten unter `data/results`.",
        "- RCP bleibt ausgeschlossen, solange keine dokumentierte Probability-Transformation existiert.",
        "- Granger-Outputs werden nicht kausal interpretiert.",
        "- Monitor- und Live-Daten bleiben read-only und schreiben bounded Artefakte.",
        "",
        "## Zentrale Erkenntnisse, Begruendung und Interpretation",
        "",
        (
            "Die wichtigste inhaltliche Verbesserung ist die Trennung zwischen "
            "Ergebnis, Interpretation und Grenze. Dadurch kann der Dozent sehen, "
            "was bereits empirisch tragfaehig ist und welche Aussagen bewusst "
            "nicht gemacht werden."
        ),
        "",
        *insight_rows,
        "",
        "## Warum dieses Vorgehen methodisch sinnvoll ist",
        "",
        *method_rows,
        "",
        "## Projektstruktur",
        "",
        f"- SQLite-Datenbank: {db['table_count']} Tabellen.",
        f"- Ergebnisartefakte: {folders.get('data/results', 0)} Dateien unter `data/results`.",
        f"- Analyse-Module: {folders.get('operations/analysis', 0)} Dateien.",
        f"- Collector-Module: {folders.get('operations/collectors', 0)} Dateien.",
        f"- Tests: {folders.get('tests', 0)} Testdateien; letzter Status: {data['project']['test_summary']}.",
        "",
        "## H1 - Forecast-Qualitaet",
        "",
        f"- Beobachtungen: {h1['observation_count']}.",
        f"- Mean Brier Polymarket: {h1['brier_polymarket']:.4f}.",
        f"- Mean Brier FiveThirtyEight: {h1['brier_fivethirtyeight']:.4f}.",
        f"- Mean Brier 50-Prozent-Baseline: {h1['brier_always_50']:.4f}.",
        f"- Mean Brier Vortag-Polymarket: {h1['brier_prior_day']:.4f}.",
        (
            f"- DM p-Wert Polymarket vs FiveThirtyEight: "
            f"{h1['dm_polymarket_vs_538']:.3g}."
        ),
        (
            f"- Polymarket niedrigerer Tagesverlust als FiveThirtyEight: "
            f"{h1['pm_better_vs_538_count']} von {h1['pm_vs_538_count']} "
            f"Tagen ({h1['pm_better_vs_538_share'] * 100:.1f}%)."
        ),
        (
            f"- Mittlerer Verlustvorteil gegenueber FiveThirtyEight: "
            f"{h1['mean_loss_advantage_vs_538']:.4f} Brier-Punkte."
        ),
        (
            f"- H1-Synthesis ueber traditionelle Vergleichsquellen: "
            f"{h1['synthesis_aggregate_support_count']} von "
            f"{h1['synthesis_evidence_row_count']} Vergleichszeilen stuetzen "
            f"Polymarket im mittleren Brier; "
            f"{h1['synthesis_majority_support_count']} von "
            f"{h1['synthesis_evidence_row_count']} zeigen auch eine Mehrheit "
            f"niedrigerer Einzelfallverluste. "
            f"Breiter Viele-Faelle-Beweis: "
            f"{h1['synthesis_broad_support_count']} von "
            f"{h1['synthesis_evidence_row_count']}."
        ),
        (
            f"- H1-Claim-Evidence-Audit: "
            f"{h1['claim_audit_support_row_count']} von "
            f"{h1['claim_audit_row_count']} Audit-Zeilen stuetzen "
            f"Polymarket begrenzt, "
            f"{h1['claim_audit_contradiction_row_count']} widerspricht dem "
            f"starken Claim; bei direkt pollbezogenen Zeilen sind "
            f"{h1['claim_audit_direct_poll_support_row_count']} von "
            f"{h1['claim_audit_direct_poll_row_count']} stuetzend und "
            f"{h1['claim_audit_direct_poll_contradiction_row_count']} "
            f"widersprechend. Breiter User-Claim belegt: "
            f"{h1['claim_audit_broad_user_claim_proven']}."
        ),
        (
            f"- H1-Poll-Comparison-Result: Im primaeren <=90-Tage-"
            f"Low/Middle-Poll-Distanz-Scope hat Polymarket in "
            f"{h1['poll_result_primary_pm_count']} von "
            f"{h1['poll_result_primary_row_count']} State-Date-Zeilen "
            f"({h1['poll_result_primary_pm_share'] * 100:.1f}%) den "
            f"niedrigeren Brier-Verlust; poll-derived gewinnt "
            f"{h1['poll_result_primary_poll_count']} Zeilen. Auf State-Ebene "
            f"sind es {h1['poll_result_primary_pm_state_count']} von "
            f"{h1['poll_result_primary_state_count']} States, exakter "
            f"einseitiger p-Wert "
            f"{h1['poll_result_primary_p_value']:.4f}. Direkt pollbezogen "
            f"stuetzen {h1['poll_result_direct_poll_support_count']} von "
            f"{h1['poll_result_direct_poll_row_count']} Audit-Zeilen "
            f"Polymarket begrenzt; das Vollpanel bleibt Gegenbeleg mit "
            f"poll-derived {h1['poll_result_full_panel_poll_count']} von "
            f"{h1['poll_result_full_panel_row_count']} Zeilen. Status: "
            f"{h1['poll_result_goal_status']}."
        ),
        (
            f"- H1-Poll-Claim-Readiness: "
            f"{h1['poll_claim_supported_bounded_count']} von "
            f"{h1['poll_claim_row_count']} Claim-Zeilen stuetzen den "
            f"bounded <=90-Tage Low/Middle-Poll-Distanz-Scope, "
            f"{h1['poll_claim_counterexample_count']} sind Gegenbeispiel-"
            f"Scopes und {h1['poll_claim_mixed_mean_count']} zeigen nur "
            f"Mean-Loss-Stuetze ohne Fall- oder State-Mehrheit. Im bounded "
            f"Scope hat Polymarket {h1['poll_claim_primary_pm_count']} von "
            f"{h1['poll_claim_primary_count']} State-Date-Zeilen "
            f"({h1['poll_claim_primary_pm_share'] * 100:.1f}%) und "
            f"{h1['poll_claim_state_month_pm_count']} von "
            f"{h1['poll_claim_state_month_count']} State-Month-Einheiten "
            f"(exact p={h1['poll_claim_state_month_p_value']:.2g}, "
            f"95-Prozent-Untergrenze "
            f"{h1['poll_claim_state_month_ci_low']:.3f}) auf seiner Seite. "
            f"Bounded Claim supported: {h1['poll_claim_bounded_supported']}; "
            f"breiter Claim belegt: {h1['poll_claim_broad_proven']}."
        ),
        (
            f"- H1-Poll-Scope-Frontier: "
            f"{h1['poll_frontier_robust_scope_count']} von "
            f"{h1['poll_frontier_row_count']} Horizont-x-Poll-Distanz-Scopes "
            f"erfuellen die robuste Regel. Der groesste robuste Scope ist "
            f"{h1['poll_frontier_largest_horizon']} + "
            f"{h1['poll_frontier_largest_tier']}: Polymarket "
            f"{h1['poll_frontier_largest_pm_count']} von "
            f"{h1['poll_frontier_largest_row_count']} State-Date-Zeilen "
            f"({h1['poll_frontier_largest_pm_share'] * 100:.1f}%), "
            f"{h1['poll_frontier_largest_state_month_pm_count']} von "
            f"{h1['poll_frontier_largest_state_month_count']} State-Month-"
            f"Einheiten, exact p="
            f"{h1['poll_frontier_largest_state_month_p_value']:.3g}. "
            f"Der staerkste Scope bleibt "
            f"{h1['poll_frontier_strongest_scope_id']} mit "
            f"{h1['poll_frontier_strongest_row_count']} Zeilen und p="
            f"{h1['poll_frontier_strongest_p_value']:.2g}. "
            f"<=90 Tage ueber alle Distanzen stuetzen Polymarket zwar in "
            f"{h1['poll_frontier_lte_90_all_pm_count']} von "
            f"{h1['poll_frontier_lte_90_all_row_count']} Zeilen "
            f"({h1['poll_frontier_lte_90_all_pm_share'] * 100:.1f}%), "
            f"aber State-Month p="
            f"{h1['poll_frontier_lte_90_all_state_month_p_value']:.3g}; "
            f"das Vollpanel bleibt Gegenbeleg mit poll-derived "
            f"{h1['poll_frontier_full_panel_poll_count']} von "
            f"{h1['poll_frontier_full_panel_row_count']} Zeilen. Status "
            f"{h1['poll_frontier_goal_status']}."
        ),
        (
            f"- H1-Poll-Decision-Matrix: "
            f"{h1['poll_decision_robust_yes_count']} von "
            f"{h1['poll_decision_row_count']} Entscheidungszeilen sind robuste "
            f"bounded-Yes-Zeilen, {h1['poll_decision_mixed_mean_count']} zeigen "
            f"Mean-Loss-Stuetze ohne Fall-/Unit-Mehrheit und "
            f"{h1['poll_decision_counterexample_count']} sind Gegenbelege. "
            f"Groesster robuster Scope: Polymarket "
            f"{h1['poll_decision_largest_pm_count']} von "
            f"{h1['poll_decision_largest_row_count']} State-Date-Zeilen "
            f"({h1['poll_decision_largest_pm_share'] * 100:.1f}%), "
            f"{h1['poll_decision_largest_state_month_pm_count']} von "
            f"{h1['poll_decision_largest_state_month_count']} State-Month-"
            f"Einheiten, p={h1['poll_decision_largest_p_value']:.4f}. "
            f"Kalibrierungskontext: "
            f"{h1['poll_decision_calibration_aggregate_count']} von "
            f"{h1['poll_decision_calibration_pairwise_count']} Pairwise-Reihen "
            f"stuetzen Polymarket im mittleren Brier, aber nur "
            f"{h1['poll_decision_calibration_majority_count']} auch per "
            f"Fallmehrheit. Bounded ready "
            f"{h1['poll_decision_bounded_ready']}; breiter Claim "
            f"{h1['poll_decision_broad_proven']}; Status "
            f"{h1['poll_decision_goal_status']}."
        ),
        (
            f"- H1-Robust-Poll-Scope-Quality: "
            f"{h1['robust_quality_forecast_row_count']} Forecast-Zeilen aus "
            f"{h1['robust_quality_case_count']} State-Date-Faellen und "
            f"{h1['robust_quality_scope_count']} robusten Poll-Scopes. "
            f"Groesster robuster Scope: Polymarket "
            f"{h1['robust_quality_largest_pm_count']} von "
            f"{h1['robust_quality_largest_case_count']} Zeilen "
            f"({h1['robust_quality_largest_pm_share'] * 100:.1f}%), "
            f"Mean Brier {h1['robust_quality_largest_pm_brier']:.4f} vs "
            f"{h1['robust_quality_largest_poll_brier']:.4f}, ECE "
            f"{h1['robust_quality_largest_pm_ece']:.4f} vs "
            f"{h1['robust_quality_largest_poll_ece']:.4f}, Separation "
            f"{h1['robust_quality_largest_pm_separation']:.4f} vs "
            f"{h1['robust_quality_largest_poll_separation']:.4f}. "
            f"Staerkster robuster Scope: Polymarket "
            f"{h1['robust_quality_strongest_pm_count']} von "
            f"{h1['robust_quality_strongest_case_count']} Zeilen "
            f"({h1['robust_quality_strongest_pm_share'] * 100:.1f}%), "
            f"Mean Brier {h1['robust_quality_strongest_pm_brier']:.4f} vs "
            f"{h1['robust_quality_strongest_poll_brier']:.4f}, ECE "
            f"{h1['robust_quality_strongest_pm_ece']:.4f} vs "
            f"{h1['robust_quality_strongest_poll_ece']:.4f}. "
            f"Dort sind alle Outcomes positiv, deshalb ist Separation nicht "
            f"definiert. Breiter Claim belegt "
            f"{h1['robust_quality_broad_claim_proven']}."
        ),
        (
            f"- H1-Robust-Poll-Scope-Unit-Quality: Die robusten Scopes "
            f"bleiben auch auf weniger wiederholten Einheiten sichtbar. "
            f"Groesster robuster Scope: State-Ebene Polymarket "
            f"{h1['robust_unit_largest_state_pm_count']} von "
            f"{h1['robust_unit_largest_state_count']} (p="
            f"{h1['robust_unit_largest_state_p_value']:.3g}), State-Month "
            f"{h1['robust_unit_largest_state_month_pm_count']} von "
            f"{h1['robust_unit_largest_state_month_count']} (p="
            f"{h1['robust_unit_largest_state_month_p_value']:.4f}), "
            f"State-Horizon {h1['robust_unit_largest_state_horizon_pm_count']} "
            f"von {h1['robust_unit_largest_state_horizon_count']} (p="
            f"{h1['robust_unit_largest_state_horizon_p_value']:.3g}). "
            f"Staerkster robuster Scope: States "
            f"{h1['robust_unit_strongest_state_pm_count']} von "
            f"{h1['robust_unit_strongest_state_count']} (p="
            f"{h1['robust_unit_strongest_state_p_value']:.3g}), "
            f"State-Month {h1['robust_unit_strongest_state_month_pm_count']} "
            f"von {h1['robust_unit_strongest_state_month_count']} (p="
            f"{h1['robust_unit_strongest_state_month_p_value']:.2g}). "
            f"Medianer State-Month-Brier-Vorteil: "
            f"{h1['robust_unit_largest_state_month_median_advantage']:.4f} "
            f"im groessten und "
            f"{h1['robust_unit_strongest_state_month_median_advantage']:.4f} "
            f"im staerksten Scope. Breiter Claim belegt "
            f"{h1['robust_unit_broad_claim_proven']}."
        ),
        (
            f"- H1-Poll-Comparison-Unit-Robustness: Der primaere Scope haelt "
            f"auch nach Aggregation: Polymarket wird in "
            f"{h1['poll_unit_state_pm_count']} von "
            f"{h1['poll_unit_state_count']} States, "
            f"{h1['poll_unit_state_month_pm_count']} von "
            f"{h1['poll_unit_state_month_count']} State-Month-Einheiten und "
            f"{h1['poll_unit_state_horizon_pm_count']} von "
            f"{h1['poll_unit_state_horizon_count']} State-Horizon-Einheiten "
            f"gestuetzt; State-Month exact p="
            f"{h1['poll_unit_state_month_p_value']:.2g}, 95-Prozent-"
            f"Untergrenze {h1['poll_unit_state_month_ci_low']:.3f}. "
            f"Full-Panel-State-Month-Gegenbeleg: poll-derived "
            f"{h1['poll_unit_full_panel_state_month_poll_count']} von "
            f"{h1['poll_unit_full_panel_state_month_count']}; "
            f"Late-High-Distance-State-Month-Gegenbeleg: poll-derived "
            f"{h1['poll_unit_late_high_state_month_poll_count']} von "
            f"{h1['poll_unit_late_high_state_month_count']}, exact p="
            f"{h1['poll_unit_late_high_state_month_poll_p_value']:.4f}. Status: "
            f"{h1['poll_unit_goal_status']}."
        ),
        (
            f"- H1-Direct-Poll-Loss-Decomposition: Direkte Poll-Transform-"
            f"Vergleiche ergeben Mean Brier {h1['direct_poll_loss_pm_brier']:.4f} "
            f"fuer Polymarket vs {h1['direct_poll_loss_poll_brier']:.4f} fuer "
            f"poll-derived Comparatoren. Polymarket hat niedrigeren Verlust in "
            f"{h1['direct_poll_loss_pm_count']} von "
            f"{h1['direct_poll_loss_case_count']} Source-State-Faellen, "
            f"poll-derived in {h1['direct_poll_loss_poll_count']}; die "
            f"Polymarket-Gewinnfaelle haben aber im Mittel "
            f"{h1['direct_poll_loss_pm_win_mean_advantage']:.4f} Brier-Vorteil "
            f"gegenueber {h1['direct_poll_loss_poll_win_mean_advantage']:.4f} "
            f"bei poll-derived Gewinnfaellen. Das erklaert den aggregierten "
            f"Brier-Vorteil, ersetzt aber keinen Fallmehrheits- oder "
            f"Viele-Wahlen-Beweis."
        ),
        (
            f"- H1-Direct-Poll-State-Cluster-Diagnostic: Auf "
            f"{h1['direct_poll_state_cluster_state_count']} State-Clustern "
            f"bleibt der gleichgewichtete mittlere Verlustvorteil positiv "
            f"({h1['direct_poll_state_cluster_mean_advantage']:.4f}; "
            f"Bootstrap-95%-Intervall "
            f"{h1['direct_poll_state_cluster_bootstrap_ci_low']:.4f} bis "
            f"{h1['direct_poll_state_cluster_bootstrap_ci_high']:.4f}; "
            f"Sign-Flip-p={h1['direct_poll_state_cluster_sign_flip_p']:.4f}). "
            f"Die State-Mehrheit geht aber gegen Polymarket: "
            f"{h1['direct_poll_state_cluster_pm_state_count']} States fuer "
            f"Polymarket, {h1['direct_poll_state_cluster_poll_state_count']} "
            f"fuer poll-derived Comparatoren. Das stuetzt einen mittleren "
            f"Verlustvorteil, nicht eine State-Mehrheitsbehauptung."
        ),
        (
            f"- H1-Direct-Poll-Outlier-Robustness: Der gleiche "
            f"State-Cluster-Mean von "
            f"{h1['direct_poll_outlier_full_mean_advantage']:.4f} bleibt nach "
            f"jeder einzelnen State-Entfernung positiv; das Minimum ist "
            f"{h1['direct_poll_outlier_min_leave_one_mean']:.4f} ohne "
            f"{h1['direct_poll_outlier_most_influential_state']}. Entfernt man "
            f"die groessten positiven State-Beitraege, bleibt der Mean bis "
            f"{h1['direct_poll_outlier_top_k_positive']} entfernte States "
            f"positiv und kippt bei "
            f"{h1['direct_poll_outlier_first_nonpositive_k']} entfernten "
            f"States auf "
            f"{h1['direct_poll_outlier_first_nonpositive_mean']:.4f}. "
            f"Das zeigt: nicht ein einzelner Ausreisser, aber Konzentration in "
            f"den groessten positiven State-Beitraegen; Status "
            f"{h1['direct_poll_outlier_goal_status']}."
        ),
        (
            f"- H1-State-Source-Konsens: "
            f"{h1['state_source_consensus_case_count']} Source-State-Vergleiche "
            f"ueber {h1['state_source_consensus_state_count']} States; "
            f"Polymarket hat niedrigeren Verlust in "
            f"{h1['state_source_consensus_pm_case_count']} Source-State-Faellen, "
            f"traditionelle Comparatoren in "
            f"{h1['state_source_consensus_comparator_case_count']}. "
            f"Im All-Source-State-Konsens gewinnt Polymarket "
            f"{h1['state_source_consensus_pm_state_count']} States, "
            f"Comparatoren {h1['state_source_consensus_comparator_state_count']}, "
            f"Ties {h1['state_source_consensus_tie_state_count']}. "
            f"Bei States mit zwei direkten Poll-Transform-Quellen gewinnt "
            f"Polymarket {h1['state_source_consensus_direct_two_pm_state_count']} "
            f"von {h1['state_source_consensus_direct_two_state_count']} States."
        ),
        (
            f"- H1-Competitive-State-Diagnose: In den niedrigsten "
            f"Comparator-Distanz-Terzilen gewinnt Polymarket "
            f"{h1['competitive_state_all_low_pm_count']} von "
            f"{h1['competitive_state_all_low_case_count']} All-Source-Faellen "
            f"und {h1['competitive_state_direct_low_pm_count']} von "
            f"{h1['competitive_state_direct_low_case_count']} direkten "
            f"Poll-Transform-Faellen. In der hoechsten Distanz-Terzile gewinnt "
            f"Polymarket {h1['competitive_state_all_high_pm_count']} von "
            f"{h1['competitive_state_all_high_case_count']} All-Source-Faellen, "
            f"Comparatoren {h1['competitive_state_all_high_comparator_count']} "
            f"von {h1['competitive_state_all_high_case_count']}. Das stuetzt "
            f"eine begrenzte Competitive-State-Ausnahme, aber keinen breiten "
            f"Viele-Faelle-Beweis."
        ),
        (
            f"- H1-State-Date-Competitiveness-x-Horizon: Im <=90-Tage-Fenster "
            f"und in Low/Middle-Poll-Distanz-Terzilen hat Polymarket in "
            f"{h1['panel_comp_late_non_safe_pm_count']} von "
            f"{h1['panel_comp_late_non_safe_row_count']} State-Date-Zeilen "
            f"niedrigeren Verlust und in "
            f"{h1['panel_comp_late_non_safe_state_support_count']} von "
            f"{h1['panel_comp_late_non_safe_state_count']} States eine "
            f"Mehrheit niedrigerer Verluste. In der spaeten High-Distance-"
            f"Terzile gewinnt Polymarket {h1['panel_comp_late_high_pm_count']} "
            f"von {h1['panel_comp_late_high_row_count']} Zeilen, poll-derived "
            f"{h1['panel_comp_late_high_poll_count']} von "
            f"{h1['panel_comp_late_high_row_count']}. Das ist ein starker "
            f"spaeter Competitive-Poll-Befund, aber wegen wiederholter "
            f"State-Date-Zeilen kein unabhaengiger Viele-Wahlen-Beweis."
        ),
        (
            f"- H1-State-Level-Signifikanzdiagnose: Fuer dieselben spaeten "
            f"Low/Middle-Poll-Distanz-Faelle stuetzt Polymarket "
            f"{h1['state_sign_late_non_safe_pm_state_count']} von "
            f"{h1['state_sign_late_non_safe_state_count']} States; der "
            f"exakte einseitige Binomial-p-Wert betraegt "
            f"{h1['state_sign_late_non_safe_p_value']:.4f}, die exakte "
            f"95-Prozent-Untergrenze der Support-Quote "
            f"{h1['state_sign_late_non_safe_ci_low']:.3f}. Die spaeten "
            f"High-Distance-States bleiben ein Gegenbeleg: poll-derived "
            f"{h1['state_sign_late_high_poll_state_count']} von "
            f"{h1['state_sign_late_high_state_count']} States."
        ),
        (
            f"- H1-Kalibrierungsdiagnostik: "
            f"{h1['calibration_forecast_case_rows']} Forecast-Case-Zeilen "
            f"aus {h1['calibration_forecast_source_count']} Quellen und "
            f"{h1['calibration_pairwise_count']} Pairwise-Reihen; "
            f"{h1['calibration_aggregate_support_count']} von "
            f"{h1['calibration_pairwise_count']} zeigen niedrigeren mittleren "
            f"Polymarket-Brier, {h1['calibration_majority_support_count']} "
            f"von {h1['calibration_pairwise_count']} auch eine Mehrheit "
            f"niedrigerer Einzelfallverluste, breiter Viele-Faelle-Beweis "
            f"{h1['calibration_broad_support_count']} von "
            f"{h1['calibration_pairwise_count']}."
        ),
        (
            f"- 50-State-Kalibrierung: Polymarket Mean Brier "
            f"{h1['calibration_pm_state_brier']:.4f} und Fixed-Bin-ECE "
            f"{h1['calibration_pm_state_ece']:.4f}; Rieke ECE "
            f"{h1['calibration_rieke_state_ece']:.4f}, 270toWin/JHK ECE "
            f"{h1['calibration_270_state_ece']:.4f}. Das ist ein "
            f"Forecast-Qualitaets-, aber kein klarer Kalibrierungssieg."
        ),
        (
            f"- Final-Snapshot-Erweiterung: Polymarket niedrigerer Verlust in "
            f"{h1['final_snapshot_pm_lower_loss_count']} von "
            f"{h1['final_snapshot_case_count']} geloesten 2024-Outcomes; "
            f"Mean Brier Polymarket {h1['final_snapshot_mean_pm_brier']:.4f} "
            f"vs 538 final forecast {h1['final_snapshot_mean_traditional_brier']:.4f}."
        ),
        (
            f"- State-Poll-Snapshot-Erweiterung: Polymarket niedrigerer Verlust in "
            f"{h1['state_poll_snapshot_pm_lower_loss_count']} von "
            f"{h1['state_poll_snapshot_case_count']} geloesten State-Outcomes; "
            f"Mean Brier Polymarket {h1['state_poll_snapshot_mean_pm_brier']:.4f} "
            f"vs poll-derived {h1['state_poll_snapshot_mean_poll_brier']:.4f}."
        ),
        (
            f"- 270toWin-Polling-Average-Erweiterung: "
            f"{h1['two_seventy_poll_average_case_count']} gematchte State-Outcomes; "
            f"Polymarket niedrigerer Verlust in "
            f"{h1['two_seventy_poll_average_pm_lower_loss_count']} Faellen, "
            f"poll-derived in {h1['two_seventy_poll_average_poll_lower_loss_count']}. "
            f"Mean Brier Polymarket "
            f"{h1['two_seventy_poll_average_mean_pm_brier']:.4f} vs "
            f"270toWin poll-derived "
            f"{h1['two_seventy_poll_average_mean_poll_brier']:.4f}."
        ),
        (
            f"- Popular-Vote-Erweiterung: {h1['popular_vote_case_count']} nationale "
            f"Tageszeilen fuer Trump popular vote; Polymarket niedrigerer Verlust "
            f"in {h1['popular_vote_pm_lower_loss_count']} Zeilen, poll-derived in "
            f"{h1['popular_vote_poll_lower_loss_count']}. Mean Brier Polymarket "
            f"{h1['popular_vote_mean_pm_brier']:.4f} vs poll-derived "
            f"{h1['popular_vote_mean_poll_brier']:.4f}; dieser Zusatz ist ein "
            f"Gegenbeleg zum starken Claim."
        ),
        (
            f"- Margin-Threshold-Readiness: "
            f"{h1['margin_threshold_candidate_count']} Trump-State-Margin-Maerkte "
            f"geprueft; {h1['margin_threshold_with_538_poll_count']} haben "
            f"538-State-Poll-Average-Zeilen, aber "
            f"{h1['margin_threshold_with_clob_overlap_count']} haben CLOB-Historie "
            f"im bewahrten 538-Fenster. H1-kompatible neue Brier-Faelle: "
            f"{h1['margin_threshold_compatible_count']}; "
            f"{h1['margin_threshold_no_overlap_count']} blockiert durch fehlende "
            f"zeitliche Ueberlappung und "
            f"{h1['margin_threshold_missing_poll_count']} durch fehlende 538-State-Polls."
        ),
        (
            f"- State-Date-Poll-Panel: "
            f"{h1['state_poll_panel_case_count']} gematchte State-Date-Zeilen "
            f"ueber {h1['state_poll_panel_state_count']} States und "
            f"{h1['state_poll_panel_date_count']} Daten; Polymarket hat nur in "
            f"{h1['state_poll_panel_pm_lower_loss_count']} Zeilen niedrigeren "
            f"Verlust, die poll-derived 538-Transformation in "
            f"{h1['state_poll_panel_poll_lower_loss_count']}. Mean Brier "
            f"Polymarket {h1['state_poll_panel_mean_pm_brier']:.4f} vs "
            f"poll-derived {h1['state_poll_panel_mean_poll_brier']:.4f}."
        ),
        (
            f"- Temporal-Diagnose des State-Date-Panels: In den "
            f"Polymarket-stuetzenden Monaten "
            f"{h1['state_poll_temporal_support_months']} liegen "
            f"{h1['state_poll_temporal_support_row_count']} Zeilen ueber "
            f"{h1['state_poll_temporal_support_state_count']} States vor; "
            f"Polymarket hat dort in "
            f"{h1['state_poll_temporal_support_pm_lower_loss_count']} Zeilen "
            f"niedrigeren Verlust, poll-derived in "
            f"{h1['state_poll_temporal_support_poll_lower_loss_count']}. "
            f"Mean Brier {h1['state_poll_temporal_support_mean_pm_brier']:.4f} "
            f"vs {h1['state_poll_temporal_support_mean_poll_brier']:.4f}. "
            f"Das erklaert den spaeten Polymarket-Vorteil, hebt aber den "
            f"negativen Vollpanel-Befund nicht auf."
        ),
        (
            f"- Forecast-Horizon-Diagnose: Im <=90-Tage-Fenster vor der "
            f"Wahl ({h1['state_poll_horizon_near_bins']}) liegen "
            f"{h1['state_poll_horizon_near_row_count']} Zeilen ueber "
            f"{h1['state_poll_horizon_near_state_count']} States vor; "
            f"Polymarket hat in "
            f"{h1['state_poll_horizon_near_pm_lower_loss_count']} Zeilen "
            f"niedrigeren Verlust, poll-derived in "
            f"{h1['state_poll_horizon_near_poll_lower_loss_count']}. "
            f"Mean Brier {h1['state_poll_horizon_near_mean_pm_brier']:.4f} "
            f"vs {h1['state_poll_horizon_near_mean_poll_brier']:.4f}. "
            f"Diese Horizon-Diagnose stuetzt Polymarket naeher an der Wahl, "
            f"bleibt aber ein wiederholtes Forecast-Row-Fenster."
        ),
        (
            f"- State-Level-Horizon-Diagnose: Im selben <=90-Tage-Fenster "
            f"stuetzt Polymarket {h1['state_poll_horizon_state_pm_mean_support_count']} "
            f"von {h1['state_poll_horizon_state_count']} States nach mittlerem "
            f"Brier und {h1['state_poll_horizon_state_pm_majority_support_count']} "
            f"von {h1['state_poll_horizon_state_count']} States nach Mehrheit "
            f"niedrigerer Tagesverluste; "
            f"{h1['state_poll_horizon_state_poll_support_count']} States "
            f"stuetzen Polymarket nicht."
        ),
        (
            f"- <=90-Day-Score-Quality-Diagnose: "
            f"{h1['state_poll_near_quality_forecast_row_count']} Forecast-Zeilen "
            f"aus {h1['state_poll_near_quality_case_count']} State-Date-Faellen "
            f"und zwei Quellen. Polymarket hat niedrigeren Mean Brier "
            f"{h1['state_poll_near_quality_pm_mean_brier']:.4f} vs "
            f"{h1['state_poll_near_quality_poll_mean_brier']:.4f}, niedrigeren "
            f"Fixed-Bin-ECE {h1['state_poll_near_quality_pm_ece']:.4f} vs "
            f"{h1['state_poll_near_quality_poll_ece']:.4f} und hoehere "
            f"Probability-Separation "
            f"{h1['state_poll_near_quality_pm_separation']:.4f} vs "
            f"{h1['state_poll_near_quality_poll_separation']:.4f}. "
            f"Das stuetzt Forecast-Qualitaet im spaeten Fenster, bleibt aber "
            f"ein wiederholtes State-Date-Forecast-Panel."
        ),
        (
            f"- Poll-Transform-Sensitivitaet: MAE "
            f"{h1['state_poll_sensitivity_min_mae']:.1f} bis "
            f"{h1['state_poll_sensitivity_max_mae']:.1f} Prozentpunkte; "
            f"Polymarket bleibt im mittleren Brier in allen "
            f"{h1['state_poll_sensitivity_row_count']} Parameterzeilen niedriger "
            f"und hat in {h1['state_poll_sensitivity_min_pm_lower_loss_count']} bis "
            f"{h1['state_poll_sensitivity_max_pm_lower_loss_count']} von "
            f"{h1['state_poll_snapshot_case_count']} State-Outcomes den niedrigeren "
            f"Einzelfallverlust."
        ),
        (
            f"- State-Poll-Coverage-Audit: {h1['state_poll_coverage_state_count']} "
            f"US-States geprueft, {h1['state_poll_coverage_polymarket_market_count']} "
            f"mit Polymarket-State-Markt, aber nur "
            f"{h1['state_poll_coverage_valid_pair_count']} mit REP/DEM-Zeilen "
            f"im bewahrten 538-Polling-Average-Snapshot. "
            f"{h1['state_poll_coverage_missing_poll_count']} States fallen wegen "
            f"fehlender 538-Snapshot-Pollwerte aus."
        ),
        (
            f"- Rieke-50-State-Erweiterung: {h1['rieke_state_case_count']} "
            f"geloeste State-Outcomes gegen ein unabhaengiges pollbasiertes "
            f"Rieke-Modell; Mean Brier Polymarket "
            f"{h1['rieke_state_mean_pm_brier']:.4f} vs Rieke "
            f"{h1['rieke_state_mean_rieke_brier']:.4f}. Polymarket hat nur in "
            f"{h1['rieke_state_pm_lower_loss_count']} von "
            f"{h1['rieke_state_case_count']} State-Einzelfaellen den niedrigeren "
            f"Verlust, Rieke in {h1['rieke_state_rieke_lower_loss_count']} von "
            f"{h1['rieke_state_case_count']}."
        ),
        (
            f"- 270toWin/JHK-50-State-Erweiterung: "
            f"{h1['two_seventy_state_case_count']} geloeste State-Outcomes, davon "
            f"{h1['two_seventy_state_exact_case_count']} exakt ausgewiesene "
            f"State-Wahrscheinlichkeiten und "
            f"{h1['two_seventy_state_censored_case_count']} zensierte "
            f">99.9-Prozent-Boundary-Werte; Mean Brier Polymarket "
            f"{h1['two_seventy_state_mean_pm_brier']:.4f} vs 270toWin/JHK "
            f"{h1['two_seventy_state_mean_270_brier']:.4f}. Polymarket hat in "
            f"{h1['two_seventy_state_pm_lower_loss_count']} von "
            f"{h1['two_seventy_state_case_count']} Einzelfaellen den niedrigeren "
            f"Verlust, 270toWin/JHK in "
            f"{h1['two_seventy_state_270_lower_loss_count']} von "
            f"{h1['two_seventy_state_case_count']}."
        ),
        (
            f"- H1-Zusatzchecks insgesamt: "
            f"{h1['final_snapshot_case_count'] + h1['state_poll_snapshot_case_count']} "
            f"geloeste Outcomes in den 538-nahen Zusatzchecks, davon "
            f"{h1['final_snapshot_pm_lower_loss_count'] + h1['state_poll_snapshot_pm_lower_loss_count']} "
            f"mit niedrigerem Polymarket-Verlust. Die Rieke- und 270toWin/JHK-"
            f"State-Reihen werden separat berichtet, weil sie dasselbe "
            f"Praesidentschaftsrennen mit anderen traditionellen Modellen "
            f"abdecken."
        ),
        "",
        "## H2 - Event-Window-Reaktion",
        "",
        f"- Kuratierte Ereignisse: {h2['event_count']}.",
        f"- Kompakte H2-Zeilen: {h2['summary_rows']}.",
        "- Beispielhafte Primaerfenster:",
        *[
            f"  - {row['event']}: {row['change_pp']:+.1f} Prozentpunkte"
            for row in h2["primary_examples"]
        ],
        "",
        "## H3 - Wallet-Tier-Timing",
        "",
        f"- Aligned model rows: {h3['model_rows']}.",
        f"- Tier counts: {h3['tier_counts_text']}.",
        (
            f"- Staerkste dokumentierte Lead-Lag-Korrelation: "
            f"{h3['top_correlation_label']} = {h3['top_correlation']:.4f}."
        ),
        (
            f"- Kleinster dokumentierter Granger-p-Wert: "
            f"{h3['min_granger_label']} = {h3['min_granger_p']:.4f}."
        ),
        "",
        "## Monitor-Prototyp",
        "",
        f"- Recorded replay rows: {monitor['snapshot_count']}.",
        f"- Severity counts: {monitor['severity_counts_text']}.",
        f"- Latest live dashboard markets: {monitor['live_market_count']}; alert rows: {monitor['live_alert_count']}.",
        f"- Wallet graph: {monitor['wallet_graph_nodes']} nodes, {monitor['wallet_graph_edges']} edges.",
        (
            f"- Anomaly review queue: {monitor['anomaly_queue_rows']} Cases "
            f"({monitor['anomaly_high_priority_count']} high, "
            f"{monitor['anomaly_medium_priority_count']} medium, "
            f"{monitor['anomaly_low_priority_count']} low); Status "
            f"{monitor['anomaly_review_status_counts']}."
        ),
        f"- Review limitation: {monitor['anomaly_limitation']}",
        "",
        "## Schweizer Referendum",
        "",
        f"- Kuratierte Umfragen: {swiss['poll_count']}.",
        f"- Polymarket snapshots: {swiss['snapshot_count']}.",
        f"- Bounded price-history rows: {swiss['history_rows']}.",
        f"- Latest Polymarket Yes: {swiss['latest_poly_yes_pct']:.1f} Prozent.",
        f"- Latest matched poll Yes: {swiss['latest_poll_yes_pct']:.1f} Prozent.",
        f"- Raw gap: {swiss['latest_raw_gap_pp']:+.1f} Prozentpunkte.",
        f"- Decided-voter gap: {swiss['latest_decided_gap_pp']:+.1f} Prozentpunkte.",
        "",
        "## Abbildungen",
        "",
    ]
    for figure in data["figures"]:
        rel = Path(
            figure.path.resolve().relative_to(markdown_output.parent.resolve())
            if figure.path.resolve().is_relative_to(markdown_output.parent.resolve())
            else Path(
                _relative_path(figure.path.resolve(), markdown_output.parent.resolve())
            )
        )
        lines.append(f"![{figure.caption}]({rel.as_posix()})")
        lines.append("")
        lines.append(f"*{figure.note}*")
        lines.append("")
    return "\n".join(lines)


def render_html(data: dict[str, Any], *, html_output: Path) -> str:
    """Render a standalone readable HTML companion report."""

    h1 = data["h1"]
    h2 = data["h2"]
    h3 = data["h3"]
    monitor = data["monitor"]
    swiss = data["swiss"]
    literature = data["literature"]
    source_review = data["source_review"]
    highlevel = data["thesis_highlevel"]
    project_highlevel = data["project_highlevel"]
    next_work = data["next_work"]
    execution_checklist = data["execution_checklist"]
    advisor_handoff = data["advisor_handoff"]
    submission_readiness = data["submission_readiness"]
    drafting_sequence = data["drafting_sequence"]
    bounded_chapter_draft = data["bounded_chapter_draft"]
    source_gated_drafting = data["source_gated_drafting"]
    worksheet_drafting_bridge = data["worksheet_drafting_bridge"]
    figures = "\n".join(
        _figure_html(figure, html_output=html_output)
        for figure in data["figures"]
    )
    literature_rows = "\n".join(
        "<tr>"
        f"<td><code>{escape(source['source_id'])}</code><br>"
        f"{escape(source['authors'])} ({escape(source['year'])}): "
        f"{escape(source['title'])}</td>"
        f"<td>{escape(source['role'])}</td>"
        f"<td>{escape(source['research_note'])}</td>"
        f"<td>{escape(source['status'])}</td>"
        "</tr>"
        for source in literature["sources"]
    )
    insight_rows = "\n".join(
        "<tr>"
        f"<td>{escape(row['bereich'])}</td>"
        f"<td>{escape(row['erkenntnis'])}</td>"
        f"<td>{escape(row['evidenz'])}</td>"
        f"<td>{escape(row['interpretation'])}</td>"
        f"<td>{escape(row['grenze'])}</td>"
        "</tr>"
        for row in _interpretation_rows(data)
    )
    method_decision_rows = "\n".join(
        "<tr>"
        f"<td>{escape(decision)}</td>"
        f"<td>{escape(reason)}</td>"
        f"<td>{escape(consequence)}</td>"
        "</tr>"
        for decision, reason, consequence in _method_decision_rows()
    )
    highlevel_rows = "\n".join(
        "<tr>"
        f"<td>{escape(row[0])}</td>"
        f"<td>{escape(row[1])}</td>"
        f"<td>{escape(row[2])}</td>"
        "</tr>"
        for row in highlevel["rows"]
    )
    next_work_rows = "\n".join(
        "<tr>"
        f"<td>{row['priority_order']}</td>"
        f"<td>{escape(row['workstream'])}</td>"
        f"<td>{escape(row['next_action'])}</td>"
        f"<td>{escape(row['guardrail'])}</td>"
        "</tr>"
        for row in next_work["rows"]
    )
    execution_rows = "\n".join(
        "<tr>"
        f"<td><code>{escape(row['task_id'])}</code></td>"
        f"<td>{escape(row['chapter_title'])}</td>"
        f"<td>{escape(row['draft_action_de'])}</td>"
        f"<td>{escape(row['done_when_de'])}</td>"
        f"<td>{escape(row['advisor_question_ids'])}</td>"
        "</tr>"
        for row in execution_checklist["rows"]
    )
    advisor_handoff_rows = "\n".join(
        "<tr>"
        f"<td>{row['package_order']}</td>"
        f"<td><code>{escape(row['path'])}</code></td>"
        f"<td>{escape(row['handoff_use_de'])}</td>"
        f"<td>{escape(row['advisor_decision_de'])}</td>"
        f"<td>{escape(row['boundary_de'])}</td>"
        "</tr>"
        for row in advisor_handoff["rows"]
    )
    project_highlevel_rows = "\n".join(
        "<tr>"
        f"<td>{escape(row['project_layer'])}</td>"
        f"<td>{escape(row['status_de'])}</td>"
        f"<td>{escape(row['current_decision_de'])}</td>"
        f"<td>{escape(row['next_gate_de'])}</td>"
        "</tr>"
        for row in project_highlevel["rows"]
    )
    readiness_rows = "\n".join(
        "<tr>"
        f"<td>{escape(row['gate_area'])}</td>"
        f"<td>{escape(row['current_status'])}</td>"
        f"<td>{escape(row['next_action_de'])}</td>"
        f"<td>{escape(row['blocker_or_limit_de'])}</td>"
        f"<td>{escape(row['thesis_use_de'])}</td>"
        "</tr>"
        for row in submission_readiness["rows"]
    )
    drafting_rows = "\n".join(
        "<tr>"
        f"<td>{row['priority_order']}</td>"
        f"<td>{escape(row['thesis_section'])}</td>"
        f"<td>{escape(row['draft_permission'])}</td>"
        f"<td>{escape(row['writing_action_de'])}</td>"
        f"<td>{escape(row['must_not_claim_de'])}</td>"
        "</tr>"
        for row in drafting_sequence["rows"]
    )
    bounded_chapter_rows = "\n".join(
        "<tr>"
        f"<td>{escape(row['thesis_area'])}</td>"
        f"<td><code>{escape(row['method_evidence_ids'])}</code></td>"
        f"<td><code>{escape(row['interpretation_evidence_ids'])}</code></td>"
        f"<td>{escape(row['literature_artifact_summary_de'])}</td>"
        f"<td>{escape(row['table_figure_de'])}</td>"
        f"<td>{escape(row['source_review_gate_summary_de'])}</td>"
        "</tr>"
        for row in bounded_chapter_draft["chapter_rows"]
    )
    source_gated_chapter_rows = "\n".join(
        "<tr>"
        f"<td>{escape(row['thesis_area'])}</td>"
        f"<td>{row['step_count']}: {escape(row['step_summary_de'])}</td>"
        f"<td>{escape(row['manual_review_summary_de'])}</td>"
        f"<td>{escape(row['table_figure_de'])}</td>"
        f"<td>{escape(row['status_de'])}</td>"
        "</tr>"
        for row in source_gated_drafting["chapter_rows"]
    )
    source_gated_step_rows = "\n".join(
        "<tr>"
        f"<td>{row['draft_sequence_order']}</td>"
        f"<td>{escape(row['thesis_area'])}</td>"
        f"<td>{escape(row['draft_section_de'])}</td>"
        f"<td>{escape(row['writer_action_de'])}</td>"
        f"<td>{escape(row['final_gate_short_de'])}</td>"
        "</tr>"
        for row in source_gated_drafting["step_rows"]
    )
    worksheet_bridge_rows = "\n".join(
        "<tr>"
        f"<td>{escape(row['thesis_area'])}</td>"
        f"<td>{escape(row['worksheet_summary_de'])}</td>"
        f"<td>{escape(row['drafting_summary_de'])}</td>"
        f"<td>{escape(row['table_figure_de'])}</td>"
        f"<td>{escape(row['gate_short_de'])}</td>"
        "</tr>"
        for row in worksheet_drafting_bridge["chapter_rows"]
    )
    h2_rows = "\n".join(
        f"<tr><td>{escape(row['event'])}</td><td>{row['change_pp']:+.1f} pp</td></tr>"
        for row in h2["primary_examples"]
    )
    swiss_source_rows = "\n".join(
        "<tr>"
        f"<td>{escape(row['source'])}</td>"
        f"<td>{escape(row['poll_id'])}</td>"
        f"<td>{row['poll_yes']:.1f}%</td>"
        f"<td>{row['raw_gap_pp']:+.1f} pp</td>"
        "</tr>"
        for row in swiss["latest_source_rows"]
    )
    return f"""<!doctype html>
<html lang="de">
<head>
  <meta charset="utf-8">
  <title>Dozentenbericht BA Thesis</title>
  <style>
    body {{ font-family: Calibri, Arial, sans-serif; margin: 36px auto; max-width: 1040px; color: #1d1d1f; line-height: 1.45; }}
    h1 {{ font-size: 30px; margin-bottom: 4px; color: #0b2545; }}
    h2 {{ margin-top: 34px; color: #2e74b5; border-bottom: 1px solid #d8e1ec; padding-bottom: 5px; }}
    h3 {{ color: #1f4d78; }}
    .subtitle {{ color: #555; font-size: 16px; margin-top: 0; }}
    .callout {{ background: #f4f6f9; border-left: 5px solid #2e74b5; padding: 14px 16px; margin: 18px 0; }}
    .grid {{ display: grid; grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 12px; margin: 16px 0; }}
    .metric {{ background: #f2f4f7; padding: 12px; border-radius: 6px; }}
    .metric strong {{ display: block; color: #0b2545; font-size: 20px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 12px 0 18px; }}
    th, td {{ border: 1px solid #ccd5df; padding: 8px 10px; vertical-align: top; }}
    th {{ background: #f2f4f7; color: #0b2545; text-align: left; }}
    figure {{ margin: 26px 0; }}
    figure img {{ max-width: 100%; border: 1px solid #e2e6ea; }}
    figcaption {{ color: #444; font-size: 14px; margin-top: 6px; }}
    .small {{ color: #666; font-size: 13px; }}
  </style>
</head>
<body>
  <h1>Dozentenbericht zur Bachelorarbeit</h1>
  <p class="subtitle">{escape(data['project']['working_title'])}</p>
  <p class="small">Erstellt aus lokalen deterministischen Artefakten: {escape(data['generated_at_utc'])}</p>
  <div class="callout"><strong>Kurzfazit.</strong> Das Projekt prueft informationelle Effizienz ueber drei reproduzierbare Analyseebenen: H1 Forecast-Qualitaet, H2 taegliche Event-Window-Reaktion und H3 Wallet-Tier-Timing. Der Monitor ist ein read-only Forschungsprototyp mit deterministischer Anomaly-Review-Queue; der Swiss-Referendum-Vergleich laeuft als separater Datensammlungs-Track.</div>

  <h2>Aufbau wie in der Bachelorarbeit</h2>
  <p>Der Bericht folgt der Logik einer Bachelorarbeit: Forschungsfrage, Theorie- und Literaturrahmen, Methodik, Datenbasis, empirische Ergebnisse, Diskussion der Grenzen und naechste Arbeitsschritte.</p>
  <ul>
    <li>H1 prueft Forecast-Qualitaet von Polymarket gegen traditionelle Probability-Forecasts und poll-derived Vergleichswerte.</li>
    <li>H2 prueft taegliche Reaktionen auf vorab kuratierte politische Ereignisse.</li>
    <li>H3 prueft aggregierte Wallet-Tier-Aktivitaet als Timing-Diagnostik.</li>
    <li>Monitor und Swiss-Referendum-Track bleiben read-only Forschungs- und Vergleichserweiterungen.</li>
  </ul>

  <h2>Highlevel-Projektstand</h2>
  <p>Der Review-Access bleibt pausiert. Der aktuelle Fortschritt liegt in der Thesis-Konsolidierung: Methoden, Interpretationen, Quellen, Tabellen und Figuren sind auf deterministische Artefakte gemappt.</p>
  <div class="grid">
    <div class="metric"><strong>{highlevel['core_tables']}</strong>Kern-Tabellen</div>
    <div class="metric"><strong>{highlevel['core_figures']}</strong>Kern-Figuren</div>
    <div class="metric"><strong>{highlevel['citation_packets']}</strong>Citation-Pakete</div>
    <div class="metric"><strong>{highlevel['chapter_rows']}</strong>Kapitelplan-Zeilen</div>
  </div>
  <p>Aktive Phase: {escape(highlevel['active_phase'])}. Das Thesis-Paket umfasst {highlevel['caption_rows']} Caption-Zeilen und wird aus <code>data/results/thesis_table_figure_captions.csv</code> gespeist.</p>
  <table><tr><th>Ebene</th><th>Stand</th><th>Konsequenz fuer die Thesis</th></tr>{highlevel_rows}</table>

  <h2>Dozentenpaket und Uebergabereihenfolge</h2>
  <p>Das Advisor-Handoff-Paket ordnet {advisor_handoff['row_count']} Dateien fuer die naechste Betreuung. Zuerst kommt <code>{escape(advisor_handoff['first_deliverable'])}</code>, danach folgen Word-Bericht, Absprache-Checklist, Submission Readiness Board, Drafting Sequence, Feedback-Log und die Arbeitsdateien fuer Kapitel und Source Review.</p>
  <table><tr><th>Reihenfolge</th><th>Datei</th><th>Verwendung</th><th>Entscheidung</th><th>Grenze</th></tr>{advisor_handoff_rows}</table>

  <h2>Projektmatrix fuer die naechste Abstimmung</h2>
  <p>Die Projektmatrix fasst {project_highlevel['row_count']} Ebenen als Status-, Entscheidungs- und Gate-Sicht zusammen. Sie zeigt explizit, dass Review-Access pausiert bleibt und Agenten nur dokumentierter Ausblick sind.</p>
  <table><tr><th>Projektteil</th><th>Status</th><th>Entscheidung</th><th>Naechstes Gate</th></tr>{project_highlevel_rows}</table>

  <h2>Submission Readiness und finale Gates</h2>
  <p>Das Submission Readiness Board trennt {submission_readiness['row_count']} Gates in Draft-Arbeit, finale Blocker und Future Work. Draft-ready Gates: {submission_readiness['draft_ready_count']}; final blockierte Gates: {submission_readiness['final_blocked_count']}. Source Review steht auf <code>{escape(submission_readiness['source_review_status'])}</code>, Swiss steht auf <code>{escape(submission_readiness['swiss_status'])}</code>, Agenten stehen auf <code>{escape(submission_readiness['agent_status'])}</code>.</p>
  <table><tr><th>Gate</th><th>Status</th><th>Naechste Aktion</th><th>Grenze</th><th>Thesis-Nutzung</th></tr>{readiness_rows}</table>

  <h2>Schreibsequenz fuer den naechsten Entwurf</h2>
  <p>Die Drafting Sequence ordnet {drafting_sequence['row_count']} Schritte vom Quellenreview bis zur finalen QA. Erste Sequenz ist <code>{escape(drafting_sequence['first_step'])}</code>, letzte Sequenz ist <code>{escape(drafting_sequence['final_step'])}</code>. Bounded write-now: {drafting_sequence['bounded_write_now_count']}; final blockiert: {drafting_sequence['final_blocked_count']}; Future-work-only: {drafting_sequence['future_work_only_count']}.</p>
  <table><tr><th>Prioritaet</th><th>Thesis-Abschnitt</th><th>Erlaubnis</th><th>Schreibaktion</th><th>Nicht behaupten</th></tr>{drafting_rows}</table>

  <h2>Bounded H1-H2-H3 Kapitelentwurf</h2>
  <p>Der H1-H2-H3 Bounded Chapter Draft liefert {bounded_chapter_draft['row_count']} geordnete Prosa-Bausteine: {bounded_chapter_draft['rows_per_chapter']} je H1, H2 und H3. Bounded-draft-ready: {bounded_chapter_draft['bounded_ready_count']}; final-submission-ready: {bounded_chapter_draft['final_ready_count']}.</p>
  <p>Jede empirische Methode und jede Interpretation bleibt an Evidence-IDs, Literatur-IDs, deterministische Artefakte, kuratierte Tabellen/Figuren, Limitationen und Source-Review-Gates gebunden.</p>
  <table><tr><th>Kapitel</th><th>Methoden</th><th>Interpretationen</th><th>Literatur/Artefakte</th><th>Tabelle/Figur</th><th>Gate</th></tr>{bounded_chapter_rows}</table>
  <p class="small">Schreibvorlage, kein finaler Zitations- oder Abgabeclaim: keine neuen Kennzahlen, keine Rohartefakt-Dumps, keine Quellenstatus-Hochstufung und keine Runtime-Agenten.</p>

  <h2>Source-Gated H1-H2-H3 Drafting Sequence</h2>
  <p>Der source-gated Thesis-Drafting-Pass bringt den H1-H2-H3-Kern in {source_gated_drafting['row_count']} paragraphenweise Schreibschritte: {source_gated_drafting['rows_per_chapter']} je Kapitel. Bounded-draft-ready: {source_gated_drafting['bounded_ready_count']}; final-submission-ready: {source_gated_drafting['final_ready_count']}. Manual Source Review: {source_gated_drafting['manual_rows_linked']} Rows verlinkt, {source_gated_drafting['manual_pending_rows']} pending, {source_gated_drafting['manual_final_ready_rows']} final-ready.</p>
  <p>Fuer den Dozenten ist das die konkrete Schreibreihenfolge nach dem Bounded Chapter Draft: Methode/Resultat setzen, Interpretation und Limitation setzen, Tabelle/Figur einbauen, Manual Source Review ausfuehren, Finalgate und Future-Agent-Grenze sichtbar lassen.</p>
  <p>{escape(source_gated_drafting['review_control_de'])}</p>
  <table><tr><th>Kapitel</th><th>Schritte</th><th>Manual Source Review</th><th>Tabelle/Figur</th><th>Status</th></tr>{source_gated_chapter_rows}</table>
  <table><tr><th>Ordnung</th><th>Kapitel</th><th>Schreibschritt</th><th>Writer Action</th><th>Finalgate</th></tr>{source_gated_step_rows}</table>
  <p class="small">Auch diese Sequenz ist kein finaler Zitations- oder Abgabeclaim: Page-/Section-Note, Claim-Support, Blocked-Wording und Citation-Use bleiben manuelle Gates; Agenten bleiben documentation-only Future Work.</p>

  <h2>Worksheet-to-Drafting Bridge fuer H1-H2-H3</h2>
  <p>Die Worksheet-Drafting-Bridge verbindet {worksheet_drafting_bridge['worksheet_rows']} manuelle H1-H2-H3 Worksheet-Zeilen mit {worksheet_drafting_bridge['drafting_steps']} source-gated Schreibschritten. Method rows: {worksheet_drafting_bridge['method_rows']}; Interpretation rows: {worksheet_drafting_bridge['interpretation_rows']}; Source/artifact gaps: {worksheet_drafting_bridge['source_artifact_gap_rows']}; final-release rows: {worksheet_drafting_bridge['final_release_ready_rows']}.</p>
  <p>{escape(worksheet_drafting_bridge['source_artifact_rule_de'])}</p>
  <p>Damit sieht der Dozent in einer kleinen Bruecke, welche Quellen- und Artefaktpflicht vor jedem Absatz gilt und warum der Haupttext bei T2/F1, T3/F2 und T4/F3 bleibt.</p>
  <table><tr><th>Kapitel</th><th>Worksheets</th><th>Drafting</th><th>Tabelle/Figur</th><th>Gate</th></tr>{worksheet_bridge_rows}</table>
  <p class="small">{escape(worksheet_drafting_bridge['future_agent_boundary_de'])}</p>

  <h2>Naechste Arbeitsschritte</h2>
  <p>Der Next-Work-Plan ordnet {next_work['row_count']} Workstreams. Erste Prioritaet ist <code>{escape(next_work['first_workstream'])}</code>, letzte QA-Prioritaet ist <code>{escape(next_work['final_workstream'])}</code>.</p>
  <table><tr><th>Prioritaet</th><th>Workstream</th><th>Naechste Aktion</th><th>Guardrail</th></tr>{next_work_rows}</table>

  <h2>Kapitelweise Umsetzungscheckliste</h2>
  <p>Die Execution-Checkliste uebersetzt die Highlevel-View in {execution_checklist['row_count']} Kapitelaufgaben. Erste Aufgabe ist <code>{escape(execution_checklist['first_task'])}</code>, letzte Aufgabe ist <code>{escape(execution_checklist['final_task'])}</code>. Sie ist kein neues empirisches Ergebnis, sondern eine Schreib- und Abnahmelogik fuer den naechsten Entwurf.</p>
  <table><tr><th>Task</th><th>Kapitel</th><th>Schreibaktion</th><th>Fertig wenn</th><th>Advisor-Fragen</th></tr>{execution_rows}</table>

  <h2>Forschungsfrage und Design</h2>
  <p>Die Leitfrage lautet, inwiefern Polymarket-Preise Informationen waehrend politischer Ereignisse abbilden, anders als traditionelle Prognosequellen reagieren und ob aggregierte Wallet-Aktivitaet fruehe Timing-Signale zeigt. Informationelle Effizienz wird deshalb nicht direkt behauptet, sondern ueber reproduzierbare Proxies operationalisiert.</p>
  <p>Alle Kennzahlen stammen aus Python-Artefakten unter <code>data/results</code>. RCP bleibt ausgeschlossen, solange keine dokumentierte Probability-Transformation existiert; Granger-Outputs werden nicht kausal interpretiert; Monitor- und Live-Daten bleiben read-only und bounded.</p>

  <h2>Wissenschaftlicher Quellenrahmen</h2>
  <p>Der lokale Literaturindex umfasst {literature['source_count']} Quellen; fuer diesen Bericht werden {literature['selected_source_count']} wissenschaftlich relevante Kernquellen als Rahmen verwendet. Statusverteilung: {escape(literature['status_counts_text'])}.</p>
  <p>Das Source-Review-Worksheet enthaelt {source_review['worksheet_rows']} manuelle Review-Zeilen, davon {source_review['priority_1_rows']} Priority-1-Methodenquellen und {source_review['blocked_rows']} blockierte oder Future-Work-Quelle. Alle Reviewer-Entscheide bleiben pending.</p>
  <table><tr><th>Quelle</th><th>Rolle in der Arbeit</th><th>Beitrag zur Interpretation</th><th>Status</th></tr>{literature_rows}</table>
  <p class="small">{escape(literature['citation_boundary'])}</p>

  <h2>Zentrale Erkenntnisse, Begruendung und Interpretation</h2>
  <p>Die Resultate werden nicht als Rohzahlen stehen gelassen. Jede zentrale Erkenntnis wird mit Evidenz, vorsichtiger Interpretation und Grenze ausgewiesen.</p>
  <table><tr><th>Bereich</th><th>Erkenntnis</th><th>Evidenz</th><th>Interpretation</th><th>Grenze</th></tr>{insight_rows}</table>

  <h2>Warum dieses Vorgehen methodisch sinnvoll ist</h2>
  <table><tr><th>Entscheidung</th><th>Begruendung</th><th>Konsequenz</th></tr>{method_decision_rows}</table>

  <h2>Projektstruktur</h2>
  <div class="grid">
    <div class="metric"><strong>{data['project']['database']['table_count']}</strong>Datenbanktabellen</div>
    <div class="metric"><strong>{data['project']['folder_inventory'].get('data/results', 0)}</strong>Result-Artefakte</div>
    <div class="metric"><strong>{data['project']['folder_inventory'].get('operations/analysis', 0)}</strong>Analyse-Module</div>
    <div class="metric"><strong>{data['project']['folder_inventory'].get('tests', 0)}</strong>Testdateien</div>
  </div>
  <p>Letzter Teststatus: <code>{escape(data['project']['test_summary'])}</code>.</p>

  <h2>H1 - Forecast-Qualitaet</h2>
  <table><tr><th>Aspekt</th><th>Wert</th></tr>
    <tr><td>Beobachtungen</td><td>{h1['observation_count']}</td></tr>
    <tr><td>Polymarket Mean Brier</td><td>{h1['brier_polymarket']:.4f}</td></tr>
    <tr><td>FiveThirtyEight Mean Brier</td><td>{h1['brier_fivethirtyeight']:.4f}</td></tr>
    <tr><td>DM p-Wert Polymarket vs FiveThirtyEight</td><td>{h1['dm_polymarket_vs_538']:.3g}</td></tr>
    <tr><td>Polymarket niedrigerer Tagesverlust</td><td>{h1['pm_better_vs_538_count']} von {h1['pm_vs_538_count']} Tagen ({h1['pm_better_vs_538_share'] * 100:.1f}%)</td></tr>
    <tr><td>Mittlerer Verlustvorteil vs FiveThirtyEight</td><td>{h1['mean_loss_advantage_vs_538']:.4f}</td></tr>
    <tr><td>H1-Synthesis</td><td>{h1['synthesis_aggregate_support_count']} von {h1['synthesis_evidence_row_count']} Vergleichszeilen stuetzen Polymarket im mittleren Brier; {h1['synthesis_majority_support_count']} von {h1['synthesis_evidence_row_count']} zeigen eine Mehrheit niedrigerer Einzelfallverluste; breiter Viele-Faelle-Beweis {h1['synthesis_broad_support_count']} von {h1['synthesis_evidence_row_count']}.</td></tr>
    <tr><td>H1-Claim-Evidence-Audit</td><td>{h1['claim_audit_support_row_count']} von {h1['claim_audit_row_count']} Audit-Zeilen stuetzen Polymarket begrenzt; {h1['claim_audit_contradiction_row_count']} widerspricht dem starken Claim; direkt pollbezogen {h1['claim_audit_direct_poll_support_row_count']} von {h1['claim_audit_direct_poll_row_count']} stuetzend; breiter User-Claim belegt {h1['claim_audit_broad_user_claim_proven']}.</td></tr>
    <tr><td>H1-Poll-Comparison-Result</td><td>Primaerer &lt;=90-Tage-Low/Middle-Poll-Distanz-Scope: Polymarket {h1['poll_result_primary_pm_count']} von {h1['poll_result_primary_row_count']} State-Date-Zeilen ({h1['poll_result_primary_pm_share'] * 100:.1f}%), poll-derived {h1['poll_result_primary_poll_count']}; State-Ebene Polymarket {h1['poll_result_primary_pm_state_count']} von {h1['poll_result_primary_state_count']}, exakter einseitiger p-Wert {h1['poll_result_primary_p_value']:.4f}. Direkt pollbezogen {h1['poll_result_direct_poll_support_count']} von {h1['poll_result_direct_poll_row_count']} Audit-Zeilen stuetzend; Vollpanel-Gegenbeleg poll-derived {h1['poll_result_full_panel_poll_count']} von {h1['poll_result_full_panel_row_count']}; Status {h1['poll_result_goal_status']}.</td></tr>
    <tr><td>H1-Poll-Claim-Readiness</td><td>{h1['poll_claim_supported_bounded_count']} von {h1['poll_claim_row_count']} Claim-Zeilen stuetzen den bounded &lt;=90-Tage Low/Middle-Poll-Distanz-Scope; {h1['poll_claim_counterexample_count']} Gegenbeispiel-Scopes und {h1['poll_claim_mixed_mean_count']} Mean-Loss-Stuetze-ohne-Mehrheit-Zeilen bleiben als Grenzen. Bounded Scope: Polymarket {h1['poll_claim_primary_pm_count']} von {h1['poll_claim_primary_count']} State-Date-Zeilen ({h1['poll_claim_primary_pm_share'] * 100:.1f}%) und {h1['poll_claim_state_month_pm_count']} von {h1['poll_claim_state_month_count']} State-Month-Einheiten, exact p={h1['poll_claim_state_month_p_value']:.2g}, 95-Prozent-Untergrenze {h1['poll_claim_state_month_ci_low']:.3f}. Bounded Claim supported {h1['poll_claim_bounded_supported']}; breiter Claim belegt {h1['poll_claim_broad_proven']}; Status {h1['poll_claim_goal_status']}.</td></tr>
    <tr><td>H1-Poll-Scope-Frontier</td><td>{h1['poll_frontier_robust_scope_count']} von {h1['poll_frontier_row_count']} Horizont-x-Poll-Distanz-Scopes erfuellen die robuste Regel. Groesster robuster Scope: {h1['poll_frontier_largest_horizon']} + {h1['poll_frontier_largest_tier']}, Polymarket {h1['poll_frontier_largest_pm_count']} von {h1['poll_frontier_largest_row_count']} State-Date-Zeilen ({h1['poll_frontier_largest_pm_share'] * 100:.1f}%), {h1['poll_frontier_largest_state_month_pm_count']} von {h1['poll_frontier_largest_state_month_count']} State-Month-Einheiten, exact p={h1['poll_frontier_largest_state_month_p_value']:.3g}. Staerkster Scope {h1['poll_frontier_strongest_scope_id']}: {h1['poll_frontier_strongest_row_count']} Zeilen, p={h1['poll_frontier_strongest_p_value']:.2g}. &lt;=90 Tage alle Distanzen: Polymarket {h1['poll_frontier_lte_90_all_pm_count']} von {h1['poll_frontier_lte_90_all_row_count']} Zeilen ({h1['poll_frontier_lte_90_all_pm_share'] * 100:.1f}%), State-Month p={h1['poll_frontier_lte_90_all_state_month_p_value']:.3g}; Vollpanel-Gegenbeleg poll-derived {h1['poll_frontier_full_panel_poll_count']} von {h1['poll_frontier_full_panel_row_count']}; Status {h1['poll_frontier_goal_status']}.</td></tr>
    <tr><td>H1-Poll-Decision-Matrix</td><td>{h1['poll_decision_robust_yes_count']} von {h1['poll_decision_row_count']} Entscheidungszeilen sind robuste bounded-Yes-Zeilen; {h1['poll_decision_mixed_mean_count']} Mean-Loss-Stuetze-ohne-Mehrheit-Zeilen und {h1['poll_decision_counterexample_count']} Gegenbelege bleiben als Grenzen. Groesster robuster Scope: Polymarket {h1['poll_decision_largest_pm_count']} von {h1['poll_decision_largest_row_count']} State-Date-Zeilen ({h1['poll_decision_largest_pm_share'] * 100:.1f}%), {h1['poll_decision_largest_state_month_pm_count']} von {h1['poll_decision_largest_state_month_count']} State-Month-Einheiten, p={h1['poll_decision_largest_p_value']:.4f}. Kalibrierungskontext: {h1['poll_decision_calibration_aggregate_count']} von {h1['poll_decision_calibration_pairwise_count']} Pairwise-Reihen stuetzen Polymarket im mittleren Brier; {h1['poll_decision_calibration_majority_count']} von {h1['poll_decision_calibration_pairwise_count']} auch per Fallmehrheit. Bounded ready {h1['poll_decision_bounded_ready']}; breiter Claim {h1['poll_decision_broad_proven']}; Status {h1['poll_decision_goal_status']}.</td></tr>
    <tr><td>H1-Robust-Poll-Scope-Quality</td><td>{h1['robust_quality_forecast_row_count']} Forecast-Zeilen aus {h1['robust_quality_case_count']} State-Date-Faellen und {h1['robust_quality_scope_count']} robusten Poll-Scopes. Groesster robuster Scope: Polymarket {h1['robust_quality_largest_pm_count']} von {h1['robust_quality_largest_case_count']} Zeilen ({h1['robust_quality_largest_pm_share'] * 100:.1f}%), Mean Brier {h1['robust_quality_largest_pm_brier']:.4f} vs poll-derived {h1['robust_quality_largest_poll_brier']:.4f}, ECE {h1['robust_quality_largest_pm_ece']:.4f} vs {h1['robust_quality_largest_poll_ece']:.4f}, Probability-Separation {h1['robust_quality_largest_pm_separation']:.4f} vs {h1['robust_quality_largest_poll_separation']:.4f}. Staerkster robuster Scope: Polymarket {h1['robust_quality_strongest_pm_count']} von {h1['robust_quality_strongest_case_count']} Zeilen ({h1['robust_quality_strongest_pm_share'] * 100:.1f}%), Mean Brier {h1['robust_quality_strongest_pm_brier']:.4f} vs {h1['robust_quality_strongest_poll_brier']:.4f}, ECE {h1['robust_quality_strongest_pm_ece']:.4f} vs {h1['robust_quality_strongest_poll_ece']:.4f}; alle Outcomes dort positiv, Separation nicht definiert. Breiter Claim belegt {h1['robust_quality_broad_claim_proven']}.</td></tr>
    <tr><td>H1-Robust-Poll-Scope-Unit-Quality</td><td>{h1['robust_unit_summary_row_count']} Aggregationszeilen ueber robuste Poll-Scopes. Groesster robuster Scope: State {h1['robust_unit_largest_state_pm_count']} von {h1['robust_unit_largest_state_count']} (p={h1['robust_unit_largest_state_p_value']:.3g}), State-Month {h1['robust_unit_largest_state_month_pm_count']} von {h1['robust_unit_largest_state_month_count']} (poll-derived {h1['robust_unit_largest_state_month_poll_count']}, p={h1['robust_unit_largest_state_month_p_value']:.4f}, 95-Prozent-Untergrenze {h1['robust_unit_largest_state_month_ci_low']:.3f}), State-Horizon {h1['robust_unit_largest_state_horizon_pm_count']} von {h1['robust_unit_largest_state_horizon_count']} (p={h1['robust_unit_largest_state_horizon_p_value']:.3g}). Staerkster robuster Scope: State {h1['robust_unit_strongest_state_pm_count']} von {h1['robust_unit_strongest_state_count']} (p={h1['robust_unit_strongest_state_p_value']:.3g}), State-Month {h1['robust_unit_strongest_state_month_pm_count']} von {h1['robust_unit_strongest_state_month_count']} (p={h1['robust_unit_strongest_state_month_p_value']:.2g}, 95-Prozent-Untergrenze {h1['robust_unit_strongest_state_month_ci_low']:.3f}), State-Horizon {h1['robust_unit_strongest_state_horizon_pm_count']} von {h1['robust_unit_strongest_state_horizon_count']} (p={h1['robust_unit_strongest_state_horizon_p_value']:.2g}). Medianer State-Month-Brier-Vorteil {h1['robust_unit_largest_state_month_median_advantage']:.4f} im groessten und {h1['robust_unit_strongest_state_month_median_advantage']:.4f} im staerksten Scope. Breiter Claim belegt {h1['robust_unit_broad_claim_proven']}.</td></tr>
    <tr><td>H1-Poll-Comparison-Unit-Robustness</td><td>Primaerer Scope nach Aggregation: Polymarket {h1['poll_unit_state_pm_count']} von {h1['poll_unit_state_count']} States, {h1['poll_unit_state_month_pm_count']} von {h1['poll_unit_state_month_count']} State-Month-Einheiten und {h1['poll_unit_state_horizon_pm_count']} von {h1['poll_unit_state_horizon_count']} State-Horizon-Einheiten; State-Month exact p={h1['poll_unit_state_month_p_value']:.2g}, 95-Prozent-Untergrenze {h1['poll_unit_state_month_ci_low']:.3f}. Full-Panel-State-Month-Gegenbeleg: poll-derived {h1['poll_unit_full_panel_state_month_poll_count']} von {h1['poll_unit_full_panel_state_month_count']}; Late-High-Distance-State-Month-Gegenbeleg: poll-derived {h1['poll_unit_late_high_state_month_poll_count']} von {h1['poll_unit_late_high_state_month_count']}, exact p={h1['poll_unit_late_high_state_month_poll_p_value']:.4f}; Status {h1['poll_unit_goal_status']}.</td></tr>
    <tr><td>H1-Direct-Poll-Loss-Decomposition</td><td>Direkte Poll-Transform-Vergleiche: Mean Brier Polymarket {h1['direct_poll_loss_pm_brier']:.4f} vs poll-derived {h1['direct_poll_loss_poll_brier']:.4f}; Polymarket niedrigerer Verlust in {h1['direct_poll_loss_pm_count']} von {h1['direct_poll_loss_case_count']} Source-State-Faellen, poll-derived in {h1['direct_poll_loss_poll_count']}. Polymarket-Gewinnfaelle haben mittleren Brier-Vorteil {h1['direct_poll_loss_pm_win_mean_advantage']:.4f}, poll-derived Gewinnfaelle {h1['direct_poll_loss_poll_win_mean_advantage']:.4f}; Total-Margin-Ratio {h1['direct_poll_loss_margin_ratio']:.1f}. Fallmehrheit belegt: {h1['direct_poll_loss_case_majority_supports_pm']}.</td></tr>
    <tr><td>H1-Direct-Poll-State-Cluster</td><td>State-Cluster-Diagnostik ueber {h1['direct_poll_state_cluster_state_count']} States: gleichgewichteter mittlerer Verlustvorteil {h1['direct_poll_state_cluster_mean_advantage']:.4f}, Bootstrap-95%-Intervall {h1['direct_poll_state_cluster_bootstrap_ci_low']:.4f} bis {h1['direct_poll_state_cluster_bootstrap_ci_high']:.4f}, Sign-Flip-p={h1['direct_poll_state_cluster_sign_flip_p']:.4f}. State-Mehrheit: Polymarket {h1['direct_poll_state_cluster_pm_state_count']} States, poll-derived {h1['direct_poll_state_cluster_poll_state_count']}; Polymarket-State-Mehrheit belegt: {h1['direct_poll_state_cluster_majority_supports_pm']}.</td></tr>
    <tr><td>H1-Direct-Poll-Outlier-Robustness</td><td>Outlier-Diagnostik ueber {h1['direct_poll_outlier_state_count']} State-Cluster: voller Mean {h1['direct_poll_outlier_full_mean_advantage']:.4f}; alle Leave-one-state-out Means positiv {h1['direct_poll_outlier_leave_one_all_positive']}, Minimum {h1['direct_poll_outlier_min_leave_one_mean']:.4f} ohne {h1['direct_poll_outlier_most_influential_state']}. Top-positive Exclusion: Mean bleibt bis {h1['direct_poll_outlier_top_k_positive']} entfernte States positiv und kippt bei {h1['direct_poll_outlier_first_nonpositive_k']} entfernten States auf {h1['direct_poll_outlier_first_nonpositive_mean']:.4f}; groesster positiver State {h1['direct_poll_outlier_largest_positive_state']} ({h1['direct_poll_outlier_largest_positive_advantage']:.4f}). Status {h1['direct_poll_outlier_goal_status']}.</td></tr>
    <tr><td>H1-State-Source-Konsens</td><td>{h1['state_source_consensus_case_count']} Source-State-Vergleiche ueber {h1['state_source_consensus_state_count']} States; Polymarket niedrigerer Verlust in {h1['state_source_consensus_pm_case_count']} Source-State-Faellen, Comparatoren in {h1['state_source_consensus_comparator_case_count']}. All-Source-State-Konsens: Polymarket {h1['state_source_consensus_pm_state_count']} States, Comparatoren {h1['state_source_consensus_comparator_state_count']}, Ties {h1['state_source_consensus_tie_state_count']}. Zwei direkte Poll-Transform-Quellen: Polymarket {h1['state_source_consensus_direct_two_pm_state_count']} von {h1['state_source_consensus_direct_two_state_count']} States.</td></tr>
    <tr><td>H1-Competitive-State-Diagnose</td><td>Niedrigste Comparator-Distanz-Terzile: Polymarket {h1['competitive_state_all_low_pm_count']} von {h1['competitive_state_all_low_case_count']} All-Source-Faellen und {h1['competitive_state_direct_low_pm_count']} von {h1['competitive_state_direct_low_case_count']} direkten Poll-Transform-Faellen; hoechste Distanz-Terzile: Polymarket {h1['competitive_state_all_high_pm_count']} von {h1['competitive_state_all_high_case_count']}, Comparatoren {h1['competitive_state_all_high_comparator_count']} von {h1['competitive_state_all_high_case_count']}. Begrenzte Competitive-State-Ausnahme, kein breiter Viele-Faelle-Beweis.</td></tr>
    <tr><td>H1-State-Date-Competitiveness-x-Horizon</td><td>&lt;=90 Tage und Low/Middle-Poll-Distanz: Polymarket {h1['panel_comp_late_non_safe_pm_count']} von {h1['panel_comp_late_non_safe_row_count']} State-Date-Zeilen und {h1['panel_comp_late_non_safe_state_support_count']} von {h1['panel_comp_late_non_safe_state_count']} States; spaete High-Distance-Zeilen: Polymarket {h1['panel_comp_late_high_pm_count']} von {h1['panel_comp_late_high_row_count']}, poll-derived {h1['panel_comp_late_high_poll_count']} von {h1['panel_comp_late_high_row_count']}. Starker spaeter Competitive-Poll-Befund, aber kein unabhaengiger Viele-Wahlen-Beweis.</td></tr>
    <tr><td>H1-State-Level-Signifikanzdiagnose</td><td>Spaete Low/Middle-Poll-Distanz: Polymarket {h1['state_sign_late_non_safe_pm_state_count']} von {h1['state_sign_late_non_safe_state_count']} States; exakter einseitiger Binomial-p-Wert {h1['state_sign_late_non_safe_p_value']:.4f}; exakte 95-Prozent-Untergrenze {h1['state_sign_late_non_safe_ci_low']:.3f}. Spaete High-Distance-States: poll-derived {h1['state_sign_late_high_poll_state_count']} von {h1['state_sign_late_high_state_count']} States.</td></tr>
    <tr><td>H1-Kalibrierungsdiagnostik</td><td>{h1['calibration_forecast_case_rows']} Forecast-Case-Zeilen aus {h1['calibration_forecast_source_count']} Quellen und {h1['calibration_pairwise_count']} Pairwise-Reihen; {h1['calibration_aggregate_support_count']} von {h1['calibration_pairwise_count']} zeigen niedrigeren mittleren Polymarket-Brier; {h1['calibration_majority_support_count']} von {h1['calibration_pairwise_count']} auch eine Mehrheit niedrigerer Einzelfallverluste; breiter Viele-Faelle-Beweis {h1['calibration_broad_support_count']} von {h1['calibration_pairwise_count']}.</td></tr>
    <tr><td>50-State-Kalibrierung</td><td>Polymarket Mean Brier {h1['calibration_pm_state_brier']:.4f} und Fixed-Bin-ECE {h1['calibration_pm_state_ece']:.4f}; Rieke ECE {h1['calibration_rieke_state_ece']:.4f}; 270toWin/JHK ECE {h1['calibration_270_state_ece']:.4f}. Forecast-Qualitaets-, aber kein klarer Kalibrierungssieg.</td></tr>
    <tr><td>Final-Snapshot-Erweiterung</td><td>{h1['final_snapshot_pm_lower_loss_count']} von {h1['final_snapshot_case_count']} geloesten 2024-Outcomes mit niedrigerem Polymarket-Verlust; Mean Brier {h1['final_snapshot_mean_pm_brier']:.4f} vs {h1['final_snapshot_mean_traditional_brier']:.4f}</td></tr>
    <tr><td>State-Poll-Snapshot-Erweiterung</td><td>{h1['state_poll_snapshot_pm_lower_loss_count']} von {h1['state_poll_snapshot_case_count']} geloesten State-Outcomes mit niedrigerem Polymarket-Verlust; Mean Brier {h1['state_poll_snapshot_mean_pm_brier']:.4f} vs {h1['state_poll_snapshot_mean_poll_brier']:.4f}</td></tr>
    <tr><td>270toWin-Polling-Average-Erweiterung</td><td>{h1['two_seventy_poll_average_case_count']} gematchte State-Outcomes; Polymarket niedrigerer Verlust in {h1['two_seventy_poll_average_pm_lower_loss_count']} Faellen, poll-derived in {h1['two_seventy_poll_average_poll_lower_loss_count']}; Mean Brier {h1['two_seventy_poll_average_mean_pm_brier']:.4f} vs {h1['two_seventy_poll_average_mean_poll_brier']:.4f}</td></tr>
    <tr><td>Popular-Vote-Erweiterung</td><td>{h1['popular_vote_case_count']} nationale Tageszeilen fuer Trump popular vote; Polymarket niedrigerer Verlust in {h1['popular_vote_pm_lower_loss_count']} Zeilen, poll-derived in {h1['popular_vote_poll_lower_loss_count']}; Mean Brier {h1['popular_vote_mean_pm_brier']:.4f} vs {h1['popular_vote_mean_poll_brier']:.4f}. Gegenbeleg zum starken Claim.</td></tr>
    <tr><td>Margin-Threshold-Readiness</td><td>{h1['margin_threshold_candidate_count']} Trump-State-Margin-Maerkte geprueft; {h1['margin_threshold_with_538_poll_count']} mit 538-State-Poll-Average-Zeilen, {h1['margin_threshold_with_clob_overlap_count']} mit CLOB-Historie im bewahrten 538-Fenster, {h1['margin_threshold_compatible_count']} neue H1-Brier-Faelle. {h1['margin_threshold_no_overlap_count']} sind durch fehlende zeitliche Ueberlappung blockiert, {h1['margin_threshold_missing_poll_count']} durch fehlende 538-State-Polls.</td></tr>
    <tr><td>State-Date-Poll-Panel</td><td>{h1['state_poll_panel_case_count']} gematchte State-Date-Zeilen ueber {h1['state_poll_panel_state_count']} States und {h1['state_poll_panel_date_count']} Daten; Polymarket niedrigerer Verlust in {h1['state_poll_panel_pm_lower_loss_count']} Zeilen, poll-derived niedrigerer Verlust in {h1['state_poll_panel_poll_lower_loss_count']}; Mean Brier {h1['state_poll_panel_mean_pm_brier']:.4f} vs {h1['state_poll_panel_mean_poll_brier']:.4f}.</td></tr>
    <tr><td>Temporal-Diagnose State-Date-Panel</td><td>Polymarket-stuetzende Monate {h1['state_poll_temporal_support_months']}: {h1['state_poll_temporal_support_pm_lower_loss_count']} von {h1['state_poll_temporal_support_row_count']} Zeilen mit niedrigerem Polymarket-Verlust ueber {h1['state_poll_temporal_support_state_count']} States; poll-derived niedrigerer Verlust in {h1['state_poll_temporal_support_poll_lower_loss_count']}; Mean Brier {h1['state_poll_temporal_support_mean_pm_brier']:.4f} vs {h1['state_poll_temporal_support_mean_poll_brier']:.4f}.</td></tr>
    <tr><td>Forecast-Horizon-Diagnose</td><td>&lt;=90-Tage-Fenster ({h1['state_poll_horizon_near_bins']}): {h1['state_poll_horizon_near_pm_lower_loss_count']} von {h1['state_poll_horizon_near_row_count']} Zeilen mit niedrigerem Polymarket-Verlust ueber {h1['state_poll_horizon_near_state_count']} States; poll-derived niedrigerer Verlust in {h1['state_poll_horizon_near_poll_lower_loss_count']}; Mean Brier {h1['state_poll_horizon_near_mean_pm_brier']:.4f} vs {h1['state_poll_horizon_near_mean_poll_brier']:.4f}.</td></tr>
    <tr><td>State-Level-Horizon-Diagnose</td><td>Im &lt;=90-Tage-Fenster stuetzt Polymarket {h1['state_poll_horizon_state_pm_mean_support_count']} von {h1['state_poll_horizon_state_count']} States nach mittlerem Brier und {h1['state_poll_horizon_state_pm_majority_support_count']} von {h1['state_poll_horizon_state_count']} States nach Mehrheit niedrigerer Tagesverluste; {h1['state_poll_horizon_state_poll_support_count']} States stuetzen Polymarket nicht.</td></tr>
    <tr><td>&lt;=90-Day Score Quality</td><td>{h1['state_poll_near_quality_forecast_row_count']} Forecast-Zeilen aus {h1['state_poll_near_quality_case_count']} State-Date-Faellen und zwei Quellen: Polymarket Mean Brier {h1['state_poll_near_quality_pm_mean_brier']:.4f} vs poll-derived {h1['state_poll_near_quality_poll_mean_brier']:.4f}; Fixed-Bin-ECE {h1['state_poll_near_quality_pm_ece']:.4f} vs {h1['state_poll_near_quality_poll_ece']:.4f}; Probability-Separation {h1['state_poll_near_quality_pm_separation']:.4f} vs {h1['state_poll_near_quality_poll_separation']:.4f}.</td></tr>
    <tr><td>Poll-Transform-Sensitivitaet</td><td>MAE {h1['state_poll_sensitivity_min_mae']:.1f} bis {h1['state_poll_sensitivity_max_mae']:.1f} Prozentpunkte; Polymarket bleibt in allen {h1['state_poll_sensitivity_row_count']} Parameterzeilen im mittleren Brier niedriger; Lower-Loss-Spanne {h1['state_poll_sensitivity_min_pm_lower_loss_count']} bis {h1['state_poll_sensitivity_max_pm_lower_loss_count']} von {h1['state_poll_snapshot_case_count']} State-Outcomes.</td></tr>
    <tr><td>State-Poll-Coverage-Audit</td><td>{h1['state_poll_coverage_state_count']} US-States geprueft; {h1['state_poll_coverage_polymarket_market_count']} mit Polymarket-State-Markt; {h1['state_poll_coverage_valid_pair_count']} valide H1-Brier-Paare; {h1['state_poll_coverage_missing_poll_count']} wegen fehlender 538-Snapshot-Pollwerte ausgeschlossen.</td></tr>
    <tr><td>Rieke-50-State-Erweiterung</td><td>{h1['rieke_state_case_count']} geloeste State-Outcomes gegen Rieke poll-based model; Mean Brier {h1['rieke_state_mean_pm_brier']:.4f} vs {h1['rieke_state_mean_rieke_brier']:.4f}; Polymarket niedrigerer Einzelfallverlust in {h1['rieke_state_pm_lower_loss_count']} von {h1['rieke_state_case_count']}, Rieke in {h1['rieke_state_rieke_lower_loss_count']} von {h1['rieke_state_case_count']}.</td></tr>
    <tr><td>270toWin/JHK-50-State-Erweiterung</td><td>{h1['two_seventy_state_case_count']} geloeste State-Outcomes gegen 270toWin/JHK; {h1['two_seventy_state_exact_case_count']} exakt ausgewiesene Wahrscheinlichkeiten und {h1['two_seventy_state_censored_case_count']} zensierte &gt;99.9-Prozent-Boundary-Werte; Mean Brier {h1['two_seventy_state_mean_pm_brier']:.4f} vs {h1['two_seventy_state_mean_270_brier']:.4f}; Polymarket niedrigerer Einzelfallverlust in {h1['two_seventy_state_pm_lower_loss_count']} von {h1['two_seventy_state_case_count']}, 270toWin/JHK in {h1['two_seventy_state_270_lower_loss_count']} von {h1['two_seventy_state_case_count']}.</td></tr>
  </table>
  <p>Interpretation: H1 spricht fuer tiefere Forecast-Verluste von Polymarket im getesteten Fenster. Die neue Kalibrierungsdiagnostik ersetzt die schwache Ein-Outcome-Reliability-Kurve durch geloeste Fallartefakte und zeigt die Grenze klarer: niedrigere mittlere Brier-Verluste ja, klarer Kalibrierungssieg oder breiter Viele-Faelle-Beweis nein. Das neue State-Date-Poll-Panel ist der groesste poll-derived Vergleich und spricht gegen den starken Polymarket-Claim: {h1['state_poll_panel_poll_lower_loss_count']} von {h1['state_poll_panel_case_count']} Zeilen liegen bei der poll-derived Transformation niedriger, Polymarket nur in {h1['state_poll_panel_pm_lower_loss_count']}. Die Temporal-Diagnose macht die Nuance sichtbar: In den Polymarket-stuetzenden Monaten {h1['state_poll_temporal_support_months']} liegt Polymarket in {h1['state_poll_temporal_support_pm_lower_loss_count']} von {h1['state_poll_temporal_support_row_count']} Zeilen niedriger, aber die frueheren Monate dominieren den Vollpanel-Befund. Die Forecast-Horizon-Diagnose zeigt dieselbe Grenze methodisch klarer: Im &lt;=90-Tage-Fenster vor der Wahl liegt Polymarket in {h1['state_poll_horizon_near_pm_lower_loss_count']} von {h1['state_poll_horizon_near_row_count']} Zeilen niedriger; auf State-Ebene stuetzen {h1['state_poll_horizon_state_pm_mean_support_count']} von {h1['state_poll_horizon_state_count']} States Polymarket nach mittlerem Brier und Mehrheit niedrigerer Tagesverluste. Die &lt;=90-Day-Score-Quality-Diagnose verdichtet dieses Fenster auf {h1['state_poll_near_quality_forecast_row_count']} Forecast-Zeilen und zeigt Polymarket mit niedrigerem Mean Brier ({h1['state_poll_near_quality_pm_mean_brier']:.4f} vs {h1['state_poll_near_quality_poll_mean_brier']:.4f}), niedrigerem Fixed-Bin-ECE ({h1['state_poll_near_quality_pm_ece']:.4f} vs {h1['state_poll_near_quality_poll_ece']:.4f}) und hoeherer Probability-Separation ({h1['state_poll_near_quality_pm_separation']:.4f} vs {h1['state_poll_near_quality_poll_separation']:.4f}). Ueber 90 Tage vor der Wahl dominiert die poll-derived Transformation. Die Final-Snapshot-Erweiterung stuetzt dieselbe Richtung in {h1['final_snapshot_pm_lower_loss_count']} von {h1['final_snapshot_case_count']} geloesten 2024-Outcomes. Die State-Poll-Snapshot-Erweiterung zeigt Polymarket in {h1['state_poll_snapshot_pm_lower_loss_count']} von {h1['state_poll_snapshot_case_count']} Faellen mit niedrigerem Verlust, wird aber vom groesseren Panel relativiert. Die 270toWin-Polling-Average-Erweiterung bringt {h1['two_seventy_poll_average_case_count']} weitere direkt pollbasierte State-Faelle: Polymarket hat den niedrigeren mittleren Brier ({h1['two_seventy_poll_average_mean_pm_brier']:.4f} vs {h1['two_seventy_poll_average_mean_poll_brier']:.4f}), aber nur in {h1['two_seventy_poll_average_pm_lower_loss_count']} von {h1['two_seventy_poll_average_case_count']} States den niedrigeren Einzelfallverlust. Die Sensitivitaet variiert nur die Poll-Fehlerannahme und zeigt fuer MAE {h1['state_poll_sensitivity_min_mae']:.1f} bis {h1['state_poll_sensitivity_max_mae']:.1f} Prozentpunkte in allen {h1['state_poll_sensitivity_row_count']} Parameterzeilen einen niedrigeren mittleren Polymarket-Brier im Snapshot. Der Coverage-Audit zeigt zugleich, dass die oeffentlich auffindbaren Polymarket-State-Maerkte nicht automatisch H1-Brier-Paare sind: von {h1['state_poll_coverage_polymarket_market_count']} Polymarket-State-Maerkten bleiben wegen der bewahrten 538-Snapshot-Abdeckung nur {h1['state_poll_coverage_valid_pair_count']} valide Snapshot-Paare. Die Rieke-Erweiterung deckt alle {h1['rieke_state_case_count']} US-State-Outcomes ab und zeigt im Aggregat ebenfalls einen niedrigeren mittleren Polymarket-Brier ({h1['rieke_state_mean_pm_brier']:.4f} vs {h1['rieke_state_mean_rieke_brier']:.4f}), aber nicht eine Mehrheit niedrigerer Einzelfallverluste: Polymarket liegt in {h1['rieke_state_pm_lower_loss_count']} von {h1['rieke_state_case_count']} States vorne, Rieke in {h1['rieke_state_rieke_lower_loss_count']} von {h1['rieke_state_case_count']}. Die 270toWin/JHK-Erweiterung stuetzt den aggregierten Brier-Befund erneut ({h1['two_seventy_state_mean_pm_brier']:.4f} vs {h1['two_seventy_state_mean_270_brier']:.4f}), zeigt aber wegen sicherer zensierter States ebenfalls keine Mehrheit niedrigerer Einzelfallverluste fuer Polymarket ({h1['two_seventy_state_pm_lower_loss_count']} von {h1['two_seventy_state_case_count']}). Deshalb ist die belastbare Aussage eine gemischte H1-Evidenz, kein Nachweis, dass Polymarket in den meisten State-Faellen oder im groesseren Poll-Panel besser ist. Die State-Erweiterungen verwenden dokumentierte Wahrscheinlichkeitsmodelle; sie sind keine Rohpolls, kein RCP-Vergleich und kein Speed-Test.</p>

  <h2>H2 - Event-Window-Reaktion</h2>
  <p>{h2['event_count']} kuratierte Ereignisse, {h2['summary_rows']} kompakte Summary-Zeilen. Daily-Daten erlauben keine Intraday-Speed-Aussage.</p>
  <table><tr><th>Primaerfenster</th><th>Finaler Change</th></tr>{h2_rows}</table>

  <h2>H3 - Wallet-Tier-Timing</h2>
  <p>{h3['model_rows']} alignierte Modellzeilen. Tier counts: {escape(h3['tier_counts_text'])}.</p>
  <p>Staerkste dokumentierte Korrelation: {escape(h3['top_correlation_label'])} = {h3['top_correlation']:.4f}. Kleinster Granger-p-Wert: {escape(h3['min_granger_label'])} = {h3['min_granger_p']:.4f}.</p>
  <p>Grenze: BUY-only Quelle, taegliche Aggregation und Multiple-Testing-Sensitivitaet.</p>

  <h2>Monitor-Prototyp</h2>
  <p>Recorded replay rows: {monitor['snapshot_count']}. Severity counts: {escape(monitor['severity_counts_text'])}. Latest live dashboard: {monitor['live_market_count']} Maerkte, {monitor['live_alert_count']} Alert-Zeilen. Wallet graph: {monitor['wallet_graph_nodes']} Nodes und {monitor['wallet_graph_edges']} Edges.</p>
  <p>Anomaly review queue: {monitor['anomaly_queue_rows']} Cases ({monitor['anomaly_high_priority_count']} high, {monitor['anomaly_medium_priority_count']} medium, {monitor['anomaly_low_priority_count']} low). Status: {escape(monitor['anomaly_review_status_counts'])}. Labels: {escape(monitor['anomaly_review_labels'])}.</p>
  <p>{escape(monitor['anomaly_allowed_interpretation'])}</p>
  <p>Der Monitor bleibt read-only und ist keine Trading- oder Profitabilitaetskomponente.</p>

  <h2>Schweizer Referendum</h2>
  <p>{swiss['poll_count']} Umfragen, {swiss['snapshot_count']} Polymarket-Snapshots und {swiss['history_rows']} bounded price-history rows.</p>
  <table><tr><th>Quelle</th><th>Neuester Poll</th><th>Poll Yes</th><th>Raw Gap</th></tr>{swiss_source_rows}</table>
  <p>Latest result: Polymarket Yes {swiss['latest_poly_yes_pct']:.1f}%, latest matched poll Yes {swiss['latest_poll_yes_pct']:.1f}%, raw gap {swiss['latest_raw_gap_pp']:+.1f} pp, decided-voter gap {swiss['latest_decided_gap_pp']:+.1f} pp.</p>

  <h2>Visualisierungen</h2>
  {figures}
</body>
</html>
"""


def _figure_html(figure: FigureSpec, *, html_output: Path) -> str:
    rel = Path(_relative_path(figure.path.resolve(), html_output.parent.resolve()))
    return (
        "<figure>"
        f"<img src=\"{escape(rel.as_posix())}\" alt=\"{escape(figure.caption)}\">"
        f"<figcaption><strong>{escape(figure.caption)}</strong><br>{escape(figure.note)}</figcaption>"
        "</figure>"
    )


def write_docx(data: dict[str, Any], output_path: Path) -> None:
    """Write the formatted DOCX report."""

    doc = Document()
    _setup_document(doc)
    _add_cover(doc, data)
    _add_toc_note(doc)
    _add_highlevel_status_section(doc, data["thesis_highlevel"])
    _add_advisor_handoff_package_section(doc, data["advisor_handoff"])
    _add_project_highlevel_matrix_section(doc, data["project_highlevel"])
    _add_submission_readiness_section(doc, data["submission_readiness"])
    _add_drafting_sequence_section(doc, data["drafting_sequence"])
    _add_bounded_chapter_draft_section(doc, data["bounded_chapter_draft"])
    _add_source_gated_drafting_section(doc, data["source_gated_drafting"])
    _add_worksheet_drafting_bridge_section(doc, data["worksheet_drafting_bridge"])
    _add_next_work_section(doc, data["next_work"])
    _add_execution_checklist_section(doc, data["execution_checklist"])
    _add_research_design_section(doc, data)
    _add_literature_section(doc, data["literature"], data["source_review"])
    _add_methodology_section(doc, data)
    _add_interpretation_section(doc, data)
    _add_project_overview(doc, data)
    _add_h1_section(doc, data["h1"])
    _add_h2_section(doc, data["h2"])
    _add_h3_section(doc, data["h3"])
    _add_monitor_section(doc, data["monitor"])
    _add_swiss_section(doc, data["swiss"])
    _add_figures_section(doc, data["figures"])
    _add_presentation_section(doc)
    _add_appendix(doc, data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def write_pipeline_overview(path: Path) -> Path:
    """Create a simple pipeline overview figure for the report."""

    labels = [
        ("Datenquellen", "SQLite, Polymarket,\nFiveThirtyEight, Umfragen"),
        ("Validierung", "Schemas, Tests,\nProjekt-Guardrails"),
        ("Python-Analyse", "H1 Brier, H2 Event-Windows,\nH3 Wallet-Tiers, Swiss Vergleich"),
        ("Artefakte", "CSV, JSON, PNG,\nHTML-Dashboards"),
        ("Interpretation", "Dozentenbericht,\nThesis-Text, spaetere LLMs nur auditiert"),
    ]
    fig, ax = plt.subplots(figsize=(13, 4.2))
    ax.axis("off")
    x_positions = [0.08, 0.29, 0.50, 0.71, 0.90]
    for idx, ((title, body), x) in enumerate(zip(labels, x_positions, strict=True)):
        ax.text(
            x,
            0.62,
            title,
            ha="center",
            va="center",
            fontsize=11.5,
            fontweight="bold",
            color="#0B2545",
            bbox=dict(boxstyle="round,pad=0.45", facecolor="#E8EEF5", edgecolor="#5B7EA4"),
        )
        ax.text(x, 0.34, body, ha="center", va="top", fontsize=9.5, color="#333333")
        if idx < len(x_positions) - 1:
            ax.annotate(
                "",
                xy=(x_positions[idx + 1] - 0.075, 0.62),
                xytext=(x + 0.075, 0.62),
                arrowprops=dict(arrowstyle="->", color="#5B7EA4", lw=2),
            )
    ax.text(
        0.5,
        0.06,
        "Kernregel: Die Analyse berechnet Metriken deterministisch in Python. "
        "Interpretation kommt erst nach validierten Artefakten.",
        ha="center",
        fontsize=10,
        color="#555555",
    )
    fig.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)
    return path


def _relative_path(path: Path, base: Path) -> str:
    return os.path.relpath(path, base)


def _thesis_highlevel_data(
    metadata: dict[str, Any],
    captions: pd.DataFrame,
) -> dict[str, Any]:
    """Summarise the current thesis consolidation layer for the advisor report."""

    outputs = metadata.get("outputs", {})
    caption_counts = metadata.get("table_figure_caption_counts", {})
    citation_counts = metadata.get("citation_review_packet_counts", {})
    truthy = captions["include_in_core_package"].astype(str).str.lower().isin(
        {"true", "1", "yes"}
    )
    core = captions[truthy]
    core_tables = core[core["package_type"] == "table"]
    core_figures = core[core["package_type"] == "figure"]
    return {
        "active_phase": "Phase 12: Thesis Consolidation And Evidence Mapping",
        "caption_rows": int(caption_counts.get("total_caption_rows", len(captions))),
        "core_tables": int(
            caption_counts.get("core_table_captions", len(core_tables))
        ),
        "core_figures": int(
            caption_counts.get("core_figure_captions", len(core_figures))
        ),
        "evidence_rows": int(outputs.get("evidence_rows", 0)),
        "core_result_rows": int(outputs.get("core_result_rows", 0)),
        "chapter_rows": int(outputs.get("chapter_rows", 0)),
        "citation_packets": int(outputs.get("citation_review_packet_rows", 0)),
        "full_review_packets": int(
            citation_counts.get("full_review_required_packets", 0)
        ),
        "rows": [
            (
                "Empirischer Kern",
                "H1 Forecast-Qualitaet, H2 Event-Windows und H3 Wallet-Timing sind die zentrale Ergebnisbasis.",
                "Die Bachelorarbeit sollte diese drei Strukturen zuerst schreiben und erst danach Monitor, Swiss und Agenten einordnen.",
            ),
            (
                "Tabellen und Figuren",
                f"{len(core_tables)} Kern-Tabellen und {len(core_figures)} Kern-Figuren sind ueber `thesis_table_figure_captions.csv` beschriftet.",
                "Der Dozent bekommt eine fokussierte Ergebnisdarstellung statt einer Rohartefakt-Sammlung.",
            ),
            (
                "Quellen und Zitation",
                "Die Citation-Review-Pakete verknuepfen Quellen mit Evidence-IDs, erlaubtem Wording und Review-Gates.",
                "Finale Thesis-Zitate brauchen noch Seiten- oder Abschnittsnachweise; candidate Quellen bleiben blockiert.",
            ),
            (
                "Monitor und Swiss",
                "Monitor bleibt Prototype/Appendix; Swiss bleibt bis zum offiziellen Ergebnis beschreibender Side-Track.",
                "Beide Teile duerfen die H1-H3-Kernaussage nicht staerker machen als die deterministischen Artefakte erlauben.",
            ),
            (
                "Agenten-Ausblick",
                "Agenten koennen spaeter bei Source Review, Drafting und Guardrail-Checks helfen, bleiben aber jetzt deaktiviert.",
                "Keine Runtime-Agenten, kein MCP, keine Modell-Router und keine LLM-Metriken vor stabilem deterministic core.",
            ),
        ],
    }


def _project_highlevel_report_data(project_highlevel: pd.DataFrame) -> dict[str, Any]:
    """Translate the generated project matrix into concise advisor-facing rows."""

    status_labels = {
        "active_thesis_core": "Aktiver Thesis-Kern",
        "thesis_facing_ready": "Thesis-ready",
        "thesis_facing_ready_with_limits": "Thesis-ready mit Grenzen",
        "active_gate": "Aktives Gate",
        "thesis_facing_package": "Thesis-Paket",
        "paused_appendix_only": "Pausiert / Appendix",
        "descriptive_pending_result": "Beschreibend bis Resultat",
        "documentation_only_deferred": "Nur Ausblick",
        "project_management_ready": "Advisor-Abstimmung",
    }
    layer_labels = {
        "project_00_current_frame": "Gesamtrahmen",
        "project_01_h1_forecast_quality": "H1 Forecast-Qualitaet",
        "project_02_h2_event_windows": "H2 Event-Windows",
        "project_03_h3_wallet_timing": "H3 Wallet-Timing",
        "project_04_source_review_gate": "Quellen und Zitation",
        "project_05_table_figure_package": "Tabellen/Figuren-Paket",
        "project_06_monitor_review_access": "Monitor Review-Access",
        "project_07_swiss_referendum": "Swiss Referendum",
        "project_08_future_agents": "Agenten-Ausblick",
        "project_09_advisor_iteration": "Dozentenabstimmung",
    }
    decision_labels = {
        "project_00_current_frame": "H1-H3 als Kern schreiben; Decision Queues 10/5/8 mit 0 final-ready rows abarbeiten.",
        "project_01_h1_forecast_quality": "H1 als begrenzte Polymarket-Stuetze schreiben; H1 Queue: 10 rows, T2/F1, 0 final-ready.",
        "project_02_h2_event_windows": "H2 als taegliche Event-Window-Evidenz schreiben; H2 Queue: 5 rows, T3/F2, Kausalclaim-Grenze.",
        "project_03_h3_wallet_timing": "H3 als Timingdiagnostik schreiben; H3 Queue: 8 rows, T4/F3, Granger-/Wallet-Grenze.",
        "project_04_source_review_gate": "H1/H2/H3 Decision Queues bilden den 23-row Manual Source Review Pfad; Status wird nicht automatisch hochgestuft.",
        "project_05_table_figure_package": "5 Kern-Tabellen und 4 Kern-Figuren mit Captions verwenden.",
        "project_06_monitor_review_access": "Review-Access bleibt pausiert; Monitor bleibt Appendix/Prototype.",
        "project_07_swiss_referendum": "Bis zum offiziellen Ergebnis nur beschreibend verwenden.",
        "project_08_future_agents": "Agenten bleiben Dokumentationsausblick, keine Runtime-Implementierung.",
        "project_09_advisor_iteration": "Dozent soll H1-Wording, Quellenreview-Tiefe, Swiss und Appendix-Scope absegnen.",
    }
    gate_labels = {
        "project_00_current_frame": "Decision Queues, Source Review und Kapiteldraft.",
        "project_01_h1_forecast_quality": "H1 Queue-Felder vor finalem Wording fuellen.",
        "project_02_h2_event_windows": "H2 Queue-Felder inklusive Kausalclaim-Grenze fuellen.",
        "project_03_h3_wallet_timing": "H3 Queue-Felder inklusive Granger-/Wallet-Grenze fuellen.",
        "project_04_source_review_gate": "Page-/Section-Notes, Claim-Support, Blocked-Wording und Citation-Use eintragen.",
        "project_05_table_figure_package": "In Thesis-Layout integrieren; Nummerierung spaeter finalisieren.",
        "project_06_monitor_review_access": "Nur mit Human Review oder separatem freigegebenem Goal fortsetzen.",
        "project_07_swiss_referendum": "Nach offiziellem 14. Juni 2026 Resultat Artefakte neu generieren.",
        "project_08_future_agents": "Separates Goal, Tests und llm_audit_log vor Aktivierung.",
        "project_09_advisor_iteration": "Feedback loggen und in kleinen Commit-Plan uebersetzen.",
    }
    rows = []
    for row in project_highlevel.sort_values("view_id").to_dict(orient="records"):
        view_id = str(row["view_id"])
        rows.append(
            {
                "view_id": view_id,
                "project_layer": layer_labels.get(view_id, str(row["project_layer"])),
                "status_de": status_labels.get(str(row["status"]), str(row["status"])),
                "current_decision_de": decision_labels.get(view_id, str(row["current_decision"])),
                "next_gate_de": gate_labels.get(view_id, str(row["next_gate"])),
                "thesis_use": str(row["thesis_use"]),
            }
        )
    return {
        "row_count": len(rows),
        "rows": rows,
    }


def _next_work_report_data(next_work_plan: pd.DataFrame) -> dict[str, Any]:
    ordered = next_work_plan.sort_values("priority_order")
    german_rows = {
        "work_01_source_review": (
            "Quellenreview Kernquellen",
            "Die 11 Priority-1-Quellen mit Seiten- oder Abschnittsnotizen pruefen.",
            "Quellenstatus nicht automatisch hochstufen.",
        ),
        "work_02_method_chapters": (
            "Einleitung, Theorie und Methodik schreiben",
            "Die 8 Kapitelplan-Zeilen fuer Front Matter und Methodenkapitel nutzen.",
            "RCP, H2-Kuration, H3-Tiers und Agent-Deferral explizit halten.",
        ),
        "work_03_h1_results": (
            "H1-Ergebniskapitel schreiben",
            "H1 als begrenzte Polymarket-Stuetze plus breiten Claim-Grenze formulieren.",
            "Keine universelle Polymarket-Ueberlegenheit behaupten.",
        ),
        "work_04_h2_h3_results": (
            "H2- und H3-Ergebniskapitel schreiben",
            "H2 als taegliche Event-Window-Reaktion und H3 als Wallet-Timingdiagnostik schreiben.",
            "Keine Intraday-, Kausalitaets-, Private-Information- oder Profitabilitaetsclaims.",
        ),
        "work_05_table_figure_integration": (
            "Kompakte Tabellen und Figuren integrieren",
            "5 Kern-Tabellen und 4 Kern-Figuren mit generierten Captions verwenden.",
            "Keine Rohartefakte ohne Update von Evidence Map und Kapitelplan einfuegen.",
        ),
        "work_06_monitor_appendix": (
            "Monitor als Appendix-Prototyp halten",
            "Monitor nur als read-only Review-Workflow und nicht als Kernbeweis darstellen.",
            "Keine Wallet-Adress-Exposition und keine Order- oder Trading-Pfade.",
        ),
        "work_07_swiss_result_gate": (
            "Swiss Side-Track nach Ergebnis finalisieren",
            "Bis zum offiziellen Ergebnis beschreibend bleiben und danach Artefakte neu generieren.",
            "Poll-Anteile nicht als Gewinnwahrscheinlichkeiten behandeln.",
        ),
        "work_08_agent_outlook": (
            "Agenten-Pipeline als Future Work halten",
            "Die 7 Protokollzeilen nur als Zukunftsdesign nutzen.",
            "Keine Runtime-Agenten, kein MCP, kein Model Routing und keine LLM-Metriken.",
        ),
        "work_09_advisor_iteration": (
            "Dozentenfeedback zur Scope-Engfuehrung nutzen",
            "H1-Wording, Source-Review-Tiefe, Swiss-Platzierung und Appendix-Scope klaeren.",
            "Empirischen Scope nicht erweitern, bevor der Kern geschrieben ist.",
        ),
        "work_10_final_qa": (
            "Finale Thesis-QA",
            "Tests, Review-Checks, Citation Checks, Tabellen/Figuren und Swiss-Spelling pruefen.",
            "Keine finale Aussage ueber Artefakte und gepruefte Quellen hinaus.",
        ),
    }
    rows = []
    for row in ordered.to_dict(orient="records"):
        workstream, next_action, guardrail = german_rows.get(
            str(row["workstream_id"]),
            (str(row["workstream"]), str(row["next_action"]), str(row["guardrail"])),
        )
        rows.append(
            {
                "priority_order": int(row["priority_order"]),
                "workstream": workstream,
                "next_action": next_action,
                "guardrail": guardrail,
            }
        )
    return {
        "row_count": len(rows),
        "first_workstream": str(ordered.iloc[0]["workstream_id"]),
        "final_workstream": str(ordered.iloc[-1]["workstream_id"]),
        "rows": rows,
    }


def _execution_checklist_report_data(execution_checklist: pd.DataFrame) -> dict[str, Any]:
    """Translate chapter execution tasks into advisor-facing rows."""

    required_columns = {
        "task_id",
        "chapter_title",
        "draft_action_de",
        "done_when_de",
        "advisor_question_ids",
    }
    missing = sorted(required_columns.difference(execution_checklist.columns))
    if missing:
        raise ValueError(f"execution checklist missing required columns: {missing}")

    ordered = execution_checklist.sort_values("task_id")
    rows = [
        {
            "task_id": str(row["task_id"]),
            "chapter_title": str(row["chapter_title"]),
            "draft_action_de": str(row["draft_action_de"]),
            "done_when_de": str(row["done_when_de"]),
            "advisor_question_ids": str(row["advisor_question_ids"]),
        }
        for row in ordered.to_dict(orient="records")
    ]
    return {
        "row_count": len(rows),
        "first_task": rows[0]["task_id"] if rows else "",
        "final_task": rows[-1]["task_id"] if rows else "",
        "rows": rows,
    }


def _advisor_handoff_report_data(advisor_handoff: pd.DataFrame) -> dict[str, Any]:
    """Translate advisor handoff package rows into report rows."""

    required_columns = {
        "package_order",
        "deliverable_id",
        "path",
        "handoff_use_de",
        "advisor_decision_de",
        "boundary_de",
    }
    missing = sorted(required_columns.difference(advisor_handoff.columns))
    if missing:
        raise ValueError(f"advisor handoff package missing required columns: {missing}")

    ordered = advisor_handoff.sort_values("package_order")
    rows = [
        {
            "package_order": int(row["package_order"]),
            "deliverable_id": str(row["deliverable_id"]),
            "path": str(row["path"]),
            "handoff_use_de": str(row["handoff_use_de"]),
            "advisor_decision_de": str(row["advisor_decision_de"]),
            "boundary_de": str(row["boundary_de"]),
        }
        for row in ordered.to_dict(orient="records")
    ]
    return {
        "row_count": len(rows),
        "first_deliverable": rows[0]["deliverable_id"] if rows else "",
        "final_deliverable": rows[-1]["deliverable_id"] if rows else "",
        "rows": rows,
    }


def _submission_readiness_report_data(readiness: pd.DataFrame) -> dict[str, Any]:
    """Translate submission-readiness gates into advisor-facing rows."""

    required_columns = {
        "gate_id",
        "gate_area",
        "current_status",
        "next_action_de",
        "blocker_or_limit_de",
        "thesis_use_de",
    }
    missing = sorted(required_columns.difference(readiness.columns))
    if missing:
        raise ValueError(f"submission readiness board missing required columns: {missing}")

    ordered = readiness.sort_values("gate_id")
    rows = [
        {
            "gate_area": str(row["gate_area"]),
            "current_status": str(row["current_status"]),
            "next_action_de": str(row["next_action_de"]),
            "blocker_or_limit_de": str(row["blocker_or_limit_de"]),
            "thesis_use_de": str(row["thesis_use_de"]),
        }
        for row in ordered.to_dict(orient="records")
    ]
    final_blocked = [
        row for row in rows if row["current_status"].startswith("final_blocked")
    ]
    draft_ready = [
        row
        for row in rows
        if row["current_status"].startswith("ready_for")
        or row["current_status"] == "advisor_discussion_now"
    ]
    return {
        "row_count": len(rows),
        "draft_ready_count": len(draft_ready),
        "final_blocked_count": len(final_blocked),
        "source_review_status": _status_for_gate(rows, "source_review"),
        "swiss_status": _status_for_gate(rows, "swiss_result_gate"),
        "agent_status": _status_for_gate(rows, "agent_future_work"),
        "rows": rows,
    }


def _drafting_sequence_report_data(sequence: pd.DataFrame) -> dict[str, Any]:
    """Translate the ordered drafting sequence into advisor-facing rows."""

    required_columns = {
        "sequence_id",
        "priority_order",
        "thesis_section",
        "draft_permission",
        "current_status",
        "writing_action_de",
        "blocker_or_gate_de",
        "must_not_claim_de",
    }
    missing = sorted(required_columns.difference(sequence.columns))
    if missing:
        raise ValueError(f"drafting sequence missing required columns: {missing}")

    ordered = sequence.sort_values("priority_order")
    rows = [
        {
            "sequence_id": str(row["sequence_id"]),
            "priority_order": int(row["priority_order"]),
            "thesis_section": str(row["thesis_section"]),
            "draft_permission": str(row["draft_permission"]),
            "current_status": str(row["current_status"]),
            "writing_action_de": str(row["writing_action_de"]),
            "blocker_or_gate_de": str(row["blocker_or_gate_de"]),
            "must_not_claim_de": str(row["must_not_claim_de"]),
        }
        for row in ordered.to_dict(orient="records")
    ]
    return {
        "row_count": len(rows),
        "first_step": rows[0]["sequence_id"] if rows else "",
        "final_step": rows[-1]["sequence_id"] if rows else "",
        "bounded_write_now_count": sum(
            row["draft_permission"] == "write_now_bounded" for row in rows
        ),
        "final_blocked_count": sum(
            "final_blocked" in row["draft_permission"] for row in rows
        ),
        "future_work_only_count": sum(
            row["draft_permission"] == "future_work_only" for row in rows
        ),
        "rows": rows,
    }


def _bounded_chapter_draft_report_data(draft: pd.DataFrame) -> dict[str, Any]:
    """Translate H1-H2-H3 bounded chapter draft rows into advisor-facing rows."""

    required_columns = {
        "thesis_area",
        "draft_order",
        "draft_step",
        "method_evidence_ids",
        "interpretation_evidence_ids",
        "literature_source_ids",
        "deterministic_artifacts",
        "selected_tables",
        "selected_figures",
        "source_review_gate_de",
        "ready_for_bounded_draft",
        "ready_for_final_submission",
    }
    missing = sorted(required_columns.difference(draft.columns))
    if missing:
        raise ValueError(f"H1-H2-H3 bounded chapter draft missing required columns: {missing}")

    ordered = draft.sort_values(["thesis_area", "draft_order"])
    rows = ordered.to_dict(orient="records")
    chapter_rows: list[dict[str, Any]] = []
    for area in ("H1", "H2", "H3"):
        area_rows = [row for row in rows if str(row["thesis_area"]) == area]
        if not area_rows:
            raise ValueError(f"H1-H2-H3 bounded chapter draft missing area: {area}")
        first = area_rows[0]
        source_gate = str(first["source_review_gate_de"])
        chapter_rows.append(
            {
                "thesis_area": area,
                "method_evidence_ids": str(first["method_evidence_ids"]),
                "interpretation_evidence_ids": str(first["interpretation_evidence_ids"]),
                "literature_artifact_summary_de": (
                    f"Literatur `{first['literature_source_ids']}`; Artefakte "
                    f"`{_first_items(str(first['deterministic_artifacts']), max_items=2)}`."
                ),
                "table_figure_de": f"{first['selected_tables']} / {first['selected_figures']}",
                "source_review_gate_summary_de": _first_sentence(source_gate),
                "step_count": len(area_rows),
            }
        )

    ready_for_bounded = sum(
        _bool_text(row["ready_for_bounded_draft"]) for row in rows
    )
    ready_for_final = sum(
        _bool_text(row["ready_for_final_submission"]) for row in rows
    )
    return {
        "row_count": len(rows),
        "rows_per_chapter": int(len(rows) / 3) if rows else 0,
        "bounded_ready_count": ready_for_bounded,
        "final_ready_count": ready_for_final,
        "chapter_rows": chapter_rows,
    }


def _source_gated_drafting_report_data(drafting: pd.DataFrame) -> dict[str, Any]:
    """Translate paragraph-level source-gated drafting rows for the advisor report."""

    required_columns = {
        "thesis_area",
        "draft_sequence_order",
        "draft_section_de",
        "method_evidence_ids",
        "interpretation_evidence_ids",
        "selected_tables",
        "selected_figures",
        "manual_execution_rows",
        "manual_execution_pending_rows",
        "manual_execution_final_ready_rows",
        "writer_action_de",
        "source_review_action_de",
        "final_gate_de",
        "ready_for_bounded_draft",
        "ready_for_final_submission",
        "draft_status",
    }
    missing = sorted(required_columns.difference(drafting.columns))
    if missing:
        raise ValueError(f"source-gated thesis drafting pass missing columns: {missing}")

    ordered = drafting.sort_values(["draft_sequence_order"])
    rows = ordered.to_dict(orient="records")
    if len(rows) != 15:
        raise ValueError("source-gated thesis drafting pass must have 15 rows for the report")

    chapter_rows: list[dict[str, Any]] = []
    for area in ("H1", "H2", "H3"):
        area_rows = [row for row in rows if str(row["thesis_area"]) == area]
        if len(area_rows) != 5:
            raise ValueError(f"source-gated thesis drafting pass missing five rows for {area}")
        first = area_rows[0]
        chapter_rows.append(
            {
                "thesis_area": area,
                "step_count": len(area_rows),
                "step_summary_de": " -> ".join(
                    str(row["draft_section_de"]) for row in area_rows
                ),
                "method_evidence_ids": str(first["method_evidence_ids"]),
                "interpretation_evidence_ids": str(first["interpretation_evidence_ids"]),
                "manual_review_summary_de": (
                    f"{int(first['manual_execution_rows'])} rows; "
                    f"{int(first['manual_execution_pending_rows'])} pending; "
                    f"{int(first['manual_execution_final_ready_rows'])} final-ready; "
                    "Manual Source Review Follow-up Overview-/Ledger-Abgleich offen"
                ),
                "table_figure_de": f"{first['selected_tables']} / {first['selected_figures']}",
                "status_de": str(first["draft_status"]),
            }
        )

    step_rows = [
        {
            "draft_sequence_order": int(row["draft_sequence_order"]),
            "thesis_area": str(row["thesis_area"]),
            "draft_section_de": str(row["draft_section_de"]),
            "writer_action_de": str(row["writer_action_de"]),
            "final_gate_short_de": _source_gate_short_text(str(row["final_gate_de"])),
        }
        for row in rows
    ]

    chapter_first_rows = (
        ordered.drop_duplicates(subset=["thesis_area"])
        .sort_values("thesis_area")
        .to_dict(orient="records")
    )
    manual_rows_linked = sum(int(row["manual_execution_rows"]) for row in chapter_first_rows)
    manual_pending_rows = sum(
        int(row["manual_execution_pending_rows"]) for row in chapter_first_rows
    )
    manual_final_ready_rows = sum(
        int(row["manual_execution_final_ready_rows"]) for row in chapter_first_rows
    )
    review_control_de = (
        "Review-Access bleibt pausiert; auf Projektebene geht es jetzt ueber "
        "Advisor-Feedback, Manual Source Review Follow-up Overview-/Ledger-Abgleich, "
        "H1-H3-Prosa, kompakte Tabellen/Figuren, Swiss-Result-Gate und finale "
        "DOCX-Render-QA weiter. Kontrollpunkt: "
        f"{manual_rows_linked} offene H1-H2-H3 Review-Rows, "
        f"{manual_pending_rows} pending, {manual_final_ready_rows} final-ready; "
        "keine finale Zitation und keine Quellenstatus-Hochstufung vor "
        "abgeschlossenem Overview-/Ledger-Abgleich."
    )

    return {
        "row_count": len(rows),
        "rows_per_chapter": int(len(rows) / 3) if rows else 0,
        "bounded_ready_count": sum(
            _bool_text(row["ready_for_bounded_draft"]) for row in rows
        ),
        "final_ready_count": sum(
            _bool_text(row["ready_for_final_submission"]) for row in rows
        ),
        "manual_rows_linked": manual_rows_linked,
        "manual_pending_rows": manual_pending_rows,
        "manual_final_ready_rows": manual_final_ready_rows,
        "review_control_de": review_control_de,
        "chapter_rows": chapter_rows,
        "step_rows": step_rows,
    }


def _worksheet_drafting_bridge_report_data(bridge: pd.DataFrame) -> dict[str, Any]:
    """Translate worksheet-to-drafting bridge rows for the advisor report."""

    required_columns = {
        "thesis_area",
        "worksheet_rows",
        "method_rows",
        "interpretation_rows",
        "unique_sources",
        "method_interpretation_source_artifact_gap_rows",
        "pending_citation_rows",
        "final_release_ready_rows",
        "drafting_steps",
        "bounded_draft_ready_steps",
        "final_submission_ready_steps",
        "selected_tables",
        "selected_figures",
        "source_artifact_rule_de",
        "final_blocker_de",
        "future_agent_boundary_de",
        "ready_for_bounded_drafting",
        "ready_for_final_release",
    }
    missing = sorted(required_columns.difference(bridge.columns))
    if missing:
        raise ValueError(f"worksheet drafting bridge missing columns: {missing}")

    ordered = bridge.sort_values("bridge_order")
    rows = ordered.to_dict(orient="records")
    if len(rows) != 4:
        raise ValueError("worksheet drafting bridge must have 4 rows for the report")
    total_rows = [row for row in rows if str(row["thesis_area"]) == "TOTAL"]
    if len(total_rows) != 1:
        raise ValueError("worksheet drafting bridge must contain one TOTAL row")
    total = total_rows[0]
    if int(total["method_interpretation_source_artifact_gap_rows"]) != 0:
        raise ValueError("worksheet drafting bridge has source/artifact gaps")
    if int(total["final_release_ready_rows"]) != 0:
        raise ValueError("worksheet drafting bridge must not be final-release-ready")

    chapter_rows: list[dict[str, Any]] = []
    for row in rows:
        area = str(row["thesis_area"])
        chapter_rows.append(
            {
                "thesis_area": area,
                "worksheet_summary_de": (
                    f"{int(row['worksheet_rows'])} rows; "
                    f"{int(row['unique_sources'])} Quellen; "
                    f"{int(row['method_rows'])} method / "
                    f"{int(row['interpretation_rows'])} interpretation"
                ),
                "drafting_summary_de": (
                    f"{int(row['drafting_steps'])} steps; "
                    f"{int(row['bounded_draft_ready_steps'])} bounded; "
                    f"{int(row['final_submission_ready_steps'])} final-ready"
                ),
                "table_figure_de": f"{row['selected_tables']} / {row['selected_figures']}",
                "gate_short_de": _first_sentence(str(row["final_blocker_de"])),
            }
        )

    return {
        "row_count": len(rows),
        "worksheet_rows": int(total["worksheet_rows"]),
        "method_rows": int(total["method_rows"]),
        "interpretation_rows": int(total["interpretation_rows"]),
        "unique_sources": int(total["unique_sources"]),
        "source_artifact_gap_rows": int(
            total["method_interpretation_source_artifact_gap_rows"]
        ),
        "pending_citation_rows": int(total["pending_citation_rows"]),
        "drafting_steps": int(total["drafting_steps"]),
        "final_release_ready_rows": int(total["final_release_ready_rows"]),
        "source_artifact_rule_de": str(total["source_artifact_rule_de"]),
        "future_agent_boundary_de": str(total["future_agent_boundary_de"]),
        "chapter_rows": chapter_rows,
    }


def _first_items(value: str, *, max_items: int) -> str:
    items = [item.strip() for item in value.split(";") if item.strip()]
    selected = items[:max_items]
    if len(items) > max_items:
        selected.append(f"plus {len(items) - max_items} weitere")
    return "; ".join(selected)


def _first_sentence(value: str) -> str:
    sentence = value.split(".", maxsplit=1)[0].strip()
    if sentence:
        return sentence + "."
    return value.strip()


def _source_gate_short_text(value: str) -> str:
    short_text = _first_sentence(value)
    overview_gate = "Manual Source Review Follow-up Overview-/Ledger-Abgleich"
    if overview_gate in value and overview_gate not in short_text:
        return (
            f"{short_text} {overview_gate} vor Citation Gate sichtbar halten."
        )
    return short_text


def _bool_text(value: object) -> bool:
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() in {"true", "1", "yes", "ja"}


def _status_for_gate(rows: Sequence[dict[str, Any]], gate_area: str) -> str:
    for row in rows:
        if row["gate_area"] == gate_area:
            return str(row["current_status"])
    raise ValueError(f"readiness gate missing in report data: {gate_area}")


def _interpretation_rows(data: dict[str, Any]) -> list[dict[str, str]]:
    """Create thesis-facing insight rows from deterministic report data."""

    h1 = data["h1"]
    h2 = data["h2"]
    h3 = data["h3"]
    monitor = data["monitor"]
    swiss = data["swiss"]
    largest_h2 = max(h2["primary_examples"], key=lambda row: abs(row["change_pp"]))
    return [
        {
            "bereich": "H1 Forecast-Qualitaet",
            "erkenntnis": (
                "Polymarket ist in den aktuellen Resultaten nicht pauschal "
                "ueberlegen, zeigt aber einen klaren Vorteil in bestimmten "
                "spaeten und kompetitiven Vergleichsfenstern."
            ),
            "evidenz": (
                f"Primary Brier: {h1['brier_polymarket']:.4f} vs "
                f"{h1['brier_fivethirtyeight']:.4f}; H1-Synthesis "
                f"{h1['synthesis_aggregate_support_count']} von "
                f"{h1['synthesis_evidence_row_count']} Zeilen mit niedrigerem "
                f"mittleren Polymarket-Brier, aber nur "
                f"{h1['synthesis_majority_support_count']} von "
                f"{h1['synthesis_evidence_row_count']} mit Fallmehrheit."
            ),
            "interpretation": (
                "Die belastbare Aussage ist ein bounded Forecast-Quality-"
                "Vorteil, besonders im <=90-Tage Low/Middle-Poll-Distanz-Scope."
            ),
            "grenze": (
                f"Der breite Viele-Faelle-Claim bleibt "
                f"{h1['poll_result_goal_status']}; das State-Date-Vollpanel "
                f"stuetzt poll-derived in {h1['state_poll_panel_poll_lower_loss_count']} "
                f"von {h1['state_poll_panel_case_count']} Zeilen."
            ),
        },
        {
            "bereich": "H2 Event-Windows",
            "erkenntnis": (
                "Polymarket bewegt sich um mehrere kuratierte politische "
                "Ereignisse sichtbar in der Tagesauflosung."
            ),
            "evidenz": (
                f"{h2['event_count']} kuratierte Ereignisse; groesstes "
                f"Primaerfenster nach Betrag: {largest_h2['event']} "
                f"{largest_h2['change_pp']:+.1f} Prozentpunkte."
            ),
            "interpretation": (
                "Das stuetzt eine These, dass oeffentliche Ereignisse in "
                "den Marktpreisen sichtbar werden koennen."
            ),
            "grenze": (
                "Taegliche Daten zeigen Reaktionsrichtung und Groessenordnung, "
                "aber keine intraday Reaktionsgeschwindigkeit."
            ),
        },
        {
            "bereich": "H3 Wallet-Timing",
            "erkenntnis": (
                "Top-Wallet-Tier-Aktivitaet zeigt eine messbare, aber vorsichtig "
                "zu formulierende Timing-Struktur."
            ),
            "evidenz": (
                f"{h3['model_rows']} alignierte Modellzeilen; staerkste "
                f"Korrelation {h3['top_correlation_label']} = "
                f"{h3['top_correlation']:.4f}; kleinster Granger-p-Wert "
                f"{h3['min_granger_p']:.4f}."
            ),
            "interpretation": (
                "Das ist als Vorhersage-/Timingdiagnostik verwendbar und "
                "motiviert weitere Sensitivitaetschecks."
            ),
            "grenze": (
                "BUY-only Quelle, Tagesaggregation und Multiple-Testing-Risiko "
                "begrenzen die Aussage."
            ),
        },
        {
            "bereich": "Monitor und Review-Queue",
            "erkenntnis": (
                "Der Monitor hat die richtige Rolle als Kontroll- und Review-"
                "Infrastruktur, nicht als Ergebnisgenerator fuer starke Claims."
            ),
            "evidenz": (
                f"{monitor['anomaly_queue_rows']} aktuelle Review-Cases, "
                f"davon {monitor['anomaly_high_priority_count']} high und "
                f"{monitor['anomaly_medium_priority_count']} medium; Status "
                f"{monitor['anomaly_review_status_counts']}."
            ),
            "interpretation": (
                "Die Review-Queue ist methodisch wichtig, weil sie auffaellige "
                "Faelle von thesis-faehiger Evidenz trennt."
            ),
            "grenze": monitor["anomaly_limitation"],
        },
        {
            "bereich": "Swiss-Referendum Side-Track",
            "erkenntnis": (
                "Der laufende Referendumsvergleich zeigt aktuell eine grosse "
                "Divergenz zwischen Marktpreis und Umfrageanteilen."
            ),
            "evidenz": (
                f"{swiss['poll_count']} Umfragen, {swiss['snapshot_count']} "
                f"Polymarket-Snapshots; latest Polymarket Yes "
                f"{swiss['latest_poly_yes_pct']:.1f}%, latest poll Yes "
                f"{swiss['latest_poll_yes_pct']:.1f}%, raw gap "
                f"{swiss['latest_raw_gap_pp']:+.1f} pp."
            ),
            "interpretation": (
                "Das ist ein anschauliches aktuelles Beispiel fuer die "
                "Trennung von Marktpreisen und traditionellen Umfragesignalen."
            ),
            "grenze": (
                "Umfrageanteile sind keine Gewinnwahrscheinlichkeiten; vor dem "
                "Abstimmungsergebnis gibt es keine finale Effizienzbewertung."
            ),
        },
    ]


def _method_decision_rows() -> list[tuple[str, str, str]]:
    """Explain why the chosen empirical steps fit the research question."""

    return [
        (
            "Brier Score und DM-Test",
            "H1 vergleicht Probability-Forecasts, deshalb braucht es einen Verlustscore und einen Test auf Verlustserien.",
            "Die Aussage bleibt Forecast-Qualitaet, nicht Reaktionsgeschwindigkeit oder Mechanismus.",
        ),
        (
            "Vorab kuratierte Events",
            "H2 soll nicht Ereignisse nach sichtbaren Kursbewegungen auswaehlen.",
            "Die Event-Auswahl ist dadurch strenger, aber weniger flexibel.",
        ),
        (
            "Verteilungsbasierte Wallet-Tiers",
            "H3 vermeidet willkuerliche Whale-Schwellen und leitet Tiers aus der beobachteten Verteilung ab.",
            "Die Quelle bleibt BUY-only und kann nicht alle Marktaktivitaet abbilden.",
        ),
        (
            "Review-Queue statt Agentenclaim",
            "Auffaellige Monitor-Faelle brauchen menschliche Quellenpruefung vor Interpretation.",
            "Aktuelle Cases bleiben Review-Cues und werden nicht automatisch thesis-faehig.",
        ),
        (
            "Swiss-Referendum als Side-Track",
            "Ein aktueller, zeitlich begrenzter Markt zeigt die Methode in einem laufenden politischen Kontext.",
            "Die Analyse bleibt bis zum Ergebnis des 14. Juni 2026 beschreibend.",
        ),
    ]


def _setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "BA Thesis Projektbericht"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_paragraph_font(header, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "Deterministischer Bericht aus lokalen Artefakten"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font(footer, size=9, color=MUTED)


def _add_cover(doc: Document, data: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("DOZENTENBERICHT")
    _set_run(run, size=11, bold=True, color=MUTED)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Projektueberblick Bachelorarbeit")
    _set_run(run, size=24, bold=True, color=RGBColor(0, 0, 0))

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(data["project"]["working_title"])
    _set_run(run, size=13, color=MUTED)

    metadata = [
        ("Stand", data["generated_at_utc"]),
        (
            "Aktive Projektphase",
            "Phase 12: Thesis Consolidation And Evidence Mapping; Monitor und Swiss bleiben bounded Side-Tracks",
        ),
        ("Teststatus", data["project"]["test_summary"]),
        ("Berichtsquelle", "Lokaler Worktree und deterministische Artefakte"),
    ]
    table = _add_table(doc, metadata, ["Feld", "Wert"], [1800, 7560])
    _shade_table_header(table)
    _add_callout(
        doc,
        "Kernaussage",
        (
            "Das Projekt ist methodisch so aufgebaut, dass die deterministische "
            "Analyse vor jeder Interpretation steht. H1-H3 bilden die empirische "
            "Basis; der Monitor ist ein read-only Forschungsprototyp mit "
            "deterministischer Anomaly-Review-Queue; die Swiss-Referendum-Spur "
            "ist eine separate, aktuelle Vergleichsanalyse."
        ),
    )


def _add_toc_note(doc: Document) -> None:
    doc.add_heading("Lesefuehrung", level=1)
    _add_bullets(
        doc,
        [
            "Der Highlevel-Block fasst den aktuellen Projektstand fuer den Dozenten zusammen.",
            "Die Projektmatrix zeigt Status, Entscheidung und Gate je Projektteil.",
            "Der Abschnitt Naechste Arbeitsschritte ordnet die kommenden Thesis-Workstreams.",
            "Abschnitt 1 formuliert Forschungsfrage, Hypothesen und BA-Aufbau.",
            "Abschnitt 2 ordnet die hinterlegten wissenschaftlichen Quellen ein.",
            "Abschnitt 3 begruendet Datenbasis, Methodik und Guardrails.",
            "Abschnitt 4 fasst Erkenntnisse, Evidenz, Interpretation und Grenzen zusammen.",
            "Abschnitte 6 bis 8 erklaeren H1, H2 und H3: Warum, Methode, Ergebnis, Grenze.",
            "Abschnitte 9 bis 10 ordnen Monitor-Prototyp und Schweizer Referendumsvergleich ein.",
            "Abschnitt 11 enthaelt Visualisierungen und einen kurzen Praesentationsplan.",
        ],
    )


def _add_highlevel_status_section(doc: Document, highlevel: dict[str, Any]) -> None:
    doc.add_heading("Highlevel-Projektstand fuer den Dozenten", level=1)
    doc.add_paragraph(
        "Der Review-Access bleibt pausiert. Der aktuelle Fortschritt liegt in "
        "der Thesis-Konsolidierung: Methoden, Interpretationen, Quellen, "
        "Tabellen und Figuren sind auf deterministische Artefakte gemappt."
    )
    rows = [
        ("Aktive Phase", highlevel["active_phase"]),
        (
            "Thesis-Paket",
            f"{highlevel['core_tables']} Kern-Tabellen, {highlevel['core_figures']} Kern-Figuren und {highlevel['caption_rows']} Caption-Zeilen.",
        ),
        (
            "Evidenz und Kapitel",
            f"{highlevel['evidence_rows']} Evidence-Zeilen, {highlevel['core_result_rows']} zentrale Resultatzeilen und {highlevel['chapter_rows']} Kapitelplan-Zeilen.",
        ),
        (
            "Citation-Gate",
            f"{highlevel['citation_packets']} Review-Pakete; {highlevel['full_review_packets']} brauchen Full-Source-Review vor finaler Zitation.",
        ),
        (
            "Agenten",
            "Nur dokumentierter Ausblick; keine Runtime-Agenten, kein MCP, keine Modell-Router und keine unlogged LLM-Interpretation.",
        ),
    ]
    table = _add_table(doc, rows, ["Ebene", "Stand"], [2100, 7260])
    _shade_table_header(table)
    table = _add_table(
        doc,
        highlevel["rows"],
        ["Ebene", "Stand", "Konsequenz fuer die Thesis"],
        [1900, 3560, 3900],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Naechste Lesart",
        (
            "Fuer den Dozenten ist der Kern jetzt: H1-H3 sind die empirische "
            "Basis, Monitor und Swiss bleiben abgegrenzt, und die Thesis soll "
            "mit wenigen guten Tabellen/Figuren statt vielen Rohartefakten "
            "geschrieben werden."
        ),
    )


def _add_advisor_handoff_package_section(
    doc: Document,
    advisor_handoff: dict[str, Any],
) -> None:
    doc.add_heading("Dozentenpaket und Uebergabereihenfolge", level=1)
    doc.add_paragraph(
        f"Das Advisor-Handoff-Paket ordnet {advisor_handoff['row_count']} "
        "Dateien fuer die naechste Betreuung. Zuerst kommt "
        f"`{advisor_handoff['first_deliverable']}`, danach folgen "
        "Word-Bericht, Absprache-Checklist, Submission Readiness Board, "
        "Drafting Sequence, Feedback-Log und die Arbeitsdateien fuer Kapitel "
        "und Source Review."
    )
    rows = [
        (
            row["package_order"],
            row["path"],
            row["handoff_use_de"],
            row["advisor_decision_de"],
            row["boundary_de"],
        )
        for row in advisor_handoff["rows"]
    ]
    table = _add_table(
        doc,
        rows,
        ["Reihenfolge", "Datei", "Verwendung", "Entscheidung", "Grenze"],
        [900, 2300, 2200, 2500, 1460],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Uebergabe",
        (
            "Fuer den Dozenten reicht als Einstieg das Word-Dokument plus "
            "Absprache-Checklist; der Uebergabetext eroeffnet die Nachricht. "
            "Die weiteren Dateien steuern Schreiben, Gates, Feedback und "
            "Quellenreview, ohne Review-Access, Agenten oder neue Empirie zu "
            "aktivieren."
        ),
    )


def _add_project_highlevel_matrix_section(
    doc: Document,
    project_highlevel: dict[str, Any],
) -> None:
    doc.add_heading("Projektmatrix fuer die naechste Abstimmung", level=1)
    doc.add_paragraph(
        f"Die Projektmatrix fasst {project_highlevel['row_count']} Ebenen als "
        "Status-, Entscheidungs- und Gate-Sicht zusammen. Sie zeigt explizit, "
        "dass Review-Access pausiert bleibt und Agenten nur dokumentierter "
        "Ausblick sind."
    )
    rows = [
        (
            row["project_layer"],
            row["status_de"],
            row["current_decision_de"],
            row["next_gate_de"],
        )
        for row in project_highlevel["rows"]
    ]
    table = _add_table(
        doc,
        rows,
        ["Projektteil", "Status", "Entscheidung", "Naechstes Gate"],
        [1900, 1700, 3360, 2400],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Highlevel-Entscheidung",
        (
            "Fuer die naechste Abstimmung reicht diese Matrix als Leitlinie: "
            "H1-H3 schreiben, Quellen reviewen, Monitor und Swiss begrenzen, "
            "Agenten erst nach separatem Goal aktivieren."
        ),
    )


def _add_submission_readiness_section(
    doc: Document,
    submission_readiness: dict[str, Any],
) -> None:
    doc.add_heading("Submission Readiness und finale Gates", level=1)
    doc.add_paragraph(
        f"Das Submission Readiness Board trennt {submission_readiness['row_count']} "
        "Gates in Draft-Arbeit, finale Blocker und Future Work. Draft-ready "
        f"Gates: {submission_readiness['draft_ready_count']}; final "
        f"blockierte Gates: {submission_readiness['final_blocked_count']}. "
        f"Source Review steht auf `{submission_readiness['source_review_status']}`, "
        f"Swiss steht auf `{submission_readiness['swiss_status']}`, Agenten "
        f"stehen auf `{submission_readiness['agent_status']}`."
    )
    rows = [
        (
            row["gate_area"],
            row["current_status"],
            row["next_action_de"],
            row["blocker_or_limit_de"],
        )
        for row in submission_readiness["rows"]
    ]
    table = _add_table(
        doc,
        rows,
        ["Gate", "Status", "Naechste Aktion", "Grenze"],
        [1800, 1900, 3100, 2560],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Finale Gate-Logik",
        (
            "Der Entwurf darf weitergeschrieben werden. Finale Abgabe bleibt "
            "aber blockiert, solange Source Review, Swiss-Resultatzuordnung "
            "oder DOCX-Render-QA offen sind."
        ),
    )


def _add_drafting_sequence_section(
    doc: Document,
    drafting_sequence: dict[str, Any],
) -> None:
    doc.add_heading("Schreibsequenz fuer den naechsten Entwurf", level=1)
    doc.add_paragraph(
        f"Die Drafting Sequence ordnet {drafting_sequence['row_count']} "
        "Schritte vom Quellenreview bis zur finalen QA. Erste Sequenz ist "
        f"`{drafting_sequence['first_step']}`, letzte Sequenz ist "
        f"`{drafting_sequence['final_step']}`. Bounded write-now: "
        f"{drafting_sequence['bounded_write_now_count']}; final blockiert: "
        f"{drafting_sequence['final_blocked_count']}; Future-work-only: "
        f"{drafting_sequence['future_work_only_count']}."
    )
    rows = [
        (
            row["priority_order"],
            row["thesis_section"],
            row["draft_permission"],
            row["writing_action_de"],
            row["must_not_claim_de"],
        )
        for row in drafting_sequence["rows"]
    ]
    table = _add_table(
        doc,
        rows,
        [
            "Prioritaet",
            "Thesis-Abschnitt",
            "Erlaubnis",
            "Schreibaktion",
            "Nicht behaupten",
        ],
        [760, 1650, 1650, 3100, 2200],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Naechste Schreibentscheidung",
        (
            "Praktisch heisst das: zuerst Quellenreview und Kapitel 1-3, dann "
            "H1-H2-H3-Ergebnisse mit kompakten Tabellen/Figuren, danach Monitor "
            "nur als Appendix, Swiss nur beschreibend und Agenten nur als "
            "Future-Work-Ausblick."
        ),
    )


def _add_bounded_chapter_draft_section(
    doc: Document,
    bounded_chapter_draft: dict[str, Any],
) -> None:
    doc.add_heading("Bounded H1-H2-H3 Kapitelentwurf", level=1)
    doc.add_paragraph(
        f"Der H1-H2-H3 Bounded Chapter Draft liefert "
        f"{bounded_chapter_draft['row_count']} geordnete Prosa-Bausteine: "
        f"{bounded_chapter_draft['rows_per_chapter']} je H1, H2 und H3. "
        f"Bounded-draft-ready: {bounded_chapter_draft['bounded_ready_count']}; "
        f"final-submission-ready: {bounded_chapter_draft['final_ready_count']}."
    )
    doc.add_paragraph(
        "Jede empirische Methode und jede Interpretation bleibt an Evidence-IDs, "
        "Literatur-IDs, deterministische Artefakte, kuratierte Tabellen/Figuren, "
        "Limitationen und Source-Review-Gates gebunden."
    )
    rows = [
        (
            row["thesis_area"],
            row["method_evidence_ids"],
            row["interpretation_evidence_ids"],
            row["literature_artifact_summary_de"],
            row["table_figure_de"],
            row["source_review_gate_summary_de"],
        )
        for row in bounded_chapter_draft["chapter_rows"]
    ]
    table = _add_table(
        doc,
        rows,
        ["Kapitel", "Methode", "Interpretation", "Quelle/Artefakt", "Tabelle/Figur", "Gate"],
        [620, 1450, 1750, 2450, 1180, 1910],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Schreibvorlage, nicht Abgabeclaim",
        (
            "Die Bausteine duerfen direkt in den empirischen BA-Kern "
            "ueberfuehrt werden. Sie ersetzen keine finale Zitation, "
            "keine Quellenpruefung und keine DOCX-Render-QA."
        ),
    )


def _add_source_gated_drafting_section(
    doc: Document,
    source_gated_drafting: dict[str, Any],
) -> None:
    doc.add_heading("Source-Gated H1-H2-H3 Drafting Sequence", level=1)
    doc.add_paragraph(
        "Der source-gated Thesis-Drafting-Pass bringt den H1-H2-H3-Kern in "
        f"{source_gated_drafting['row_count']} paragraphenweise Schreibschritte: "
        f"{source_gated_drafting['rows_per_chapter']} je Kapitel. "
        f"Bounded-draft-ready: {source_gated_drafting['bounded_ready_count']}; "
        f"final-submission-ready: {source_gated_drafting['final_ready_count']}. "
        f"Manual Source Review: {source_gated_drafting['manual_rows_linked']} "
        f"Rows verlinkt, {source_gated_drafting['manual_pending_rows']} pending, "
        f"{source_gated_drafting['manual_final_ready_rows']} final-ready."
    )
    doc.add_paragraph(
        "Fuer den Dozenten ist das die konkrete Schreibreihenfolge nach dem "
        "Bounded Chapter Draft: Methode/Resultat setzen, Interpretation und "
        "Limitation setzen, Tabelle/Figur einbauen, Manual Source Review "
        "ausfuehren, Finalgate und Future-Agent-Grenze sichtbar lassen."
    )
    doc.add_paragraph(source_gated_drafting["review_control_de"])
    chapter_rows = [
        (
            row["thesis_area"],
            f"{row['step_count']}: {row['step_summary_de']}",
            row["manual_review_summary_de"],
            row["table_figure_de"],
            row["status_de"],
        )
        for row in source_gated_drafting["chapter_rows"]
    ]
    chapter_table = _add_table(
        doc,
        chapter_rows,
        ["Kapitel", "Schritte", "Manual Source Review", "Tabelle/Figur", "Status"],
        [650, 3150, 1800, 1200, 2560],
    )
    _shade_table_header(chapter_table)
    step_rows = [
        (
            row["draft_sequence_order"],
            row["thesis_area"],
            row["draft_section_de"],
            row["writer_action_de"],
            row["final_gate_short_de"],
        )
        for row in source_gated_drafting["step_rows"]
    ]
    step_table = _add_table(
        doc,
        step_rows,
        ["Ordnung", "Kapitel", "Schreibschritt", "Writer Action", "Finalgate"],
        [650, 600, 1650, 3700, 2760],
    )
    _shade_table_header(step_table)
    _add_callout(
        doc,
        "Gating fuer den Dozenten",
        (
            "Auch diese Sequenz ist kein finaler Zitations- oder Abgabeclaim: "
            "Page-/Section-Note, Claim-Support, Blocked-Wording und Citation-Use "
            "bleiben manuelle Gates; Agenten bleiben documentation-only Future Work."
        ),
    )


def _add_worksheet_drafting_bridge_section(
    doc: Document,
    worksheet_drafting_bridge: dict[str, Any],
) -> None:
    doc.add_heading("Worksheet-to-Drafting Bridge fuer H1-H2-H3", level=1)
    doc.add_paragraph(
        "Die Worksheet-Drafting-Bridge verbindet "
        f"{worksheet_drafting_bridge['worksheet_rows']} manuelle H1-H2-H3 "
        "Worksheet-Zeilen mit "
        f"{worksheet_drafting_bridge['drafting_steps']} source-gated "
        "Schreibschritten. "
        f"Method rows: {worksheet_drafting_bridge['method_rows']}; "
        f"Interpretation rows: {worksheet_drafting_bridge['interpretation_rows']}; "
        f"Source/artifact gaps: {worksheet_drafting_bridge['source_artifact_gap_rows']}; "
        f"final-release rows: {worksheet_drafting_bridge['final_release_ready_rows']}."
    )
    doc.add_paragraph(worksheet_drafting_bridge["source_artifact_rule_de"])
    doc.add_paragraph(
        "Damit sieht der Dozent in einer kleinen Bruecke, welche Quellen- und "
        "Artefaktpflicht vor jedem Absatz gilt und warum der Haupttext bei "
        "T2/F1, T3/F2 und T4/F3 bleibt."
    )
    rows = [
        (
            row["thesis_area"],
            row["worksheet_summary_de"],
            row["drafting_summary_de"],
            row["table_figure_de"],
            row["gate_short_de"],
        )
        for row in worksheet_drafting_bridge["chapter_rows"]
    ]
    table = _add_table(
        doc,
        rows,
        ["Kapitel", "Worksheets", "Drafting", "Tabelle/Figur", "Gate"],
        [650, 2150, 1700, 1200, 3660],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Agenten-Grenze",
        worksheet_drafting_bridge["future_agent_boundary_de"],
    )


def _add_next_work_section(doc: Document, next_work: dict[str, Any]) -> None:
    doc.add_heading("Naechste Arbeitsschritte", level=1)
    doc.add_paragraph(
        f"Der Next-Work-Plan ordnet {next_work['row_count']} Workstreams. "
        f"Erste Prioritaet ist `{next_work['first_workstream']}`, letzte "
        f"QA-Prioritaet ist `{next_work['final_workstream']}`."
    )
    rows = [
        (
            row["priority_order"],
            row["workstream"],
            row["next_action"],
            row["guardrail"],
        )
        for row in next_work["rows"]
    ]
    table = _add_table(
        doc,
        rows,
        ["Prioritaet", "Workstream", "Naechste Aktion", "Guardrail"],
        [900, 2100, 3460, 2900],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Arbeitslogik",
        (
            "Die naechsten Schritte beginnen mit Source Review und Kapiteldraft. "
            "Swiss, Monitor und Agenten bleiben an ihre jeweiligen Gates gebunden."
        ),
    )


def _add_execution_checklist_section(
    doc: Document,
    execution_checklist: dict[str, Any],
) -> None:
    doc.add_heading("Kapitelweise Umsetzungscheckliste", level=1)
    doc.add_paragraph(
        f"Die Execution-Checkliste uebersetzt die Highlevel-View in "
        f"{execution_checklist['row_count']} Kapitelaufgaben. Erste Aufgabe "
        f"ist `{execution_checklist['first_task']}`, letzte Aufgabe ist "
        f"`{execution_checklist['final_task']}`. Sie ist kein neues "
        "empirisches Ergebnis, sondern eine Schreib- und Abnahmelogik fuer "
        "den naechsten Entwurf."
    )
    rows = [
        (
            row["task_id"],
            row["chapter_title"],
            row["draft_action_de"],
            row["done_when_de"],
            row["advisor_question_ids"],
        )
        for row in execution_checklist["rows"]
    ]
    table = _add_table(
        doc,
        rows,
        ["Task", "Kapitel", "Schreibaktion", "Fertig wenn", "Advisor-Fragen"],
        [900, 1700, 3000, 2800, 960],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Grenze",
        (
            "Diese Checkliste reaktiviert keinen Review-Access und keine "
            "Agenten. Sie steuert nur, wie der bestehende H1-H3-Kern mit "
            "Quellenreview, Tabellen/Figuren und Advisor-Feedback in "
            "Thesis-Prosa ueberfuehrt wird."
        ),
    )


def _add_research_design_section(doc: Document, data: dict[str, Any]) -> None:
    doc.add_heading("1. Forschungsfrage und Aufbau der Bachelorarbeit", level=1)
    doc.add_paragraph(
        "Die Arbeit untersucht, ob und wie Polymarket als dezentraler "
        "Prognosemarkt Informationen in politischen Maerkten verarbeitet. "
        "Informationelle Effizienz wird dabei nicht direkt behauptet, sondern "
        "ueber beobachtbare, reproduzierbare Proxies geprueft: Forecast-"
        "Qualitaet, Event-Reaktion und Wallet-Timing."
    )
    rows = [
        (
            "Leitfrage",
            "Inwiefern bilden Polymarket-Preise Informationen waehrend politischer Ereignisse ab, wie schneiden sie gegen traditionelle Prognosequellen ab, und ob aggregierte Wallet-Aktivitaet fruehe Timing-Signale zeigt?",
        ),
        (
            "H1",
            "Forecast-Qualitaet: Brier-Verluste und Vergleichstests zwischen Polymarket, FiveThirtyEight, 270toWin/Rieke und dokumentierten poll-derived Probability-Transformationen.",
        ),
        (
            "H2",
            "Event-Window-Reaktion: taegliche Preisbewegungen um vorab kuratierte politische Ereignisse.",
        ),
        (
            "H3",
            "Wallet-Tier-Timing: verteilungsbasierte Wallet-Tiers, Lead-Lag-Korrelationen und Granger-Diagnostik.",
        ),
        (
            "Erweiterung",
            "Read-only Politics/Geo Monitor mit deterministischer Review-Queue sowie separater Swiss-Referendum-Vergleich als aktueller Datensammlungs-Track.",
        ),
    ]
    table = _add_table(doc, rows, ["Baustein", "Inhalt"], [1700, 7660])
    _shade_table_header(table)
    _add_callout(
        doc,
        "Aufbau wie in einer BA",
        (
            "Der Bericht ist deshalb nicht nur ein Projektlog, sondern folgt "
            "der spaeteren Thesis-Logik: Einleitung, Theorie, Methodik, "
            "Datenbasis, empirische Ergebnisse, Diskussion und Ausblick."
        ),
    )


def _add_literature_section(
    doc: Document,
    literature: dict[str, Any],
    source_review: dict[str, Any],
) -> None:
    doc.add_heading("2. Wissenschaftlicher Quellenrahmen", level=1)
    doc.add_paragraph(
        f"Der lokale Literaturindex umfasst {literature['source_count']} Quellen. "
        f"Fuer diesen Dozentenbericht werden {literature['selected_source_count']} "
        "wissenschaftlich relevante Kernquellen als Rahmen verwendet. "
        f"Statusverteilung: {literature['status_counts_text']}."
    )
    rows = [
        (
            f"{source['source_id']}: {source['authors']} ({source['year']})",
            source["title"],
            f"{source['role']} {source['research_note']}",
            source["status"],
        )
        for source in literature["sources"]
    ]
    table = _add_table(
        doc,
        rows,
        ["Quelle", "Titel", "Rolle und Beitrag", "Status"],
        [1700, 2800, 3560, 1300],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Quellenreview-Worksheet",
        (
            f"Das Worksheet enthaelt {source_review['worksheet_rows']} manuelle "
            f"Review-Zeilen, davon {source_review['priority_1_rows']} "
            f"Priority-1-Methodenquellen und {source_review['blocked_rows']} "
            "blockierte oder Future-Work-Quelle. Alle Reviewer-Entscheide "
            "bleiben pending."
        ),
    )
    _add_callout(doc, "Zitationsgrenze", literature["citation_boundary"])


def _add_methodology_section(doc: Document, data: dict[str, Any]) -> None:
    doc.add_heading("3. Methodisches Design und Datenbasis", level=1)
    doc.add_paragraph(
        "Die empirische Logik ist bewusst konservativ. Alle statistischen "
        "Kennzahlen werden in Python berechnet und als CSV-, JSON-, PNG- oder "
        "HTML-Artefakte abgelegt. LLMs oder Agenten duerfen die Metriken nicht "
        "berechnen; sie duerfen spaeter nur vorab berechnete, bounded Outputs "
        "interpretieren und muessen auditiert werden."
    )
    rows = [
        (
            "Datenquellen",
            f"SQLite-Datenbank mit {data['project']['database']['table_count']} Tabellen, kuratierte Event-CSV, Polymarket-Preise, traditionelle Forecast-/Polling-Quellen und {data['source_counts']['literature_rows']} Literaturindex-Zeilen.",
        ),
        (
            "Validierung",
            "Schema- und Output-Tests, Projekt-Guardrails, max. 50 Zeilen fuer spaetere Tool-Abfragen, keine Raw-Table-Prompts.",
        ),
        (
            "H1-Metriken",
            "Brier Score, Lower-Loss-Zaehler, Diebold-Mariano-Vergleiche, Kalibrierungsdiagnostik und Scope-/Claim-Audits.",
        ),
        (
            "H2-Metriken",
            "Vorab festgelegte Daily-Event-Windows; keine Intraday-Speed-Claims ohne Intraday-Daten.",
        ),
        (
            "H3-Metriken",
            "Wallet-Tiers aus beobachteter Verteilung, Lead-Lag-Korrelationen und Granger-Diagnostik ohne Kausalclaim.",
        ),
        (
            "Live-/Monitor-Regeln",
            "Read-only, public data only, filebasierte bounded Outputs, keine Orders, keine Credentials, keine Profit- oder Tradingbehauptung.",
        ),
    ]
    table = _add_table(doc, rows, ["Teil", "Begruendung"], [1700, 7660])
    _shade_table_header(table)


def _add_interpretation_section(doc: Document, data: dict[str, Any]) -> None:
    doc.add_heading("4. Zentrale Erkenntnisse, Begruendung und Interpretation", level=1)
    doc.add_paragraph(
        "Dieser Abschnitt ist die eigentliche Ergebnissynthese fuer den "
        "Dozentenbericht. Er trennt Rohbefund, Interpretation und Grenze, "
        "damit die Arbeit wissenschaftlich argumentiert und nicht nur Outputs "
        "auflistet."
    )
    rows = [
        (
            row["bereich"],
            f"{row['erkenntnis']} Evidenz: {row['evidenz']}",
            f"{row['interpretation']} Grenze: {row['grenze']}",
        )
        for row in _interpretation_rows(data)
    ]
    table = _add_table(
        doc,
        rows,
        ["Bereich", "Erkenntnis und Evidenz", "Interpretation und Grenze"],
        [1600, 3860, 3900],
    )
    _shade_table_header(table)
    method_rows = list(_method_decision_rows())
    table = _add_table(
        doc,
        method_rows,
        ["Methodische Entscheidung", "Begruendung", "Konsequenz"],
        [2100, 3860, 3400],
    )
    _shade_table_header(table)
    _add_callout(
        doc,
        "Gesamtinterpretation",
        (
            "Der aktuelle Stand ist keine einfache Ja/Nein-Antwort auf "
            "informationelle Effizienz. Die Arbeit zeigt vielmehr, in welchen "
            "Scopes Polymarket stark ist, wo traditionelle Vergleichssignale "
            "besser bleiben, und warum Wallet- und Monitorbefunde als "
            "diagnostische Review-Ebene formuliert werden muessen."
        ),
    )


def _add_project_overview(doc: Document, data: dict[str, Any]) -> None:
    doc.add_heading("5. Projektlandkarte und Stand der Umsetzung", level=1)
    doc.add_paragraph(
        "Die Repository-Struktur trennt Forschungssteuerung, Daten, deterministische "
        "Analyse, Collector-Code, Tests, Dokumentation und geparkte Legacy-Agenten. "
        "Diese Trennung ist wichtig, weil die Bachelorarbeit nicht auf frei "
        "formulierter KI-Interpretation basiert, sondern auf reproduzierbaren "
        "Python-Outputs."
    )
    db = data["project"]["database"]
    rows = [
        ("Datenbank", f"{db['table_count']} Tabellen, u.a. polymarket_prices, poll_forecasts, whale_trades."),
        ("Datenartefakte", f"{data['project']['folder_inventory'].get('data/results', 0)} Dateien unter data/results."),
        ("Analysecode", f"{data['project']['folder_inventory'].get('operations/analysis', 0)} Module fuer H1-H3, Monitor und Swiss-Analyse."),
        ("Collector-Code", f"{data['project']['folder_inventory'].get('operations/collectors', 0)} read-only oder lokale Collector/Refresh-Module."),
        ("Tests", f"{data['project']['folder_inventory'].get('tests', 0)} Testdateien; {data['project']['test_summary']}."),
        ("Legacy/Agenten", "Agent- und MCP-Einstiege bleiben guard-railed und sind nicht Teil der aktiven Laufzeit."),
    ]
    table = _add_table(doc, rows, ["Bereich", "Bedeutung"], [1900, 7460])
    _shade_table_header(table)
    doc.add_paragraph(
        "Wichtige Steuerdateien sind AGENTS.md, GOAL.md, ROADMAP.md und STATUS.md. "
        "Sie definieren Scope, aktive Phase, Sicherheitsregeln, Teststatus und "
        "naechste Commits."
    )


def _add_h1_section(doc: Document, h1: dict[str, Any]) -> None:
    doc.add_heading("6. H1 - Forecast-Qualitaet", level=1)
    doc.add_paragraph(
        "H1 beantwortet die Frage, ob Polymarket im untersuchten Ueberlappungsfenster "
        "eine niedrigere Prognoseabweichung zeigt als vergleichbare traditionelle "
        "Probability-Forecasts. Der Zweck war nicht, Reaktionsgeschwindigkeit zu "
        "messen, sondern Forecast-Qualitaet."
    )
    rows = [
        ("Warum", "Vergleich der Prognosequalitaet zwischen Polymarket, FiveThirtyEight und einfachen Baselines."),
        ("Methode", "Brier Score und Diebold-Mariano-Vergleiche aus Python-Outputs."),
        ("Datenform", f"{h1['observation_count']} taegliche Ueberlappungsbeobachtungen."),
        ("Ergebnis", f"Polymarket Mean Brier {h1['brier_polymarket']:.4f}; FiveThirtyEight {h1['brier_fivethirtyeight']:.4f}; 50-Prozent-Baseline {h1['brier_always_50']:.4f}."),
        ("Head-to-head", f"Polymarket niedrigerer Tagesverlust in {h1['pm_better_vs_538_count']} von {h1['pm_vs_538_count']} Tagen ({h1['pm_better_vs_538_share'] * 100:.1f}%)."),
        ("H1 Synthesis", f"{h1['synthesis_aggregate_support_count']} von {h1['synthesis_evidence_row_count']} Vergleichszeilen stuetzen Polymarket im mittleren Brier; {h1['synthesis_majority_support_count']} von {h1['synthesis_evidence_row_count']} zeigen eine Mehrheit niedrigerer Einzelfallverluste; breiter Viele-Faelle-Beweis {h1['synthesis_broad_support_count']} von {h1['synthesis_evidence_row_count']}."),
        ("Claim Evidence Audit", f"{h1['claim_audit_support_row_count']} von {h1['claim_audit_row_count']} Audit-Zeilen stuetzen Polymarket begrenzt; {h1['claim_audit_contradiction_row_count']} widerspricht dem starken Claim; direkt pollbezogen {h1['claim_audit_direct_poll_support_row_count']} von {h1['claim_audit_direct_poll_row_count']} stuetzend; breiter User-Claim belegt {h1['claim_audit_broad_user_claim_proven']}."),
        ("Poll Comparison Result", f"Primaerer <=90-Tage-Low/Middle-Poll-Distanz-Scope: Polymarket {h1['poll_result_primary_pm_count']} von {h1['poll_result_primary_row_count']} State-Date-Zeilen ({h1['poll_result_primary_pm_share'] * 100:.1f}%), poll-derived {h1['poll_result_primary_poll_count']}; State-Ebene Polymarket {h1['poll_result_primary_pm_state_count']} von {h1['poll_result_primary_state_count']}, exakter einseitiger p-Wert {h1['poll_result_primary_p_value']:.4f}, 95-Prozent-Untergrenze {h1['poll_result_primary_ci_low']:.3f}. Direkt pollbezogen {h1['poll_result_direct_poll_support_count']} von {h1['poll_result_direct_poll_row_count']} Audit-Zeilen stuetzend; Vollpanel-Gegenbeleg poll-derived {h1['poll_result_full_panel_poll_count']} von {h1['poll_result_full_panel_row_count']}; Status {h1['poll_result_goal_status']}."),
        ("Poll Claim Readiness", f"{h1['poll_claim_supported_bounded_count']} von {h1['poll_claim_row_count']} Claim-Zeilen stuetzen den bounded <=90-Tage Low/Middle-Poll-Distanz-Scope; {h1['poll_claim_counterexample_count']} Gegenbeispiel-Scopes und {h1['poll_claim_mixed_mean_count']} Mean-Loss-Stuetze-ohne-Mehrheit-Zeilen bleiben als Grenzen. Bounded Scope: Polymarket {h1['poll_claim_primary_pm_count']} von {h1['poll_claim_primary_count']} State-Date-Zeilen ({h1['poll_claim_primary_pm_share'] * 100:.1f}%) und {h1['poll_claim_state_month_pm_count']} von {h1['poll_claim_state_month_count']} State-Month-Einheiten, exact p={h1['poll_claim_state_month_p_value']:.2g}, 95-Prozent-Untergrenze {h1['poll_claim_state_month_ci_low']:.3f}. Bounded Claim supported {h1['poll_claim_bounded_supported']}; breiter Claim belegt {h1['poll_claim_broad_proven']}; Status {h1['poll_claim_goal_status']}."),
        ("Poll Scope Frontier", f"{h1['poll_frontier_robust_scope_count']} von {h1['poll_frontier_row_count']} Horizont-x-Poll-Distanz-Scopes erfuellen die robuste Regel. Groesster robuster Scope: {h1['poll_frontier_largest_horizon']} + {h1['poll_frontier_largest_tier']}, Polymarket {h1['poll_frontier_largest_pm_count']} von {h1['poll_frontier_largest_row_count']} State-Date-Zeilen ({h1['poll_frontier_largest_pm_share'] * 100:.1f}%), {h1['poll_frontier_largest_state_month_pm_count']} von {h1['poll_frontier_largest_state_month_count']} State-Month-Einheiten, exact p={h1['poll_frontier_largest_state_month_p_value']:.3g}. Staerkster Scope {h1['poll_frontier_strongest_scope_id']}: {h1['poll_frontier_strongest_row_count']} Zeilen, p={h1['poll_frontier_strongest_p_value']:.2g}. <=90 Tage alle Distanzen: Polymarket {h1['poll_frontier_lte_90_all_pm_count']} von {h1['poll_frontier_lte_90_all_row_count']} Zeilen ({h1['poll_frontier_lte_90_all_pm_share'] * 100:.1f}%), State-Month p={h1['poll_frontier_lte_90_all_state_month_p_value']:.3g}; Vollpanel-Gegenbeleg poll-derived {h1['poll_frontier_full_panel_poll_count']} von {h1['poll_frontier_full_panel_row_count']}; Status {h1['poll_frontier_goal_status']}."),
        ("Poll Decision Matrix", f"{h1['poll_decision_robust_yes_count']} von {h1['poll_decision_row_count']} Entscheidungszeilen sind robuste bounded-Yes-Zeilen; {h1['poll_decision_mixed_mean_count']} Mean-Loss-Stuetze-ohne-Mehrheit-Zeilen und {h1['poll_decision_counterexample_count']} Gegenbelege bleiben als Grenzen. Groesster robuster Scope: Polymarket {h1['poll_decision_largest_pm_count']} von {h1['poll_decision_largest_row_count']} State-Date-Zeilen ({h1['poll_decision_largest_pm_share'] * 100:.1f}%), {h1['poll_decision_largest_state_month_pm_count']} von {h1['poll_decision_largest_state_month_count']} State-Month-Einheiten, p={h1['poll_decision_largest_p_value']:.4f}. Kalibrierungskontext: {h1['poll_decision_calibration_aggregate_count']} von {h1['poll_decision_calibration_pairwise_count']} Pairwise-Reihen stuetzen Polymarket im mittleren Brier; {h1['poll_decision_calibration_majority_count']} von {h1['poll_decision_calibration_pairwise_count']} auch per Fallmehrheit. Bounded ready {h1['poll_decision_bounded_ready']}; breiter Claim {h1['poll_decision_broad_proven']}; Status {h1['poll_decision_goal_status']}."),
        ("Robust Poll-Scope Quality", f"{h1['robust_quality_forecast_row_count']} Forecast-Zeilen aus {h1['robust_quality_case_count']} State-Date-Faellen und {h1['robust_quality_scope_count']} robusten Poll-Scopes. Groesster robuster Scope: Polymarket {h1['robust_quality_largest_pm_count']} von {h1['robust_quality_largest_case_count']} Zeilen ({h1['robust_quality_largest_pm_share'] * 100:.1f}%), Mean Brier {h1['robust_quality_largest_pm_brier']:.4f} vs poll-derived {h1['robust_quality_largest_poll_brier']:.4f}, ECE {h1['robust_quality_largest_pm_ece']:.4f} vs {h1['robust_quality_largest_poll_ece']:.4f}, Probability-Separation {h1['robust_quality_largest_pm_separation']:.4f} vs {h1['robust_quality_largest_poll_separation']:.4f}. Staerkster robuster Scope: Polymarket {h1['robust_quality_strongest_pm_count']} von {h1['robust_quality_strongest_case_count']} Zeilen ({h1['robust_quality_strongest_pm_share'] * 100:.1f}%), Mean Brier {h1['robust_quality_strongest_pm_brier']:.4f} vs {h1['robust_quality_strongest_poll_brier']:.4f}, ECE {h1['robust_quality_strongest_pm_ece']:.4f} vs {h1['robust_quality_strongest_poll_ece']:.4f}. Der staerkste Scope hat positive Rate {h1['robust_quality_strongest_positive_rate']:.1f}; Separation ist dort nicht definiert. Breiter Claim belegt {h1['robust_quality_broad_claim_proven']}."),
        ("Robust Poll-Scope Unit Quality", f"{h1['robust_unit_summary_row_count']} Aggregationszeilen ueber robuste Poll-Scopes. Groesster robuster Scope: Polymarket {h1['robust_unit_largest_state_pm_count']} von {h1['robust_unit_largest_state_count']} States, {h1['robust_unit_largest_state_month_pm_count']} von {h1['robust_unit_largest_state_month_count']} State-Month-Einheiten und {h1['robust_unit_largest_state_horizon_pm_count']} von {h1['robust_unit_largest_state_horizon_count']} State-Horizon-Einheiten; State-Month p={h1['robust_unit_largest_state_month_p_value']:.4f}, 95-Prozent-Untergrenze {h1['robust_unit_largest_state_month_ci_low']:.3f}. Staerkster robuster Scope: Polymarket {h1['robust_unit_strongest_state_pm_count']} von {h1['robust_unit_strongest_state_count']} States, {h1['robust_unit_strongest_state_month_pm_count']} von {h1['robust_unit_strongest_state_month_count']} State-Month-Einheiten und {h1['robust_unit_strongest_state_horizon_pm_count']} von {h1['robust_unit_strongest_state_horizon_count']} State-Horizon-Einheiten; State-Month p={h1['robust_unit_strongest_state_month_p_value']:.2g}, 95-Prozent-Untergrenze {h1['robust_unit_strongest_state_month_ci_low']:.3f}. Medianer State-Month-Brier-Vorteil {h1['robust_unit_largest_state_month_median_advantage']:.4f} im groessten und {h1['robust_unit_strongest_state_month_median_advantage']:.4f} im staerksten Scope. Breiter Claim belegt {h1['robust_unit_broad_claim_proven']}."),
        ("Poll Comparison Unit Robustness", f"Primaerer Scope nach Aggregation: Polymarket {h1['poll_unit_state_pm_count']} von {h1['poll_unit_state_count']} States, {h1['poll_unit_state_month_pm_count']} von {h1['poll_unit_state_month_count']} State-Month-Einheiten, {h1['poll_unit_state_horizon_pm_count']} von {h1['poll_unit_state_horizon_count']} State-Horizon-Einheiten und {h1['poll_unit_horizon_tier_pm_count']} von {h1['poll_unit_horizon_tier_count']} Horizon-Tier-Einheiten. State-Month exact p={h1['poll_unit_state_month_p_value']:.2g}, 95-Prozent-Untergrenze {h1['poll_unit_state_month_ci_low']:.3f}; Full-Panel-State-Month-Gegenbeleg: poll-derived {h1['poll_unit_full_panel_state_month_poll_count']} von {h1['poll_unit_full_panel_state_month_count']}; Late-High-Distance-State-Month-Gegenbeleg: poll-derived {h1['poll_unit_late_high_state_month_poll_count']} von {h1['poll_unit_late_high_state_month_count']}, exact p={h1['poll_unit_late_high_state_month_poll_p_value']:.4f}; Status {h1['poll_unit_goal_status']}."),
        ("Direct Poll Loss Decomposition", f"Direkte Poll-Transform-Vergleiche: Mean Brier Polymarket {h1['direct_poll_loss_pm_brier']:.4f} vs poll-derived {h1['direct_poll_loss_poll_brier']:.4f}; Polymarket niedrigerer Verlust in {h1['direct_poll_loss_pm_count']} von {h1['direct_poll_loss_case_count']} Source-State-Faellen, poll-derived in {h1['direct_poll_loss_poll_count']}. Polymarket-Gewinnfaelle haben mittleren Brier-Vorteil {h1['direct_poll_loss_pm_win_mean_advantage']:.4f}, poll-derived Gewinnfaelle {h1['direct_poll_loss_poll_win_mean_advantage']:.4f}; Total-Margin-Ratio {h1['direct_poll_loss_margin_ratio']:.1f}."),
        ("Direct Poll State Cluster", f"State-Cluster-Diagnostik ueber {h1['direct_poll_state_cluster_state_count']} States: gleichgewichteter mittlerer Verlustvorteil {h1['direct_poll_state_cluster_mean_advantage']:.4f}, Bootstrap-95%-Intervall {h1['direct_poll_state_cluster_bootstrap_ci_low']:.4f} bis {h1['direct_poll_state_cluster_bootstrap_ci_high']:.4f}, Sign-Flip-p={h1['direct_poll_state_cluster_sign_flip_p']:.4f}. State-Mehrheit: Polymarket {h1['direct_poll_state_cluster_pm_state_count']} States, poll-derived {h1['direct_poll_state_cluster_poll_state_count']}; Polymarket-State-Mehrheit belegt {h1['direct_poll_state_cluster_majority_supports_pm']}."),
        ("Direct Poll Outlier Robustness", f"Outlier-Diagnostik ueber {h1['direct_poll_outlier_state_count']} State-Cluster: voller Mean {h1['direct_poll_outlier_full_mean_advantage']:.4f}; alle Leave-one-state-out Means positiv {h1['direct_poll_outlier_leave_one_all_positive']}, Minimum {h1['direct_poll_outlier_min_leave_one_mean']:.4f} ohne {h1['direct_poll_outlier_most_influential_state']}. Entfernt man die groessten positiven State-Beitraege, bleibt der Mean bis {h1['direct_poll_outlier_top_k_positive']} entfernte States positiv und kippt bei {h1['direct_poll_outlier_first_nonpositive_k']} entfernten States auf {h1['direct_poll_outlier_first_nonpositive_mean']:.4f}. Groesster positiver State: {h1['direct_poll_outlier_largest_positive_state']} ({h1['direct_poll_outlier_largest_positive_advantage']:.4f}); Status {h1['direct_poll_outlier_goal_status']}."),
        ("State-Source-Konsens", f"{h1['state_source_consensus_case_count']} Source-State-Vergleiche ueber {h1['state_source_consensus_state_count']} States: Polymarket niedrigerer Verlust in {h1['state_source_consensus_pm_case_count']} Source-State-Faellen, Comparatoren in {h1['state_source_consensus_comparator_case_count']}. All-Source-State-Konsens: Polymarket {h1['state_source_consensus_pm_state_count']} States, Comparatoren {h1['state_source_consensus_comparator_state_count']}, Ties {h1['state_source_consensus_tie_state_count']}. Bei States mit zwei direkten Poll-Transform-Quellen gewinnt Polymarket {h1['state_source_consensus_direct_two_pm_state_count']} von {h1['state_source_consensus_direct_two_state_count']} States."),
        ("Competitive-State-Diagnose", f"Niedrigste Comparator-Distanz-Terzile: Polymarket {h1['competitive_state_all_low_pm_count']} von {h1['competitive_state_all_low_case_count']} All-Source-Faellen und {h1['competitive_state_direct_low_pm_count']} von {h1['competitive_state_direct_low_case_count']} direkten Poll-Transform-Faellen; hoechste Distanz-Terzile: Polymarket {h1['competitive_state_all_high_pm_count']} von {h1['competitive_state_all_high_case_count']}, Comparatoren {h1['competitive_state_all_high_comparator_count']} von {h1['competitive_state_all_high_case_count']}. Begrenzte Competitive-State-Ausnahme, kein breiter Viele-Faelle-Beweis."),
        ("State-Date Competitiveness x Horizon", f"<=90 Tage und Low/Middle-Poll-Distanz: Polymarket {h1['panel_comp_late_non_safe_pm_count']} von {h1['panel_comp_late_non_safe_row_count']} State-Date-Zeilen und {h1['panel_comp_late_non_safe_state_support_count']} von {h1['panel_comp_late_non_safe_state_count']} States; spaete High-Distance-Zeilen: Polymarket {h1['panel_comp_late_high_pm_count']} von {h1['panel_comp_late_high_row_count']}, poll-derived {h1['panel_comp_late_high_poll_count']} von {h1['panel_comp_late_high_row_count']}. Starker spaeter Competitive-Poll-Befund, aber kein unabhaengiger Viele-Wahlen-Beweis."),
        ("State-Level Signifikanzdiagnose", f"Spaete Low/Middle-Poll-Distanz: Polymarket {h1['state_sign_late_non_safe_pm_state_count']} von {h1['state_sign_late_non_safe_state_count']} States; exakter einseitiger Binomial-p-Wert {h1['state_sign_late_non_safe_p_value']:.4f}; exakte 95-Prozent-Untergrenze {h1['state_sign_late_non_safe_ci_low']:.3f}. Spaete High-Distance-States: poll-derived {h1['state_sign_late_high_poll_state_count']} von {h1['state_sign_late_high_state_count']} States."),
        ("Kalibrierungsdiagnostik", f"{h1['calibration_forecast_case_rows']} Forecast-Case-Zeilen aus {h1['calibration_forecast_source_count']} Quellen und {h1['calibration_pairwise_count']} Pairwise-Reihen: {h1['calibration_aggregate_support_count']} von {h1['calibration_pairwise_count']} mit niedrigerem mittleren Polymarket-Brier, {h1['calibration_majority_support_count']} von {h1['calibration_pairwise_count']} mit Mehrheit niedrigerer Einzelfallverluste, breiter Viele-Faelle-Beweis {h1['calibration_broad_support_count']} von {h1['calibration_pairwise_count']}."),
        ("50-State-Kalibrierung", f"Polymarket Mean Brier {h1['calibration_pm_state_brier']:.4f} und Fixed-Bin-ECE {h1['calibration_pm_state_ece']:.4f}; Rieke ECE {h1['calibration_rieke_state_ece']:.4f}; 270toWin/JHK ECE {h1['calibration_270_state_ece']:.4f}. Forecast-Qualitaets-, aber kein klarer Kalibrierungssieg."),
        ("Final Snapshot", f"{h1['final_snapshot_case_count']} geloeste 2024-Final-Snapshot-Outcomes gegen 538 final forecast: Polymarket niedrigerer Verlust in {h1['final_snapshot_pm_lower_loss_count']} von {h1['final_snapshot_case_count']} Faellen; Mean Brier {h1['final_snapshot_mean_pm_brier']:.4f} vs {h1['final_snapshot_mean_traditional_brier']:.4f}."),
        ("State Poll Snapshot", f"{h1['state_poll_snapshot_case_count']} geloeste State-Outcomes gegen dokumentiert transformierte 538 Polling-Averages: Polymarket niedrigerer Verlust in {h1['state_poll_snapshot_pm_lower_loss_count']} von {h1['state_poll_snapshot_case_count']} Faellen; Mean Brier {h1['state_poll_snapshot_mean_pm_brier']:.4f} vs {h1['state_poll_snapshot_mean_poll_brier']:.4f}."),
        ("270toWin Polling Average", f"{h1['two_seventy_poll_average_case_count']} gematchte State-Outcomes gegen dokumentiert transformierte 270toWin-Polling-Averages: Polymarket niedrigerer Verlust in {h1['two_seventy_poll_average_pm_lower_loss_count']} Faellen, poll-derived in {h1['two_seventy_poll_average_poll_lower_loss_count']}; Mean Brier {h1['two_seventy_poll_average_mean_pm_brier']:.4f} vs {h1['two_seventy_poll_average_mean_poll_brier']:.4f}."),
        ("Popular Vote", f"{h1['popular_vote_case_count']} nationale Tageszeilen fuer Trump popular vote: Polymarket niedrigerer Verlust in {h1['popular_vote_pm_lower_loss_count']} Zeilen, poll-derived in {h1['popular_vote_poll_lower_loss_count']}; Mean Brier {h1['popular_vote_mean_pm_brier']:.4f} vs {h1['popular_vote_mean_poll_brier']:.4f}. Gegenbeleg zum starken Claim."),
        ("Margin Threshold Readiness", f"{h1['margin_threshold_candidate_count']} Trump-State-Margin-Maerkte geprueft: {h1['margin_threshold_with_538_poll_count']} mit 538-State-Poll-Average-Zeilen, {h1['margin_threshold_with_clob_overlap_count']} mit CLOB-Historie im bewahrten 538-Fenster, {h1['margin_threshold_compatible_count']} neue H1-Brier-Faelle; {h1['margin_threshold_no_overlap_count']} durch fehlende zeitliche Ueberlappung blockiert, {h1['margin_threshold_missing_poll_count']} durch fehlende 538-State-Polls."),
        ("State-Date Poll Panel", f"{h1['state_poll_panel_case_count']} gematchte State-Date-Zeilen ueber {h1['state_poll_panel_state_count']} States und {h1['state_poll_panel_date_count']} Daten: Polymarket niedrigerer Verlust in {h1['state_poll_panel_pm_lower_loss_count']} Zeilen, poll-derived niedrigerer Verlust in {h1['state_poll_panel_poll_lower_loss_count']}; Mean Brier {h1['state_poll_panel_mean_pm_brier']:.4f} vs {h1['state_poll_panel_mean_poll_brier']:.4f}."),
        ("State-Date Temporal Diagnostic", f"Polymarket-stuetzende Monate {h1['state_poll_temporal_support_months']}: {h1['state_poll_temporal_support_pm_lower_loss_count']} von {h1['state_poll_temporal_support_row_count']} Zeilen mit niedrigerem Polymarket-Verlust ueber {h1['state_poll_temporal_support_state_count']} States; poll-derived niedrigerer Verlust in {h1['state_poll_temporal_support_poll_lower_loss_count']}; Mean Brier {h1['state_poll_temporal_support_mean_pm_brier']:.4f} vs {h1['state_poll_temporal_support_mean_poll_brier']:.4f}."),
        ("State-Date Horizon Diagnostic", f"<=90-Tage-Fenster ({h1['state_poll_horizon_near_bins']}): {h1['state_poll_horizon_near_pm_lower_loss_count']} von {h1['state_poll_horizon_near_row_count']} Zeilen mit niedrigerem Polymarket-Verlust ueber {h1['state_poll_horizon_near_state_count']} States; poll-derived niedrigerer Verlust in {h1['state_poll_horizon_near_poll_lower_loss_count']}; Mean Brier {h1['state_poll_horizon_near_mean_pm_brier']:.4f} vs {h1['state_poll_horizon_near_mean_poll_brier']:.4f}."),
        ("State-Level Horizon Support", f"Im <=90-Tage-Fenster stuetzt Polymarket {h1['state_poll_horizon_state_pm_mean_support_count']} von {h1['state_poll_horizon_state_count']} States nach mittlerem Brier und {h1['state_poll_horizon_state_pm_majority_support_count']} von {h1['state_poll_horizon_state_count']} States nach Mehrheit niedrigerer Tagesverluste; {h1['state_poll_horizon_state_poll_support_count']} States stuetzen Polymarket nicht."),
        ("<=90-Day Score Quality", f"{h1['state_poll_near_quality_forecast_row_count']} Forecast-Zeilen aus {h1['state_poll_near_quality_case_count']} State-Date-Faellen und zwei Quellen: Polymarket Mean Brier {h1['state_poll_near_quality_pm_mean_brier']:.4f} vs poll-derived {h1['state_poll_near_quality_poll_mean_brier']:.4f}; Fixed-Bin-ECE {h1['state_poll_near_quality_pm_ece']:.4f} vs {h1['state_poll_near_quality_poll_ece']:.4f}; Probability-Separation {h1['state_poll_near_quality_pm_separation']:.4f} vs {h1['state_poll_near_quality_poll_separation']:.4f}."),
        ("Poll Transform Sensitivitaet", f"MAE {h1['state_poll_sensitivity_min_mae']:.1f} bis {h1['state_poll_sensitivity_max_mae']:.1f} Prozentpunkte: Polymarket bleibt in allen {h1['state_poll_sensitivity_row_count']} Parameterzeilen im mittleren Brier niedriger; Lower-Loss-Spanne {h1['state_poll_sensitivity_min_pm_lower_loss_count']} bis {h1['state_poll_sensitivity_max_pm_lower_loss_count']} von {h1['state_poll_snapshot_case_count']} State-Outcomes."),
        ("Coverage Audit", f"{h1['state_poll_coverage_state_count']} US-States geprueft; {h1['state_poll_coverage_polymarket_market_count']} Polymarket-State-Maerkte; {h1['state_poll_coverage_valid_pair_count']} valide H1-Brier-Paare mit REP/DEM-Zeilen im 538-Snapshot; {h1['state_poll_coverage_missing_poll_count']} States wegen fehlender 538-Snapshot-Pollwerte ausgeschlossen."),
        ("Rieke 50-State Forecast", f"{h1['rieke_state_case_count']} geloeste State-Outcomes gegen ein unabhaengiges pollbasiertes Rieke-Modell: Mean Brier {h1['rieke_state_mean_pm_brier']:.4f} vs {h1['rieke_state_mean_rieke_brier']:.4f}; Polymarket niedrigerer Einzelfallverlust in {h1['rieke_state_pm_lower_loss_count']} von {h1['rieke_state_case_count']}, Rieke in {h1['rieke_state_rieke_lower_loss_count']} von {h1['rieke_state_case_count']}."),
        ("270toWin/JHK 50-State Forecast", f"{h1['two_seventy_state_case_count']} geloeste State-Outcomes gegen 270toWin/JHK: {h1['two_seventy_state_exact_case_count']} exakt ausgewiesene Wahrscheinlichkeiten und {h1['two_seventy_state_censored_case_count']} zensierte >99.9-Prozent-Boundary-Werte; Mean Brier {h1['two_seventy_state_mean_pm_brier']:.4f} vs {h1['two_seventy_state_mean_270_brier']:.4f}; Polymarket niedrigerer Einzelfallverlust in {h1['two_seventy_state_pm_lower_loss_count']} von {h1['two_seventy_state_case_count']}, 270toWin/JHK in {h1['two_seventy_state_270_lower_loss_count']} von {h1['two_seventy_state_case_count']}."),
        ("Grenze", "H1 ist kein Speed-Test; die Tagespaare gehoeren zu einem geloesten Wahl-Outcome. Die neuen State-Faelle sind poll-derived beziehungsweise modellbasiert, keine Rohpolls und kein RCP-Vergleich."),
    ]
    table = _add_table(doc, rows, ["Frage", "Antwort"], [1700, 7660])
    _shade_table_header(table)
    _add_callout(
        doc,
        "Interpretation",
        (
            f"Im getesteten Fenster spricht H1 fuer tiefere Forecast-Verluste "
            f"von Polymarket gegenueber FiveThirtyEight. Der DM-p-Wert fuer "
            f"Polymarket vs FiveThirtyEight liegt bei {h1['dm_polymarket_vs_538']:.3g}. "
            f"Im Head-to-head-Vergleich hat Polymarket in "
            f"{h1['pm_better_vs_538_count']} von {h1['pm_vs_538_count']} "
            f"taeglichen Paaren den niedrigeren Brier-Verlust. "
            f"Die H1-Synthesis ueber alle aktuellen traditionellen "
            f"Vergleichszeilen zeigt {h1['synthesis_aggregate_support_count']} "
            f"von {h1['synthesis_evidence_row_count']} Zeilen mit niedrigerem "
            f"mittleren Polymarket-Brier, aber nur "
            f"{h1['synthesis_majority_support_count']} von "
            f"{h1['synthesis_evidence_row_count']} Zeilen mit Mehrheit "
            f"niedrigerer Einzelfallverluste und "
            f"{h1['synthesis_broad_support_count']} von "
            f"{h1['synthesis_evidence_row_count']} Zeilen, die den breiten "
            f"Viele-Faelle-Anspruch tragen. "
            f"Die fokussierte Poll-Comparison-Scorecard belegt die belastbarste "
            f"direkte Poll-Aussage: Polymarket hat im primaeren spaeten "
            f"Low/Middle-Poll-Distanz-Scope in "
            f"{h1['poll_result_primary_pm_count']} von "
            f"{h1['poll_result_primary_row_count']} State-Date-Zeilen "
            f"({h1['poll_result_primary_pm_share'] * 100:.1f}%) und "
            f"{h1['poll_result_primary_pm_state_count']} von "
            f"{h1['poll_result_primary_state_count']} States den niedrigeren "
            f"Brier-Verlust; der breite Zielstatus bleibt "
            f"{h1['poll_result_goal_status']}. "
            f"Die neue Unit-Robustness-Aggregation reduziert die Abhaengigkeit "
            f"von wiederholten Tageszeilen: Im selben primaeren Scope wird "
            f"Polymarket in {h1['poll_unit_state_month_pm_count']} von "
            f"{h1['poll_unit_state_month_count']} State-Month-Einheiten und "
            f"{h1['poll_unit_state_horizon_pm_count']} von "
            f"{h1['poll_unit_state_horizon_count']} State-Horizon-Einheiten "
            f"gestuetzt; fuer State-Month-Einheiten betraegt der exakte "
            f"einseitige p-Wert {h1['poll_unit_state_month_p_value']:.2g} "
            f"und die exakte 95-Prozent-Untergrenze "
            f"{h1['poll_unit_state_month_ci_low']:.3f}. "
            f"Die Grenzen bleiben sichtbar: Im Full Panel stuetzen "
            f"{h1['poll_unit_full_panel_state_month_poll_count']} von "
            f"{h1['poll_unit_full_panel_state_month_count']} State-Month-"
            f"Einheiten poll-derived, im Late-High-Distance-Scope "
            f"{h1['poll_unit_late_high_state_month_poll_count']} von "
            f"{h1['poll_unit_late_high_state_month_count']}. "
            f"Die neue Kalibrierungsdiagnostik nutzt "
            f"{h1['calibration_forecast_case_rows']} Forecast-Case-Zeilen und "
            f"zeigt fuer die 50-State-Fallreihe einen Polymarket-ECE von "
            f"{h1['calibration_pm_state_ece']:.4f} gegenueber "
            f"{h1['calibration_rieke_state_ece']:.4f} fuer Rieke und "
            f"{h1['calibration_270_state_ece']:.4f} fuer 270toWin/JHK; "
            f"das ist kein klarer Kalibrierungssieg. "
            f"Die kuratierte Final-Snapshot-Erweiterung zeigt zusaetzlich "
            f"{h1['final_snapshot_pm_lower_loss_count']} von "
            f"{h1['final_snapshot_case_count']} geloesten 2024-Outcomes mit "
            f"niedrigerem Polymarket-Verlust; der mittlere Brier liegt bei "
            f"{h1['final_snapshot_mean_pm_brier']:.4f} vs "
            f"{h1['final_snapshot_mean_traditional_brier']:.4f}. "
            f"Die State-Poll-Snapshot-Erweiterung fuegt "
            f"{h1['state_poll_snapshot_case_count']} State-Outcomes hinzu; "
            f"Polymarket hat dort in "
            f"{h1['state_poll_snapshot_pm_lower_loss_count']} Faellen den "
            f"niedrigeren Verlust, Mean Brier "
            f"{h1['state_poll_snapshot_mean_pm_brier']:.4f} vs "
            f"{h1['state_poll_snapshot_mean_poll_brier']:.4f}. "
            f"Die 270toWin-Polling-Average-Erweiterung erhoeht die direkt "
            f"pollbasierte State-Abdeckung auf "
            f"{h1['two_seventy_poll_average_case_count']} gematchte States; "
            f"Polymarket hat dort niedrigeren mittleren Brier "
            f"{h1['two_seventy_poll_average_mean_pm_brier']:.4f} vs "
            f"{h1['two_seventy_poll_average_mean_poll_brier']:.4f}, aber nur "
            f"{h1['two_seventy_poll_average_pm_lower_loss_count']} von "
            f"{h1['two_seventy_poll_average_case_count']} niedrigere "
            f"Einzelfallverluste. "
            f"Das groessere State-Date-Poll-Panel relativiert diesen "
            f"Snapshot-Befund: Es enthaelt "
            f"{h1['state_poll_panel_case_count']} gematchte Zeilen, davon "
            f"{h1['state_poll_panel_pm_lower_loss_count']} mit niedrigerem "
            f"Polymarket-Verlust und "
            f"{h1['state_poll_panel_poll_lower_loss_count']} mit niedrigerem "
            f"poll-derived Verlust; Mean Brier "
            f"{h1['state_poll_panel_mean_pm_brier']:.4f} vs "
            f"{h1['state_poll_panel_mean_poll_brier']:.4f}. "
            f"Die Temporal-Diagnose zeigt aber eine spaete Gegenbewegung: "
            f"In {h1['state_poll_temporal_support_months']} hat Polymarket "
            f"in {h1['state_poll_temporal_support_pm_lower_loss_count']} von "
            f"{h1['state_poll_temporal_support_row_count']} Zeilen den "
            f"niedrigeren Verlust, Mean Brier "
            f"{h1['state_poll_temporal_support_mean_pm_brier']:.4f} vs "
            f"{h1['state_poll_temporal_support_mean_poll_brier']:.4f}. "
            f"Die Forecast-Horizon-Diagnose zeigt fuer das <=90-Tage-Fenster "
            f"{h1['state_poll_horizon_near_pm_lower_loss_count']} von "
            f"{h1['state_poll_horizon_near_row_count']} Zeilen mit niedrigerem "
            f"Polymarket-Verlust, Mean Brier "
            f"{h1['state_poll_horizon_near_mean_pm_brier']:.4f} vs "
            f"{h1['state_poll_horizon_near_mean_poll_brier']:.4f}. "
            f"Auf State-Ebene sind es "
            f"{h1['state_poll_horizon_state_pm_mean_support_count']} von "
            f"{h1['state_poll_horizon_state_count']} States mit niedrigerem "
            f"mittleren Polymarket-Brier und "
            f"{h1['state_poll_horizon_state_pm_majority_support_count']} von "
            f"{h1['state_poll_horizon_state_count']} States mit Mehrheit "
            f"niedrigerer Polymarket-Tagesverluste. "
            f"Die <=90-Day-Score-Quality-Diagnose verdichtet dasselbe Fenster "
            f"zu {h1['state_poll_near_quality_forecast_row_count']} "
            f"Forecast-Zeilen und zeigt Polymarket mit niedrigerem Mean Brier "
            f"{h1['state_poll_near_quality_pm_mean_brier']:.4f} vs "
            f"{h1['state_poll_near_quality_poll_mean_brier']:.4f}, niedrigerem "
            f"Fixed-Bin-ECE {h1['state_poll_near_quality_pm_ece']:.4f} vs "
            f"{h1['state_poll_near_quality_poll_ece']:.4f} und hoeherer "
            f"Probability-Separation "
            f"{h1['state_poll_near_quality_pm_separation']:.4f} vs "
            f"{h1['state_poll_near_quality_poll_separation']:.4f}. "
            f"Die Poll-Transform-Sensitivitaet variiert MAE von "
            f"{h1['state_poll_sensitivity_min_mae']:.1f} bis "
            f"{h1['state_poll_sensitivity_max_mae']:.1f} Prozentpunkten; "
            f"Polymarket bleibt in allen "
            f"{h1['state_poll_sensitivity_row_count']} Parameterzeilen im "
            f"mittleren Brier niedriger und liegt je nach Annahme in "
            f"{h1['state_poll_sensitivity_min_pm_lower_loss_count']} bis "
            f"{h1['state_poll_sensitivity_max_pm_lower_loss_count']} von "
            f"{h1['state_poll_snapshot_case_count']} State-Outcomes vorne. "
            f"Der Coverage-Audit prueft {h1['state_poll_coverage_state_count']} "
            f"US-States und findet {h1['state_poll_coverage_polymarket_market_count']} "
            f"Polymarket-State-Maerkte, aber nur "
            f"{h1['state_poll_coverage_valid_pair_count']} valide H1-Paare "
            f"mit REP/DEM-Zeilen im 538-Snapshot. "
            f"Die Rieke-Erweiterung deckt alle "
            f"{h1['rieke_state_case_count']} State-Outcomes ab und zeigt "
            f"einen niedrigeren mittleren Polymarket-Brier "
            f"({h1['rieke_state_mean_pm_brier']:.4f} vs "
            f"{h1['rieke_state_mean_rieke_brier']:.4f}), aber nur "
            f"{h1['rieke_state_pm_lower_loss_count']} von "
            f"{h1['rieke_state_case_count']} Einzelfaellen mit niedrigerem "
            f"Polymarket-Verlust; Rieke liegt in "
            f"{h1['rieke_state_rieke_lower_loss_count']} von "
            f"{h1['rieke_state_case_count']} vorne. "
            f"Zusammen ergeben die 538-nahen Zusatzchecks "
            f"{h1['final_snapshot_case_count'] + h1['state_poll_snapshot_case_count']} "
            f"geloeste Outcomes, davon "
            f"{h1['final_snapshot_pm_lower_loss_count'] + h1['state_poll_snapshot_pm_lower_loss_count']} "
            f"mit niedrigerem Polymarket-Verlust; die Quellen bleiben "
            f"methodisch getrennt. "
            f"Die 270toWin/JHK-Erweiterung zeigt ebenfalls einen niedrigeren "
            f"mittleren Polymarket-Brier ({h1['two_seventy_state_mean_pm_brier']:.4f} "
            f"vs {h1['two_seventy_state_mean_270_brier']:.4f}), aber nur "
            f"{h1['two_seventy_state_pm_lower_loss_count']} von "
            f"{h1['two_seventy_state_case_count']} Einzelfaellen mit niedrigerem "
            f"Polymarket-Verlust; 270toWin/JHK liegt in "
            f"{h1['two_seventy_state_270_lower_loss_count']} von "
            f"{h1['two_seventy_state_case_count']} vorne. "
            "Das erklaert aber nicht den Mechanismus, beweist keine schnellere "
            "Informationsverarbeitung; die Rieke-Zahlen sind aggregierte "
            "Forecast-Qualitaetsstuetze und die 270toWin/JHK-Zahlen bestaetigen "
            "dieselbe Aggregatgrenze, aber keinen Mehrheit-der-States-Beweis."
        ),
    )


def _add_h2_section(doc: Document, h2: dict[str, Any]) -> None:
    doc.add_heading("7. H2 - Taegliche Event-Window-Reaktion", level=1)
    doc.add_paragraph(
        "H2 prueft, ob Polymarket-Wahrscheinlichkeiten um vorab kuratierte "
        "politische Ereignisse in plausibler Weise reagieren. Die Ereignisse "
        "wurden vor der Analyse fixiert, damit keine nachtraegliche Auswahl nach "
        "sichtbaren Preisbewegungen entsteht."
    )
    rows = [
        ("Warum", "Oeffentliche Informationen koennen sich in Prediction-Market-Preisen zeigen."),
        ("Methode", "Daily Event-Windows mit primaerem Fenster [0d,+1d] und Sensitivitaetsfenster [-1d,+3d]."),
        ("Datenform", f"{h2['event_count']} kuratierte Ereignisse, {h2['summary_rows']} kompakte Summary-Zeilen."),
        ("Ergebnis", "Mehrere Ereignisse zeigen sichtbare taegliche Bewegungen, z.B. Trump shooting +7.2 pp im Primaerfenster."),
        ("Grenze", "Daily-Daten erlauben keine Intraday-Speed-Aussage."),
    ]
    table = _add_table(doc, rows, ["Frage", "Antwort"], [1700, 7660])
    _shade_table_header(table)
    primary_rows = [
        (row["event"], f"{row['change_pp']:+.1f} pp")
        for row in h2["primary_examples"]
    ]
    table = _add_table(doc, primary_rows, ["Primaerfenster", "Finaler Change"], [7060, 2300])
    _shade_table_header(table)


def _add_h3_section(doc: Document, h3: dict[str, Any]) -> None:
    doc.add_heading("8. H3 - Wallet-Tier-Timing-Diagnostik", level=1)
    doc.add_paragraph(
        "H3 untersucht, ob aggregierte Wallet-Aktivitaet zeitliche Muster vor "
        "oder um Polymarket-Preisbewegungen zeigt. Die Wallet-Tiers sind "
        "dataset-relativ und aus der beobachteten Wallet-Verteilung abgeleitet, "
        "nicht aus einem willkuerlichen USD-Schwellenwert."
    )
    rows = [
        ("Warum", "Wallet-Aktivitaet kann ein fruehes Signal fuer Informationsverarbeitung sein."),
        ("Methode", "Wallet-Level cumulative amount_usd percentiles, Lead-Lag-Korrelationen und Granger-Diagnostik."),
        ("Datenform", f"{h3['model_rows']} alignierte Modellzeilen; Tier counts: {h3['tier_counts_text']}."),
        ("Ergebnis", f"Top-Tier zeigt staerkste dokumentierte Korrelation bei {h3['top_correlation_label']} = {h3['top_correlation']:.4f}; kleinster Granger-p-Wert {h3['min_granger_p']:.4f}."),
        ("Grenze", "BUY-only Quelle, taegliche Aggregation und Multiple-Testing-Sensitivitaet."),
    ]
    table = _add_table(doc, rows, ["Frage", "Antwort"], [1700, 7660])
    _shade_table_header(table)
    _add_callout(
        doc,
        "Wording fuer die Verteidigung",
        (
            "Formuliere H3 als Timing-Diagnostik, nicht als Beweis fuer private "
            "Information, Fehlverhalten, Kausalitaet oder Profitabilitaet."
        ),
    )


def _add_monitor_section(doc: Document, monitor: dict[str, Any]) -> None:
    doc.add_heading("9. Politics/Geo Monitor-Prototyp", level=1)
    doc.add_paragraph(
        "Der Monitor ist eine Forschungs-Erweiterung, nicht der deterministische "
        "Kern der Thesis. Er soll spaeter auffaellige Kombinationen aus Marktbewegung, "
        "Wallet-Tier-Aktivitaet, Konzentration und Event-Kontext sichtbar machen. "
        "Er bleibt read-only und darf keine Orders, Trading-Credentials oder "
        "Profitversprechen enthalten."
    )
    rows = [
        ("Warum", "Pruefen, ob H1-H3 eine spaetere Anomalie- und Signalhypothesen-Schicht motivieren koennen."),
        ("Methode", "Robuste rolling Baselines, Rule C mit combined-family confirmation, bounded summaries."),
        ("Recorded Replay", f"{monitor['snapshot_count']} Snapshot/Alert-Zeilen; Severity: {monitor['severity_counts_text']}."),
        ("Live Dashboard", f"{monitor['live_market_count']} Maerkte, {monitor['live_alert_count']} Alert-Zeilen, {monitor['live_scoring_rows']} Scoring-Zeilen."),
        ("Wallet Graph", f"{monitor['wallet_graph_nodes']} Nodes und {monitor['wallet_graph_edges']} Co-Activity-Edges als lokaler Review-Layer."),
        (
            "Anomaly Review Queue",
            f"{monitor['anomaly_queue_rows']} Cases: {monitor['anomaly_high_priority_count']} high, {monitor['anomaly_medium_priority_count']} medium, {monitor['anomaly_low_priority_count']} low; Status {monitor['anomaly_review_status_counts']}; Labels {monitor['anomaly_review_labels']}.",
        ),
        ("Grenze", "Keine PnL-Backtests, keine autonome Ausfuehrung, keine kausalen oder tradingbezogenen Schlussfolgerungen."),
    ]
    table = _add_table(doc, rows, ["Aspekt", "Stand"], [1900, 7460])
    _shade_table_header(table)
    _add_callout(
        doc,
        "Interpretationsgrenze",
        (
            f"{monitor['anomaly_allowed_interpretation']} "
            f"{monitor['anomaly_limitation']}"
        ),
    )


def _add_swiss_section(doc: Document, swiss: dict[str, Any]) -> None:
    doc.add_heading("10. Laufender Side-Track - Schweizer 10-Millionen-Referendum", level=1)
    doc.add_paragraph(
        "Die aktuelle Phase ist ein eigener, aktueller Vergleich: Polymarket-"
        "Wahrscheinlichkeit fuer die Annahme der Initiative gegen kuratierte "
        "Umfragewerte von SRG/gfs.bern, Tamedia/LeeWas und YouGov Schweiz. "
        "BFS/admin.ch dienen nur als Kontextquellen, nicht als Umfragequelle."
    )
    rows = [
        ("Warum", "Aktueller realer Referendumsmarkt erlaubt einen bounded Vergleich zwischen Polymarket und traditionellen Umfragen."),
        ("Methode", "Curated poll catalog, read-only Gamma Snapshot, bounded CLOB price-history um Poll-Releases."),
        ("Datenform", f"{swiss['poll_count']} Umfragen, {swiss['snapshot_count']} Snapshots, {swiss['history_rows']} Price-History-Zeilen."),
        ("Latest Result", f"Polymarket Yes {swiss['latest_poly_yes_pct']:.1f}%; latest poll Yes {swiss['latest_poll_yes_pct']:.1f}%; raw gap {swiss['latest_raw_gap_pp']:+.1f} pp."),
        ("Information Response", swiss["information_response_counts_text"]),
        ("Grenze", "Umfrageanteile sind keine Modell-Wahrscheinlichkeiten; keine Kausalitaet, Tradeability oder Mispricing-Behauptung."),
    ]
    table = _add_table(doc, rows, ["Aspekt", "Stand"], [1900, 7460])
    _shade_table_header(table)
    source_rows = [
        (item["source"], item["poll_id"], f"{item['poll_yes']:.1f}%", f"{item['raw_gap_pp']:+.1f} pp")
        for item in swiss["latest_source_rows"]
    ]
    table = _add_table(doc, source_rows, ["Quelle", "Neuester Poll", "Poll Yes", "Raw Gap"], [2800, 3300, 1500, 1760])
    _shade_table_header(table)


def _add_figures_section(doc: Document, figures: list[FigureSpec]) -> None:
    doc.add_heading("11. Visualisierungen", level=1)
    doc.add_paragraph(
        "Die folgenden Abbildungen sind bestehende oder lokal generierte Artefakte. "
        "Sie fuehren keine neuen statistischen Metriken ein, sondern visualisieren "
        "bereits vorhandene deterministische Outputs."
    )
    for idx, figure in enumerate(figures, start=1):
        if idx > 1:
            doc.add_section(WD_SECTION.CONTINUOUS)
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(3)
        run = heading.add_run(f"Abbildung {idx}: {figure.caption}")
        _set_run(run, bold=True, color=DARK_BLUE)
        if figure.path.exists():
            try:
                doc.add_picture(str(figure.path), width=Inches(6.25))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph(f"[Abbildung konnte nicht eingebettet werden: {figure.path}]")
        doc.add_paragraph(figure.note)


def _add_presentation_section(doc: Document) -> None:
    doc.add_heading("12. Vorschlag fuer die Dozentenpraesentation", level=1)
    rows = [
        ("1", "Problem", "Prediction Markets koennen Informationen aggregieren; die Arbeit prueft beobachtbare Effizienz-Proxies."),
        ("2", "Methodische Regel", "Alle Metriken werden deterministisch in Python berechnet; keine LLM-Rohdateninterpretation."),
        ("3", "H1-H3", "Forecast-Qualitaet, taegliche Event-Reaktionen, Wallet-Timing-Diagnostik."),
        ("4", "Ergebnisse", "Polymarket zeigt tiefere H1-Brier-Verluste; H2 findet Event-Bewegungen; H3 findet Timing-Diagnostik mit klaren Grenzen."),
        ("5", "Aktueller Zusatz", "Anomaly-Review-Queue bleibt human-review only; Swiss-Referendum-Track sammelt bis zum Abstimmungsergebnis weiter Daten."),
        ("6", "Naechster Schritt", "Quellen voll reviewen, Thesis-Kapitel aus den Artefakten schreiben und Sensitivitaets-/Limitationsabschnitte finalisieren."),
    ]
    table = _add_table(doc, rows, ["#", "Teil", "Kernaussage"], [700, 1900, 6760])
    _shade_table_header(table)


def _add_appendix(doc: Document, data: dict[str, Any]) -> None:
    doc.add_heading("13. Wichtige Artefakte fuer Rueckfragen", level=1)
    rows = [
        ("Steuerung", "GOAL.md, ROADMAP.md, STATUS.md, docs/project/WORK_LOG.md"),
        ("Literatur", "data/literature/literature_index.csv, docs/research/LITERATURE_MAP.md"),
        ("H1", "data/results/thesis_h1_summary.csv, h1_brier_scores.csv, h1_diebold_mariano.json, h1_forecast_quality_pairwise.csv, h1_forecast_quality_synthesis.csv, h1_claim_evidence_audit.csv, h1_claim_evidence_audit_summary.csv, h1_poll_comparison_result.csv, h1_poll_comparison_result_summary.csv, h1_poll_comparison_result.png, h1_poll_comparison_result_metadata.json, h1_poll_claim_readiness.csv, h1_poll_claim_readiness_summary.csv, h1_poll_claim_readiness.png, h1_poll_claim_readiness_metadata.json, h1_poll_scope_frontier.csv, h1_poll_scope_frontier_summary.csv, h1_poll_scope_frontier.png, h1_poll_scope_frontier_metadata.json, h1_poll_decision_matrix.csv, h1_poll_decision_matrix_summary.csv, h1_poll_decision_matrix.png, h1_poll_decision_matrix_metadata.json, h1_robust_poll_scope_quality_rows.csv, h1_robust_poll_scope_quality_bins.csv, h1_robust_poll_scope_quality_summary.csv, h1_robust_poll_scope_quality_pairwise.csv, h1_robust_poll_scope_quality.png, h1_robust_poll_scope_quality_metadata.json, h1_robust_poll_scope_unit_quality_units.csv, h1_robust_poll_scope_unit_quality_summary.csv, h1_robust_poll_scope_unit_quality.png, h1_robust_poll_scope_unit_quality_metadata.json, h1_poll_comparison_unit_robustness_units.csv, h1_poll_comparison_unit_robustness_summary.csv, h1_poll_comparison_unit_robustness.png, h1_poll_comparison_unit_robustness_metadata.json, h1_direct_poll_loss_decomposition_cases.csv, h1_direct_poll_loss_decomposition_summary.csv, h1_direct_poll_loss_decomposition.png, h1_direct_poll_loss_decomposition_metadata.json, h1_direct_poll_state_cluster_diagnostic_states.csv, h1_direct_poll_state_cluster_diagnostic_summary.csv, h1_direct_poll_state_cluster_diagnostic.png, h1_direct_poll_state_cluster_diagnostic_metadata.json, h1_direct_poll_outlier_robustness_scenarios.csv, h1_direct_poll_outlier_robustness_summary.csv, h1_direct_poll_outlier_robustness.png, h1_direct_poll_outlier_robustness_metadata.json, h1_calibration_diagnostic_summary.csv, h1_calibration_diagnostic_pairwise.csv, h1_evidence_scope.csv, h1_expansion_readiness.csv, h1_final_snapshot_cases.csv, h1_final_snapshot_summary.csv, h1_state_poll_snapshot_cases.csv, h1_state_poll_snapshot_summary.csv, h1_270towin_poll_average_cases.csv, h1_270towin_poll_average_summary.csv, h1_270towin_poll_average.png, h1_popular_vote_cases.csv, h1_popular_vote_summary.csv, h1_state_poll_panel_cases.csv, h1_state_poll_panel_summary.csv, h1_state_poll_panel_state_summary.csv, h1_state_poll_panel_temporal_summary.csv, h1_state_poll_panel_temporal_claim_audit.csv, h1_state_poll_panel_horizon_summary.csv, h1_state_poll_panel_horizon_claim_audit.csv, h1_state_poll_panel_horizon_state_support.csv, h1_state_poll_panel_horizon_state_support_summary.csv, h1_state_poll_panel_near_window_quality_summary.csv, h1_state_poll_panel_near_window_quality_bins.csv, h1_state_poll_panel_near_window_quality_rows.csv, h1_state_poll_panel_competitiveness_grid.csv, h1_state_poll_panel_competitiveness_summary.csv, h1_state_poll_panel_competitiveness.png, h1_state_poll_panel_state_significance.csv, h1_state_poll_panel_state_significance_summary.csv, h1_state_poll_panel_state_significance.png, h1_state_poll_snapshot_sensitivity.csv, h1_state_poll_snapshot_coverage.csv, h1_rieke_state_forecast_cases.csv, h1_rieke_state_forecast_summary.csv, h1_270towin_state_forecast_cases.csv, h1_270towin_state_forecast_summary.csv, h1_state_source_consensus_cases.csv, h1_state_source_consensus_summary.csv, h1_competitive_state_diagnostic_cases.csv, h1_competitive_state_diagnostic_summary.csv"),
        ("H2", "data/events_timeline_seed.csv, h2_event_window_summary.csv, thesis_h2_event_window_car.png"),
        ("H3", "h3_wallet_distribution_inventory.json, h3_granger_results.csv, thesis_h3_summary.csv"),
        ("Monitor", "monitor_v2_bounded_summary.csv, monitor_v2_polymarket_dashboard.html, wallet_graph_dashboard.html, monitor_anomaly_review_queue.csv, monitor_anomaly_review_summary.csv, monitor_anomaly_case_review_packets.csv, monitor_anomaly_review_decision_readiness.csv"),
        ("Swiss", "data/swiss_referendum_10mio_polls.csv, swiss_referendum_10mio_latest_summary.md, dashboard HTML"),
        ("Tests", "tests/ und operations/project/review_check.py"),
    ]
    table = _add_table(doc, rows, ["Bereich", "Artefakte"], [1700, 7660])
    _shade_table_header(table)


def _h1_data(
    summary: pd.DataFrame,
    pairwise: pd.DataFrame,
    synthesis: pd.DataFrame,
    claim_audit_summary: pd.DataFrame,
    poll_comparison_result: pd.DataFrame,
    poll_claim_readiness: pd.DataFrame,
    poll_scope_frontier: pd.DataFrame,
    poll_decision_matrix: pd.DataFrame,
    robust_poll_scope_quality_pairwise: pd.DataFrame,
    robust_poll_scope_quality_summary: pd.DataFrame,
    robust_poll_scope_unit_quality: pd.DataFrame,
    poll_comparison_unit_robustness: pd.DataFrame,
    direct_poll_loss_decomposition: pd.DataFrame,
    direct_poll_state_cluster: pd.DataFrame,
    direct_poll_outlier_robustness: pd.DataFrame,
    calibration_summary: pd.DataFrame,
    calibration_pairwise: pd.DataFrame,
    final_snapshot: pd.DataFrame,
    state_poll_snapshot: pd.DataFrame,
    popular_vote: pd.DataFrame,
    margin_threshold_readiness: pd.DataFrame,
    state_poll_panel: pd.DataFrame,
    state_poll_panel_temporal: pd.DataFrame,
    state_poll_panel_horizon: pd.DataFrame,
    state_poll_panel_horizon_state: pd.DataFrame,
    state_poll_panel_near_quality: pd.DataFrame,
    state_poll_sensitivity: pd.DataFrame,
    state_poll_coverage: pd.DataFrame,
    rieke_state_forecast: pd.DataFrame,
    two_seventy_state_forecast: pd.DataFrame,
    two_seventy_poll_average: pd.DataFrame,
    state_source_consensus: pd.DataFrame,
    competitive_state: pd.DataFrame,
    panel_competitiveness: pd.DataFrame,
    state_significance: pd.DataFrame,
) -> dict[str, Any]:
    fte = pairwise.loc[pairwise["comparator"] == "fivethirtyeight"].iloc[0]
    sensitivity = state_poll_sensitivity.sort_values("poll_error_mae_points")
    coverage_status = state_poll_coverage["coverage_status"].value_counts()
    calibration_by_source = calibration_summary.set_index("forecast_source_id")
    calibration_pm_state = calibration_by_source.loc["polymarket_state_final_50"]
    calibration_rieke_state = calibration_by_source.loc["rieke_state_final_50"]
    calibration_270_state = calibration_by_source.loc["two_seventy_state_final_50"]
    temporal_by_scope = state_poll_panel_temporal.set_index("audit_scope")
    temporal_support = temporal_by_scope.loc["polymarket_supporting_months"]
    temporal_full = temporal_by_scope.loc["full_panel"]
    horizon_by_scope = state_poll_panel_horizon.set_index("audit_scope")
    horizon_near = horizon_by_scope.loc["within_90_days_before_election"]
    near_quality = state_poll_panel_near_quality.set_index("source_id")
    near_quality_pm = near_quality.loc["polymarket"]
    near_quality_poll = near_quality.loc["poll_derived"]
    robust_quality_by_scope = robust_poll_scope_quality_pairwise.set_index("scope_id")
    robust_quality_largest = robust_quality_by_scope.loc[
        "largest_robust_lte120_low_middle"
    ]
    robust_quality_strongest = robust_quality_by_scope.loc[
        "strongest_robust_lte90_low_middle"
    ]
    robust_quality_summary = robust_poll_scope_quality_summary.set_index(
        ["scope_id", "source_id"]
    )
    robust_quality_strongest_pm = robust_quality_summary.loc[
        ("strongest_robust_lte90_low_middle", "polymarket")
    ]
    robust_unit_quality = robust_poll_scope_unit_quality.set_index(
        ["scope_id", "unit_type"]
    )
    robust_unit_largest_state = robust_unit_quality.loc[
        ("largest_robust_lte120_low_middle", "state")
    ]
    robust_unit_largest_state_month = robust_unit_quality.loc[
        ("largest_robust_lte120_low_middle", "state_month")
    ]
    robust_unit_largest_state_horizon = robust_unit_quality.loc[
        ("largest_robust_lte120_low_middle", "state_horizon")
    ]
    robust_unit_strongest_state = robust_unit_quality.loc[
        ("strongest_robust_lte90_low_middle", "state")
    ]
    robust_unit_strongest_state_month = robust_unit_quality.loc[
        ("strongest_robust_lte90_low_middle", "state_month")
    ]
    robust_unit_strongest_state_horizon = robust_unit_quality.loc[
        ("strongest_robust_lte90_low_middle", "state_horizon")
    ]
    margin_status = margin_threshold_readiness["status"].value_counts()
    return {
        "observation_count": int(_summary_value(summary, "h1_observation_count")),
        "brier_polymarket": _summary_value(summary, "h1_mean_brier_polymarket"),
        "brier_fivethirtyeight": _summary_value(summary, "h1_mean_brier_fivethirtyeight"),
        "brier_always_50": _summary_value(summary, "h1_mean_brier_always_50"),
        "brier_prior_day": _summary_value(summary, "h1_mean_brier_prior_day_polymarket"),
        "dm_polymarket_vs_538": _summary_value(
            summary,
            "h1_dm_polymarket_vs_fivethirtyeight",
        ),
        "pm_better_vs_538_count": int(fte["polymarket_lower_loss_count"]),
        "pm_vs_538_count": int(fte["comparison_row_count"]),
        "pm_better_vs_538_share": float(fte["polymarket_better_share"]),
        "mean_loss_advantage_vs_538": float(fte["mean_loss_advantage"]),
        "synthesis_evidence_row_count": int(len(synthesis)),
        "synthesis_aggregate_support_count": _bool_count(
            synthesis,
            "aggregate_mean_supports_polymarket",
        ),
        "synthesis_majority_support_count": _bool_count(
            synthesis,
            "majority_cases_supports_polymarket",
        ),
        "synthesis_broad_support_count": _bool_count(
            synthesis,
            "broad_many_cases_claim_supported",
        ),
        "claim_audit_row_count": int(
            _summary_value(claim_audit_summary, "audit_row_count")
        ),
        "claim_audit_support_row_count": int(
            _summary_value(claim_audit_summary, "support_row_count")
        ),
        "claim_audit_contradiction_row_count": int(
            _summary_value(claim_audit_summary, "contradiction_row_count")
        ),
        "claim_audit_direct_poll_row_count": int(
            _summary_value(claim_audit_summary, "direct_poll_audit_row_count")
        ),
        "claim_audit_direct_poll_support_row_count": int(
            _summary_value(claim_audit_summary, "direct_poll_support_row_count")
        ),
        "claim_audit_direct_poll_contradiction_row_count": int(
            _summary_value(claim_audit_summary, "direct_poll_contradiction_row_count")
        ),
        "claim_audit_broad_user_claim_proven": int(
            _summary_value(claim_audit_summary, "broad_user_claim_proven")
        ),
        "poll_result_primary_pm_count": int(
            _summary_value(poll_comparison_result, "primary_polymarket_support_count")
        ),
        "poll_result_primary_poll_count": int(
            _summary_value(poll_comparison_result, "primary_poll_support_count")
        ),
        "poll_result_primary_row_count": int(
            _summary_value(poll_comparison_result, "primary_comparison_count")
        ),
        "poll_result_primary_pm_share": _summary_value(
            poll_comparison_result,
            "primary_polymarket_support_share",
        ),
        "poll_result_primary_state_count": int(
            _summary_value(poll_comparison_result, "primary_state_count")
        ),
        "poll_result_primary_pm_state_count": int(
            _summary_value(poll_comparison_result, "primary_polymarket_state_count")
        ),
        "poll_result_primary_p_value": _summary_value(
            poll_comparison_result,
            "primary_exact_binomial_p_value",
        ),
        "poll_result_primary_ci_low": _summary_value(
            poll_comparison_result,
            "primary_exact_95_ci_low",
        ),
        "poll_result_direct_poll_support_count": int(
            _summary_value(poll_comparison_result, "direct_poll_audit_support_count")
        ),
        "poll_result_direct_poll_row_count": int(
            _summary_value(poll_comparison_result, "direct_poll_audit_row_count")
        ),
        "poll_result_full_panel_pm_count": int(
            _summary_value(
                poll_comparison_result,
                "full_panel_polymarket_support_count",
            )
        ),
        "poll_result_full_panel_poll_count": int(
            _summary_value(poll_comparison_result, "full_panel_poll_support_count")
        ),
        "poll_result_full_panel_row_count": int(
            _summary_value(poll_comparison_result, "full_panel_polymarket_support_count")
            + _summary_value(poll_comparison_result, "full_panel_poll_support_count")
        ),
        "poll_result_late_high_poll_count": int(
            _summary_value(
                poll_comparison_result,
                "late_high_distance_poll_support_count",
            )
        ),
        "poll_result_bounded_supported": int(
            _summary_value(
                poll_comparison_result,
                "bounded_polymarket_statement_supported",
            )
        ),
        "poll_result_broad_claim_proven": int(
            _summary_value(poll_comparison_result, "broad_claim_proven")
        ),
        "poll_result_goal_status": _summary_text_value(
            poll_comparison_result,
            "h1_goal_completion_status",
        ),
        "poll_claim_row_count": int(
            _summary_value(poll_claim_readiness, "claim_row_count")
        ),
        "poll_claim_supported_bounded_count": int(
            _summary_value(
                poll_claim_readiness,
                "supported_bounded_scope_row_count",
            )
        ),
        "poll_claim_mixed_mean_count": int(
            _summary_value(poll_claim_readiness, "mixed_mean_support_row_count")
        ),
        "poll_claim_counterexample_count": int(
            _summary_value(poll_claim_readiness, "counterexample_row_count")
        ),
        "poll_claim_primary_pm_count": int(
            _summary_value(
                poll_claim_readiness,
                "primary_polymarket_support_count",
            )
        ),
        "poll_claim_primary_count": int(
            _summary_value(poll_claim_readiness, "primary_comparison_count")
        ),
        "poll_claim_primary_pm_share": _summary_value(
            poll_claim_readiness,
            "primary_polymarket_support_share",
        ),
        "poll_claim_state_month_pm_count": int(
            _summary_value(
                poll_claim_readiness,
                "primary_state_month_polymarket_support_count",
            )
        ),
        "poll_claim_state_month_count": int(
            _summary_value(poll_claim_readiness, "primary_state_month_unit_count")
        ),
        "poll_claim_state_month_p_value": _summary_value(
            poll_claim_readiness,
            "primary_state_month_exact_p_value",
        ),
        "poll_claim_state_month_ci_low": _summary_value(
            poll_claim_readiness,
            "primary_state_month_exact_95_ci_low",
        ),
        "poll_claim_bounded_supported": int(
            _summary_value(poll_claim_readiness, "bounded_poll_claim_supported")
        ),
        "poll_claim_broad_proven": int(
            _summary_value(poll_claim_readiness, "broad_claim_proven")
        ),
        "poll_claim_goal_status": _summary_text_value(
            poll_claim_readiness,
            "h1_goal_completion_status",
        ),
        "poll_frontier_row_count": int(
            _summary_value(poll_scope_frontier, "frontier_row_count")
        ),
        "poll_frontier_robust_scope_count": int(
            _summary_value(poll_scope_frontier, "robust_scope_count")
        ),
        "poll_frontier_largest_scope_id": _summary_text_value(
            poll_scope_frontier,
            "largest_robust_scope_id",
        ),
        "poll_frontier_largest_horizon": _summary_text_value(
            poll_scope_frontier,
            "largest_robust_horizon_label",
        ),
        "poll_frontier_largest_tier": _summary_text_value(
            poll_scope_frontier,
            "largest_robust_tier_label",
        ),
        "poll_frontier_largest_row_count": int(
            _summary_value(poll_scope_frontier, "largest_robust_row_count")
        ),
        "poll_frontier_largest_pm_count": int(
            _summary_value(
                poll_scope_frontier,
                "largest_robust_polymarket_support_count",
            )
        ),
        "poll_frontier_largest_poll_count": int(
            _summary_value(poll_scope_frontier, "largest_robust_poll_support_count")
        ),
        "poll_frontier_largest_pm_share": _summary_value(
            poll_scope_frontier,
            "largest_robust_polymarket_support_share",
        ),
        "poll_frontier_largest_state_count": int(
            _summary_value(poll_scope_frontier, "largest_robust_state_count")
        ),
        "poll_frontier_largest_state_month_pm_count": int(
            _summary_value(
                poll_scope_frontier,
                "largest_robust_state_month_polymarket_support_count",
            )
        ),
        "poll_frontier_largest_state_month_count": int(
            _summary_value(poll_scope_frontier, "largest_robust_state_month_count")
        ),
        "poll_frontier_largest_state_month_p_value": _summary_value(
            poll_scope_frontier,
            "largest_robust_state_month_p_value",
        ),
        "poll_frontier_largest_mean_loss_advantage": _summary_value(
            poll_scope_frontier,
            "largest_robust_mean_loss_advantage",
        ),
        "poll_frontier_strongest_scope_id": _summary_text_value(
            poll_scope_frontier,
            "strongest_robust_scope_id",
        ),
        "poll_frontier_strongest_row_count": int(
            _summary_value(poll_scope_frontier, "strongest_robust_row_count")
        ),
        "poll_frontier_strongest_p_value": _summary_value(
            poll_scope_frontier,
            "strongest_robust_state_month_p_value",
        ),
        "poll_frontier_lte_90_all_row_count": int(
            _summary_value(poll_scope_frontier, "lte_90_all_row_count")
        ),
        "poll_frontier_lte_90_all_pm_count": int(
            _summary_value(
                poll_scope_frontier,
                "lte_90_all_polymarket_support_count",
            )
        ),
        "poll_frontier_lte_90_all_pm_share": _summary_value(
            poll_scope_frontier,
            "lte_90_all_polymarket_support_share",
        ),
        "poll_frontier_lte_90_all_state_month_p_value": _summary_value(
            poll_scope_frontier,
            "lte_90_all_state_month_p_value",
        ),
        "poll_frontier_full_panel_pm_count": int(
            _summary_value(poll_scope_frontier, "full_panel_polymarket_support_count")
        ),
        "poll_frontier_full_panel_poll_count": int(
            _summary_value(poll_scope_frontier, "full_panel_poll_support_count")
        ),
        "poll_frontier_full_panel_row_count": int(
            _summary_value(poll_scope_frontier, "full_panel_row_count")
        ),
        "poll_frontier_broad_claim_proven": int(
            _summary_value(poll_scope_frontier, "broad_claim_proven")
        ),
        "poll_frontier_goal_status": _summary_text_value(
            poll_scope_frontier,
            "h1_goal_completion_status",
        ),
        "poll_decision_row_count": int(
            _summary_value(poll_decision_matrix, "decision_row_count")
        ),
        "poll_decision_robust_yes_count": int(
            _summary_value(poll_decision_matrix, "robust_bounded_yes_count")
        ),
        "poll_decision_mixed_mean_count": int(
            _summary_value(poll_decision_matrix, "mixed_mean_only_count")
        ),
        "poll_decision_counterexample_count": int(
            _summary_value(poll_decision_matrix, "counterexample_count")
        ),
        "poll_decision_largest_pm_count": int(
            _summary_value(
                poll_decision_matrix,
                "largest_robust_polymarket_support_count",
            )
        ),
        "poll_decision_largest_poll_count": int(
            _summary_value(
                poll_decision_matrix,
                "largest_robust_comparator_support_count",
            )
        ),
        "poll_decision_largest_row_count": int(
            _summary_value(poll_decision_matrix, "largest_robust_row_count")
        ),
        "poll_decision_largest_pm_share": _summary_value(
            poll_decision_matrix,
            "largest_robust_polymarket_support_share",
        ),
        "poll_decision_largest_state_month_pm_count": int(
            _summary_value(
                poll_decision_matrix,
                "largest_robust_state_month_polymarket_support_count",
            )
        ),
        "poll_decision_largest_state_month_count": int(
            _summary_value(poll_decision_matrix, "largest_robust_state_month_count")
        ),
        "poll_decision_largest_p_value": _summary_value(
            poll_decision_matrix,
            "largest_robust_p_value",
        ),
        "poll_decision_strongest_scope_id": _summary_text_value(
            poll_decision_matrix,
            "strongest_robust_scope_id",
        ),
        "poll_decision_strongest_pm_count": int(
            _summary_value(
                poll_decision_matrix,
                "strongest_robust_polymarket_support_count",
            )
        ),
        "poll_decision_strongest_row_count": int(
            _summary_value(poll_decision_matrix, "strongest_robust_row_count")
        ),
        "poll_decision_strongest_p_value": _summary_value(
            poll_decision_matrix,
            "strongest_robust_p_value",
        ),
        "poll_decision_full_panel_poll_count": int(
            _summary_value(poll_decision_matrix, "full_panel_poll_support_count")
        ),
        "poll_decision_full_panel_row_count": int(
            _summary_value(poll_decision_matrix, "full_panel_row_count")
        ),
        "poll_decision_calibration_pairwise_count": int(
            _summary_value(poll_decision_matrix, "calibration_pairwise_count")
        ),
        "poll_decision_calibration_aggregate_count": int(
            _summary_value(
                poll_decision_matrix,
                "calibration_aggregate_support_count",
            )
        ),
        "poll_decision_calibration_majority_count": int(
            _summary_value(
                poll_decision_matrix,
                "calibration_majority_support_count",
            )
        ),
        "poll_decision_bounded_ready": int(
            _summary_value(poll_decision_matrix, "bounded_poll_claim_ready")
        ),
        "poll_decision_broad_proven": int(
            _summary_value(poll_decision_matrix, "broad_claim_proven")
        ),
        "poll_decision_goal_status": _summary_text_value(
            poll_decision_matrix,
            "h1_goal_completion_status",
        ),
        "robust_quality_scope_count": int(len(robust_poll_scope_quality_pairwise)),
        "robust_quality_forecast_row_count": int(
            (
                robust_quality_largest["case_count"]
                + robust_quality_strongest["case_count"]
            )
            * 2
        ),
        "robust_quality_case_count": int(
            robust_quality_largest["case_count"]
            + robust_quality_strongest["case_count"]
        ),
        "robust_quality_largest_case_count": int(
            robust_quality_largest["case_count"]
        ),
        "robust_quality_largest_pm_count": int(
            robust_quality_largest["polymarket_lower_loss_count"]
        ),
        "robust_quality_largest_poll_count": int(
            robust_quality_largest["poll_derived_lower_loss_count"]
        ),
        "robust_quality_largest_pm_share": float(
            robust_quality_largest["polymarket_lower_loss_share"]
        ),
        "robust_quality_largest_pm_brier": float(
            robust_quality_largest["mean_polymarket_brier"]
        ),
        "robust_quality_largest_poll_brier": float(
            robust_quality_largest["mean_poll_derived_brier"]
        ),
        "robust_quality_largest_mean_advantage": float(
            robust_quality_largest["mean_loss_advantage"]
        ),
        "robust_quality_largest_pm_ece": float(
            robust_quality_largest["polymarket_expected_calibration_error"]
        ),
        "robust_quality_largest_poll_ece": float(
            robust_quality_largest["poll_derived_expected_calibration_error"]
        ),
        "robust_quality_largest_pm_separation": float(
            robust_quality_largest["polymarket_probability_separation"]
        ),
        "robust_quality_largest_poll_separation": float(
            robust_quality_largest["poll_derived_probability_separation"]
        ),
        "robust_quality_strongest_case_count": int(
            robust_quality_strongest["case_count"]
        ),
        "robust_quality_strongest_pm_count": int(
            robust_quality_strongest["polymarket_lower_loss_count"]
        ),
        "robust_quality_strongest_poll_count": int(
            robust_quality_strongest["poll_derived_lower_loss_count"]
        ),
        "robust_quality_strongest_pm_share": float(
            robust_quality_strongest["polymarket_lower_loss_share"]
        ),
        "robust_quality_strongest_pm_brier": float(
            robust_quality_strongest["mean_polymarket_brier"]
        ),
        "robust_quality_strongest_poll_brier": float(
            robust_quality_strongest["mean_poll_derived_brier"]
        ),
        "robust_quality_strongest_pm_ece": float(
            robust_quality_strongest["polymarket_expected_calibration_error"]
        ),
        "robust_quality_strongest_poll_ece": float(
            robust_quality_strongest["poll_derived_expected_calibration_error"]
        ),
        "robust_quality_strongest_positive_rate": float(
            robust_quality_strongest_pm["positive_rate"]
        ),
        "robust_quality_broad_claim_proven": int(
            _bool_count(robust_poll_scope_quality_pairwise, "broad_claim_supported")
        ),
        "robust_unit_summary_row_count": int(len(robust_poll_scope_unit_quality)),
        "robust_unit_largest_state_count": int(
            robust_unit_largest_state["unit_count"]
        ),
        "robust_unit_largest_state_pm_count": int(
            robust_unit_largest_state["polymarket_support_count"]
        ),
        "robust_unit_largest_state_p_value": float(
            robust_unit_largest_state["exact_binomial_p_value_greater"]
        ),
        "robust_unit_largest_state_month_count": int(
            robust_unit_largest_state_month["unit_count"]
        ),
        "robust_unit_largest_state_month_pm_count": int(
            robust_unit_largest_state_month["polymarket_support_count"]
        ),
        "robust_unit_largest_state_month_poll_count": int(
            robust_unit_largest_state_month["poll_derived_support_count"]
        ),
        "robust_unit_largest_state_month_p_value": float(
            robust_unit_largest_state_month["exact_binomial_p_value_greater"]
        ),
        "robust_unit_largest_state_month_ci_low": float(
            robust_unit_largest_state_month["exact_95_ci_low"]
        ),
        "robust_unit_largest_state_month_median_advantage": float(
            robust_unit_largest_state_month["median_unit_loss_advantage"]
        ),
        "robust_unit_largest_state_horizon_count": int(
            robust_unit_largest_state_horizon["unit_count"]
        ),
        "robust_unit_largest_state_horizon_pm_count": int(
            robust_unit_largest_state_horizon["polymarket_support_count"]
        ),
        "robust_unit_largest_state_horizon_p_value": float(
            robust_unit_largest_state_horizon["exact_binomial_p_value_greater"]
        ),
        "robust_unit_strongest_state_count": int(
            robust_unit_strongest_state["unit_count"]
        ),
        "robust_unit_strongest_state_pm_count": int(
            robust_unit_strongest_state["polymarket_support_count"]
        ),
        "robust_unit_strongest_state_p_value": float(
            robust_unit_strongest_state["exact_binomial_p_value_greater"]
        ),
        "robust_unit_strongest_state_month_count": int(
            robust_unit_strongest_state_month["unit_count"]
        ),
        "robust_unit_strongest_state_month_pm_count": int(
            robust_unit_strongest_state_month["polymarket_support_count"]
        ),
        "robust_unit_strongest_state_month_p_value": float(
            robust_unit_strongest_state_month["exact_binomial_p_value_greater"]
        ),
        "robust_unit_strongest_state_month_ci_low": float(
            robust_unit_strongest_state_month["exact_95_ci_low"]
        ),
        "robust_unit_strongest_state_month_median_advantage": float(
            robust_unit_strongest_state_month["median_unit_loss_advantage"]
        ),
        "robust_unit_strongest_state_horizon_count": int(
            robust_unit_strongest_state_horizon["unit_count"]
        ),
        "robust_unit_strongest_state_horizon_pm_count": int(
            robust_unit_strongest_state_horizon["polymarket_support_count"]
        ),
        "robust_unit_strongest_state_horizon_p_value": float(
            robust_unit_strongest_state_horizon["exact_binomial_p_value_greater"]
        ),
        "robust_unit_broad_claim_proven": int(
            _bool_count(robust_poll_scope_unit_quality, "broad_claim_supported")
        ),
        "poll_unit_state_count": int(
            _summary_value(poll_comparison_unit_robustness, "primary_state_unit_count")
        ),
        "poll_unit_state_pm_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "primary_state_polymarket_support_count",
            )
        ),
        "poll_unit_state_month_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "primary_state_month_unit_count",
            )
        ),
        "poll_unit_state_month_pm_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "primary_state_month_polymarket_support_count",
            )
        ),
        "poll_unit_state_month_p_value": _summary_value(
            poll_comparison_unit_robustness,
            "primary_state_month_polymarket_exact_binomial_p_value_greater",
        ),
        "poll_unit_state_month_ci_low": _summary_value(
            poll_comparison_unit_robustness,
            "primary_state_month_polymarket_exact_95_ci_low",
        ),
        "poll_unit_state_horizon_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "primary_state_horizon_unit_count",
            )
        ),
        "poll_unit_state_horizon_pm_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "primary_state_horizon_polymarket_support_count",
            )
        ),
        "poll_unit_horizon_tier_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "primary_horizon_tier_unit_count",
            )
        ),
        "poll_unit_horizon_tier_pm_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "primary_horizon_tier_polymarket_support_count",
            )
        ),
        "poll_unit_full_panel_state_month_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "full_panel_state_month_unit_count",
            )
        ),
        "poll_unit_full_panel_state_month_poll_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "full_panel_state_month_poll_support_count",
            )
        ),
        "poll_unit_late_high_state_month_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "late_high_state_month_unit_count",
            )
        ),
        "poll_unit_late_high_state_month_poll_count": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "late_high_state_month_poll_support_count",
            )
        ),
        "poll_unit_late_high_state_month_poll_p_value": _summary_value(
            poll_comparison_unit_robustness,
            "late_high_state_month_poll_exact_binomial_p_value_greater",
        ),
        "poll_unit_scope_supported_across_all_units": int(
            _summary_value(
                poll_comparison_unit_robustness,
                "primary_scope_supported_across_all_units",
            )
        ),
        "poll_unit_broad_claim_proven": int(
            _summary_value(poll_comparison_unit_robustness, "broad_claim_proven")
        ),
        "poll_unit_goal_status": _summary_text_value(
            poll_comparison_unit_robustness,
            "h1_goal_completion_status",
        ),
        "direct_poll_loss_case_count": int(
            _summary_value(direct_poll_loss_decomposition, "direct_poll_case_count")
        ),
        "direct_poll_loss_pm_count": int(
            _summary_value(
                direct_poll_loss_decomposition,
                "direct_poll_polymarket_lower_loss_count",
            )
        ),
        "direct_poll_loss_poll_count": int(
            _summary_value(
                direct_poll_loss_decomposition,
                "direct_poll_comparator_lower_loss_count",
            )
        ),
        "direct_poll_loss_pm_brier": _summary_value(
            direct_poll_loss_decomposition,
            "direct_poll_mean_polymarket_brier",
        ),
        "direct_poll_loss_poll_brier": _summary_value(
            direct_poll_loss_decomposition,
            "direct_poll_mean_poll_derived_brier",
        ),
        "direct_poll_loss_mean_advantage": _summary_value(
            direct_poll_loss_decomposition,
            "direct_poll_mean_loss_advantage",
        ),
        "direct_poll_loss_pm_win_total_advantage": _summary_value(
            direct_poll_loss_decomposition,
            "polymarket_win_total_loss_advantage",
        ),
        "direct_poll_loss_poll_win_total_advantage": _summary_value(
            direct_poll_loss_decomposition,
            "comparator_win_total_loss_advantage_abs",
        ),
        "direct_poll_loss_pm_win_mean_advantage": _summary_value(
            direct_poll_loss_decomposition,
            "polymarket_win_mean_loss_advantage",
        ),
        "direct_poll_loss_poll_win_mean_advantage": _summary_value(
            direct_poll_loss_decomposition,
            "comparator_win_mean_loss_advantage_abs",
        ),
        "direct_poll_loss_margin_ratio": _summary_value(
            direct_poll_loss_decomposition,
            "polymarket_win_total_to_comparator_win_abs_ratio",
        ),
        "direct_poll_loss_aggregate_supports_pm": int(
            _summary_value(
                direct_poll_loss_decomposition,
                "direct_poll_aggregate_mean_supports_polymarket",
            )
        ),
        "direct_poll_loss_case_majority_supports_pm": int(
            _summary_value(
                direct_poll_loss_decomposition,
                "direct_poll_case_majority_supports_polymarket",
            )
        ),
        "direct_poll_state_cluster_case_count": int(
            _summary_value(direct_poll_state_cluster, "source_state_case_count")
        ),
        "direct_poll_state_cluster_state_count": int(
            _summary_value(direct_poll_state_cluster, "state_count")
        ),
        "direct_poll_state_cluster_pm_state_count": int(
            _summary_value(
                direct_poll_state_cluster,
                "state_mean_polymarket_support_count",
            )
        ),
        "direct_poll_state_cluster_poll_state_count": int(
            _summary_value(direct_poll_state_cluster, "state_mean_poll_support_count")
        ),
        "direct_poll_state_cluster_mean_advantage": _summary_value(
            direct_poll_state_cluster,
            "equal_state_mean_loss_advantage",
        ),
        "direct_poll_state_cluster_median_advantage": _summary_value(
            direct_poll_state_cluster,
            "equal_state_median_loss_advantage",
        ),
        "direct_poll_state_cluster_bootstrap_ci_low": _summary_value(
            direct_poll_state_cluster,
            "equal_state_bootstrap_95_ci_low",
        ),
        "direct_poll_state_cluster_bootstrap_ci_high": _summary_value(
            direct_poll_state_cluster,
            "equal_state_bootstrap_95_ci_high",
        ),
        "direct_poll_state_cluster_sign_flip_p": _summary_value(
            direct_poll_state_cluster,
            "equal_state_sign_flip_p_value_greater",
        ),
        "direct_poll_state_cluster_pm_state_p": _summary_value(
            direct_poll_state_cluster,
            "state_mean_polymarket_exact_binomial_p_value_greater",
        ),
        "direct_poll_state_cluster_poll_state_p": _summary_value(
            direct_poll_state_cluster,
            "state_mean_poll_exact_binomial_p_value_greater",
        ),
        "direct_poll_state_cluster_mean_supports_pm": int(
            _summary_value(
                direct_poll_state_cluster,
                "state_cluster_mean_supports_polymarket",
            )
        ),
        "direct_poll_state_cluster_majority_supports_pm": int(
            _summary_value(
                direct_poll_state_cluster,
                "state_count_majority_supports_polymarket",
            )
        ),
        "direct_poll_outlier_state_count": int(
            _summary_value(direct_poll_outlier_robustness, "state_count")
        ),
        "direct_poll_outlier_full_mean_advantage": _summary_value(
            direct_poll_outlier_robustness,
            "full_mean_loss_advantage",
        ),
        "direct_poll_outlier_min_leave_one_mean": _summary_value(
            direct_poll_outlier_robustness,
            "min_leave_one_out_mean_loss_advantage",
        ),
        "direct_poll_outlier_leave_one_all_positive": int(
            _summary_value(direct_poll_outlier_robustness, "leave_one_out_all_positive")
        ),
        "direct_poll_outlier_most_influential_state": _summary_text_value(
            direct_poll_outlier_robustness,
            "most_influential_removed_state",
        ),
        "direct_poll_outlier_top_k_positive": int(
            _summary_value(
                direct_poll_outlier_robustness,
                "max_top_positive_exclusion_k_with_positive_mean",
            )
        ),
        "direct_poll_outlier_first_nonpositive_k": int(
            _summary_value(
                direct_poll_outlier_robustness,
                "first_nonpositive_top_positive_exclusion_k",
            )
        ),
        "direct_poll_outlier_first_nonpositive_mean": _summary_value(
            direct_poll_outlier_robustness,
            "first_nonpositive_top_positive_exclusion_mean",
        ),
        "direct_poll_outlier_largest_positive_state": _summary_text_value(
            direct_poll_outlier_robustness,
            "largest_positive_state",
        ),
        "direct_poll_outlier_largest_positive_advantage": _summary_value(
            direct_poll_outlier_robustness,
            "largest_positive_state_loss_advantage",
        ),
        "direct_poll_outlier_supports_pm_mean": int(
            _summary_value(
                direct_poll_outlier_robustness,
                "outlier_robustness_supports_polymarket_mean",
            )
        ),
        "direct_poll_outlier_goal_status": _summary_text_value(
            direct_poll_outlier_robustness,
            "h1_goal_completion_status",
        ),
        "state_source_consensus_case_count": int(
            _summary_value(state_source_consensus, "source_state_case_count")
        ),
        "state_source_consensus_state_count": int(
            _summary_value(state_source_consensus, "state_count")
        ),
        "state_source_consensus_pm_case_count": int(
            _summary_value(
                state_source_consensus,
                "all_source_polymarket_lower_loss_count",
            )
        ),
        "state_source_consensus_comparator_case_count": int(
            _summary_value(
                state_source_consensus,
                "all_source_comparator_lower_loss_count",
            )
        ),
        "state_source_consensus_pm_state_count": int(
            _summary_value(
                state_source_consensus,
                "all_source_polymarket_majority_state_count",
            )
        ),
        "state_source_consensus_comparator_state_count": int(
            _summary_value(
                state_source_consensus,
                "all_source_comparator_majority_state_count",
            )
        ),
        "state_source_consensus_tie_state_count": int(
            _summary_value(state_source_consensus, "all_source_tie_state_count")
        ),
        "state_source_consensus_direct_two_state_count": int(
            _summary_value(state_source_consensus, "direct_poll_two_source_state_count")
        ),
        "state_source_consensus_direct_two_pm_state_count": int(
            _summary_value(
                state_source_consensus,
                "direct_poll_two_source_polymarket_majority_state_count",
            )
        ),
        "state_source_consensus_direct_two_comparator_state_count": int(
            _summary_value(
                state_source_consensus,
                "direct_poll_two_source_comparator_majority_state_count",
            )
        ),
        "state_source_consensus_direct_two_tie_state_count": int(
            _summary_value(
                state_source_consensus,
                "direct_poll_two_source_tie_state_count",
            )
        ),
        "competitive_state_case_count": int(
            _summary_value(competitive_state, "case_count")
        ),
        "competitive_state_all_low_case_count": int(
            _summary_value(competitive_state, "all_low_distance_case_count")
        ),
        "competitive_state_all_low_pm_count": int(
            _summary_value(
                competitive_state,
                "all_low_distance_polymarket_lower_loss_count",
            )
        ),
        "competitive_state_all_low_comparator_count": int(
            _summary_value(
                competitive_state,
                "all_low_distance_comparator_lower_loss_count",
            )
        ),
        "competitive_state_all_low_advantage": _summary_value(
            competitive_state,
            "all_low_distance_mean_loss_advantage",
        ),
        "competitive_state_all_high_case_count": int(
            _summary_value(competitive_state, "all_high_distance_case_count")
        ),
        "competitive_state_all_high_pm_count": int(
            _summary_value(
                competitive_state,
                "all_high_distance_polymarket_lower_loss_count",
            )
        ),
        "competitive_state_all_high_comparator_count": int(
            _summary_value(
                competitive_state,
                "all_high_distance_comparator_lower_loss_count",
            )
        ),
        "competitive_state_direct_low_case_count": int(
            _summary_value(competitive_state, "direct_low_distance_case_count")
        ),
        "competitive_state_direct_low_pm_count": int(
            _summary_value(
                competitive_state,
                "direct_low_distance_polymarket_lower_loss_count",
            )
        ),
        "competitive_state_direct_low_comparator_count": int(
            _summary_value(
                competitive_state,
                "direct_low_distance_comparator_lower_loss_count",
            )
        ),
        "competitive_state_direct_low_advantage": _summary_value(
            competitive_state,
            "direct_low_distance_mean_loss_advantage",
        ),
        "competitive_state_direct_high_case_count": int(
            _summary_value(competitive_state, "direct_high_distance_case_count")
        ),
        "competitive_state_direct_high_pm_count": int(
            _summary_value(
                competitive_state,
                "direct_high_distance_polymarket_lower_loss_count",
            )
        ),
        "competitive_state_direct_high_comparator_count": int(
            _summary_value(
                competitive_state,
                "direct_high_distance_comparator_lower_loss_count",
            )
        ),
        "panel_comp_row_count": int(
            _summary_value(panel_competitiveness, "panel_row_count")
        ),
        "panel_comp_late_non_safe_row_count": int(
            _summary_value(panel_competitiveness, "late_non_safe_row_count")
        ),
        "panel_comp_late_non_safe_state_count": int(
            _summary_value(panel_competitiveness, "late_non_safe_state_count")
        ),
        "panel_comp_late_non_safe_pm_count": int(
            _summary_value(
                panel_competitiveness,
                "late_non_safe_polymarket_lower_loss_count",
            )
        ),
        "panel_comp_late_non_safe_poll_count": int(
            _summary_value(panel_competitiveness, "late_non_safe_poll_lower_loss_count")
        ),
        "panel_comp_late_non_safe_state_support_count": int(
            _summary_value(
                panel_competitiveness,
                "late_non_safe_polymarket_state_support_count",
            )
        ),
        "panel_comp_late_non_safe_advantage": _summary_value(
            panel_competitiveness,
            "late_non_safe_mean_loss_advantage",
        ),
        "panel_comp_late_high_row_count": int(
            _summary_value(panel_competitiveness, "late_high_distance_row_count")
        ),
        "panel_comp_late_high_state_count": int(
            _summary_value(panel_competitiveness, "late_high_distance_state_count")
        ),
        "panel_comp_late_high_pm_count": int(
            _summary_value(
                panel_competitiveness,
                "late_high_distance_polymarket_lower_loss_count",
            )
        ),
        "panel_comp_late_high_poll_count": int(
            _summary_value(
                panel_competitiveness,
                "late_high_distance_poll_lower_loss_count",
            )
        ),
        "panel_comp_late_high_advantage": _summary_value(
            panel_competitiveness,
            "late_high_distance_mean_loss_advantage",
        ),
        "state_sign_late_non_safe_state_count": int(
            _summary_value(state_significance, "late_non_safe_state_count")
        ),
        "state_sign_late_non_safe_pm_state_count": int(
            _summary_value(
                state_significance,
                "late_non_safe_polymarket_majority_state_count",
            )
        ),
        "state_sign_late_non_safe_p_value": _summary_value(
            state_significance,
            "late_non_safe_polymarket_exact_binomial_p_value_greater",
        ),
        "state_sign_late_non_safe_ci_low": _summary_value(
            state_significance,
            "late_non_safe_polymarket_exact_95_ci_low",
        ),
        "state_sign_late_high_state_count": int(
            _summary_value(state_significance, "late_high_distance_state_count")
        ),
        "state_sign_late_high_pm_state_count": int(
            _summary_value(
                state_significance,
                "late_high_distance_polymarket_majority_state_count",
            )
        ),
        "state_sign_late_high_poll_state_count": int(
            _summary_value(
                state_significance,
                "late_high_distance_poll_majority_state_count",
            )
        ),
        "state_sign_late_high_poll_p_value": _summary_value(
            state_significance,
            "late_high_distance_poll_exact_binomial_p_value_greater",
        ),
        "calibration_forecast_case_rows": int(calibration_summary["case_count"].sum()),
        "calibration_forecast_source_count": int(len(calibration_summary)),
        "calibration_pairwise_count": int(len(calibration_pairwise)),
        "calibration_aggregate_support_count": _bool_count(
            calibration_pairwise,
            "aggregate_mean_supports_polymarket",
        ),
        "calibration_majority_support_count": _bool_count(
            calibration_pairwise,
            "majority_cases_supports_polymarket",
        ),
        "calibration_broad_support_count": _bool_count(
            calibration_pairwise,
            "broad_many_cases_claim_supported",
        ),
        "calibration_pm_state_brier": float(
            calibration_pm_state["mean_brier_loss"]
        ),
        "calibration_pm_state_ece": float(
            calibration_pm_state["expected_calibration_error"]
        ),
        "calibration_rieke_state_ece": float(
            calibration_rieke_state["expected_calibration_error"]
        ),
        "calibration_270_state_ece": float(
            calibration_270_state["expected_calibration_error"]
        ),
        "final_snapshot_case_count": int(_summary_value(final_snapshot, "case_count")),
        "final_snapshot_pm_lower_loss_count": int(
            _summary_value(final_snapshot, "polymarket_lower_loss_count")
        ),
        "final_snapshot_traditional_lower_loss_count": int(
            _summary_value(final_snapshot, "traditional_lower_loss_count")
        ),
        "final_snapshot_mean_pm_brier": _summary_value(
            final_snapshot,
            "mean_polymarket_brier",
        ),
        "final_snapshot_mean_traditional_brier": _summary_value(
            final_snapshot,
            "mean_traditional_brier",
        ),
        "final_snapshot_mean_loss_advantage": _summary_value(
            final_snapshot,
            "mean_loss_advantage",
        ),
        "state_poll_snapshot_case_count": int(
            _summary_value(state_poll_snapshot, "case_count")
        ),
        "state_poll_snapshot_pm_lower_loss_count": int(
            _summary_value(state_poll_snapshot, "polymarket_lower_loss_count")
        ),
        "state_poll_snapshot_poll_lower_loss_count": int(
            _summary_value(state_poll_snapshot, "poll_derived_lower_loss_count")
        ),
        "state_poll_snapshot_mean_pm_brier": _summary_value(
            state_poll_snapshot,
            "mean_polymarket_brier",
        ),
        "state_poll_snapshot_mean_poll_brier": _summary_value(
            state_poll_snapshot,
            "mean_poll_derived_brier",
        ),
        "state_poll_snapshot_mean_loss_advantage": _summary_value(
            state_poll_snapshot,
            "mean_loss_advantage",
        ),
        "popular_vote_case_count": int(_summary_value(popular_vote, "case_count")),
        "popular_vote_pm_lower_loss_count": int(
            _summary_value(popular_vote, "polymarket_lower_loss_count")
        ),
        "popular_vote_poll_lower_loss_count": int(
            _summary_value(popular_vote, "poll_derived_lower_loss_count")
        ),
        "popular_vote_mean_pm_brier": _summary_value(
            popular_vote,
            "mean_polymarket_brier",
        ),
        "popular_vote_mean_poll_brier": _summary_value(
            popular_vote,
            "mean_poll_derived_brier",
        ),
        "popular_vote_mean_loss_advantage": _summary_value(
            popular_vote,
            "mean_loss_advantage",
        ),
        "margin_threshold_candidate_count": int(len(margin_threshold_readiness)),
        "margin_threshold_with_538_poll_count": _bool_count(
            margin_threshold_readiness,
            "has_538_state_poll_rows",
        ),
        "margin_threshold_with_clob_overlap_count": _bool_count(
            margin_threshold_readiness,
            "has_clob_history_during_538_poll_window",
        ),
        "margin_threshold_compatible_count": _bool_count(
            margin_threshold_readiness,
            "compatible_for_h1_brier_now",
        ),
        "margin_threshold_no_overlap_count": int(
            margin_status.get("blocked_by_no_temporal_overlap", 0)
        ),
        "margin_threshold_missing_poll_count": int(
            margin_status.get("blocked_by_missing_538_state_poll_rows", 0)
        ),
        "state_poll_panel_case_count": int(
            _summary_value(state_poll_panel, "matched_case_count")
        ),
        "state_poll_panel_state_count": int(
            _summary_value(state_poll_panel, "matched_state_count")
        ),
        "state_poll_panel_date_count": int(
            _summary_value(state_poll_panel, "matched_date_count")
        ),
        "state_poll_panel_pm_lower_loss_count": int(
            _summary_value(state_poll_panel, "polymarket_lower_loss_count")
        ),
        "state_poll_panel_poll_lower_loss_count": int(
            _summary_value(state_poll_panel, "poll_derived_lower_loss_count")
        ),
        "state_poll_panel_mean_pm_brier": _summary_value(
            state_poll_panel,
            "mean_polymarket_brier",
        ),
        "state_poll_panel_mean_poll_brier": _summary_value(
            state_poll_panel,
            "mean_poll_derived_brier",
        ),
        "state_poll_panel_mean_loss_advantage": _summary_value(
            state_poll_panel,
            "mean_loss_advantage",
        ),
        "state_poll_temporal_support_months": str(
            temporal_support["included_months"]
        ).replace(",", ", "),
        "state_poll_temporal_support_row_count": int(
            temporal_support["row_count"]
        ),
        "state_poll_temporal_support_state_count": int(
            temporal_support["state_count"]
        ),
        "state_poll_temporal_support_pm_lower_loss_count": int(
            temporal_support["polymarket_lower_loss_count"]
        ),
        "state_poll_temporal_support_poll_lower_loss_count": int(
            temporal_support["poll_derived_lower_loss_count"]
        ),
        "state_poll_temporal_support_mean_pm_brier": float(
            temporal_support["mean_polymarket_brier"]
        ),
        "state_poll_temporal_support_mean_poll_brier": float(
            temporal_support["mean_poll_derived_brier"]
        ),
        "state_poll_temporal_full_pm_lower_loss_count": int(
            temporal_full["polymarket_lower_loss_count"]
        ),
        "state_poll_temporal_full_poll_lower_loss_count": int(
            temporal_full["poll_derived_lower_loss_count"]
        ),
        "state_poll_horizon_near_bins": str(
            horizon_near["included_horizon_bins"]
        ).replace(",", ", "),
        "state_poll_horizon_near_row_count": int(horizon_near["row_count"]),
        "state_poll_horizon_near_state_count": int(horizon_near["state_count"]),
        "state_poll_horizon_near_pm_lower_loss_count": int(
            horizon_near["polymarket_lower_loss_count"]
        ),
        "state_poll_horizon_near_poll_lower_loss_count": int(
            horizon_near["poll_derived_lower_loss_count"]
        ),
        "state_poll_horizon_near_mean_pm_brier": float(
            horizon_near["mean_polymarket_brier"]
        ),
        "state_poll_horizon_near_mean_poll_brier": float(
            horizon_near["mean_poll_derived_brier"]
        ),
        "state_poll_horizon_state_count": int(
            _summary_value(state_poll_panel_horizon_state, "state_count")
        ),
        "state_poll_horizon_state_pm_mean_support_count": int(
            _summary_value(
                state_poll_panel_horizon_state,
                "polymarket_mean_support_state_count",
            )
        ),
        "state_poll_horizon_state_pm_majority_support_count": int(
            _summary_value(
                state_poll_panel_horizon_state,
                "polymarket_majority_support_state_count",
            )
        ),
        "state_poll_horizon_state_poll_support_count": int(
            _summary_value(
                state_poll_panel_horizon_state,
                "poll_derived_or_no_polymarket_support_state_count",
            )
        ),
        "state_poll_near_quality_case_count": int(near_quality_pm["row_count"]),
        "state_poll_near_quality_forecast_row_count": int(
            near_quality_pm["row_count"] + near_quality_poll["row_count"]
        ),
        "state_poll_near_quality_state_count": int(near_quality_pm["state_count"]),
        "state_poll_near_quality_pm_mean_brier": float(
            near_quality_pm["mean_brier_loss"]
        ),
        "state_poll_near_quality_poll_mean_brier": float(
            near_quality_poll["mean_brier_loss"]
        ),
        "state_poll_near_quality_pm_ece": float(
            near_quality_pm["expected_calibration_error"]
        ),
        "state_poll_near_quality_poll_ece": float(
            near_quality_poll["expected_calibration_error"]
        ),
        "state_poll_near_quality_pm_separation": float(
            near_quality_pm["probability_separation"]
        ),
        "state_poll_near_quality_poll_separation": float(
            near_quality_poll["probability_separation"]
        ),
        "state_poll_sensitivity_row_count": int(len(sensitivity)),
        "state_poll_sensitivity_min_mae": float(
            sensitivity["poll_error_mae_points"].min()
        ),
        "state_poll_sensitivity_max_mae": float(
            sensitivity["poll_error_mae_points"].max()
        ),
        "state_poll_sensitivity_min_pm_lower_loss_count": int(
            sensitivity["polymarket_lower_loss_count"].min()
        ),
        "state_poll_sensitivity_max_pm_lower_loss_count": int(
            sensitivity["polymarket_lower_loss_count"].max()
        ),
        "state_poll_sensitivity_min_mean_loss_advantage": float(
            sensitivity["mean_loss_advantage"].min()
        ),
        "state_poll_sensitivity_max_mean_loss_advantage": float(
            sensitivity["mean_loss_advantage"].max()
        ),
        "state_poll_coverage_state_count": int(len(state_poll_coverage)),
        "state_poll_coverage_polymarket_market_count": int(
            state_poll_coverage["polymarket_market_available"].sum()
        ),
        "state_poll_coverage_poll_snapshot_count": int(
            state_poll_coverage["poll_snapshot_has_rep_dem"].sum()
        ),
        "state_poll_coverage_valid_pair_count": int(
            state_poll_coverage["included_in_brier_comparison"].sum()
        ),
        "state_poll_coverage_missing_poll_count": int(
            coverage_status.get("excluded_missing_538_poll_snapshot", 0)
        ),
        "state_poll_coverage_missing_both_count": int(
            coverage_status.get("excluded_missing_both_sources", 0)
        ),
        "rieke_state_case_count": int(
            _summary_value(rieke_state_forecast, "case_count")
        ),
        "rieke_state_pm_lower_loss_count": int(
            _summary_value(rieke_state_forecast, "polymarket_lower_loss_count")
        ),
        "rieke_state_rieke_lower_loss_count": int(
            _summary_value(rieke_state_forecast, "rieke_lower_loss_count")
        ),
        "rieke_state_mean_pm_brier": _summary_value(
            rieke_state_forecast,
            "mean_polymarket_brier",
        ),
        "rieke_state_mean_rieke_brier": _summary_value(
            rieke_state_forecast,
            "mean_rieke_brier",
        ),
        "rieke_state_mean_loss_advantage": _summary_value(
            rieke_state_forecast,
            "mean_loss_advantage",
        ),
        "two_seventy_state_case_count": int(
            _summary_value(two_seventy_state_forecast, "case_count")
        ),
        "two_seventy_state_exact_case_count": int(
            _summary_value(two_seventy_state_forecast, "exact_probability_case_count")
        ),
        "two_seventy_state_censored_case_count": int(
            _summary_value(two_seventy_state_forecast, "censored_boundary_case_count")
        ),
        "two_seventy_state_pm_lower_loss_count": int(
            _summary_value(two_seventy_state_forecast, "polymarket_lower_loss_count")
        ),
        "two_seventy_state_270_lower_loss_count": int(
            _summary_value(two_seventy_state_forecast, "two_seventy_lower_loss_count")
        ),
        "two_seventy_state_tie_count": int(
            _summary_value(two_seventy_state_forecast, "tie_count")
        ),
        "two_seventy_state_exact_pm_lower_loss_count": int(
            _summary_value(
                two_seventy_state_forecast,
                "exact_probability_polymarket_lower_loss_count",
            )
        ),
        "two_seventy_state_exact_270_lower_loss_count": int(
            _summary_value(
                two_seventy_state_forecast,
                "exact_probability_two_seventy_lower_loss_count",
            )
        ),
        "two_seventy_state_exact_tie_count": int(
            _summary_value(two_seventy_state_forecast, "exact_probability_tie_count")
        ),
        "two_seventy_state_mean_pm_brier": _summary_value(
            two_seventy_state_forecast,
            "mean_polymarket_brier",
        ),
        "two_seventy_state_mean_270_brier": _summary_value(
            two_seventy_state_forecast,
            "mean_two_seventy_brier",
        ),
        "two_seventy_state_mean_loss_advantage": _summary_value(
            two_seventy_state_forecast,
            "mean_loss_advantage",
        ),
        "two_seventy_state_exact_mean_pm_brier": _summary_value(
            two_seventy_state_forecast,
            "exact_probability_mean_polymarket_brier",
        ),
        "two_seventy_state_exact_mean_270_brier": _summary_value(
            two_seventy_state_forecast,
            "exact_probability_mean_two_seventy_brier",
        ),
        "two_seventy_poll_average_case_count": int(
            _summary_value(two_seventy_poll_average, "case_count")
        ),
        "two_seventy_poll_average_state_rows": int(
            _summary_value(two_seventy_poll_average, "poll_average_state_rows")
        ),
        "two_seventy_poll_average_missing_state_count": int(
            _summary_value(two_seventy_poll_average, "poll_average_missing_state_count")
        ),
        "two_seventy_poll_average_pm_lower_loss_count": int(
            _summary_value(two_seventy_poll_average, "polymarket_lower_loss_count")
        ),
        "two_seventy_poll_average_poll_lower_loss_count": int(
            _summary_value(two_seventy_poll_average, "poll_derived_lower_loss_count")
        ),
        "two_seventy_poll_average_mean_pm_brier": _summary_value(
            two_seventy_poll_average,
            "mean_polymarket_brier",
        ),
        "two_seventy_poll_average_mean_poll_brier": _summary_value(
            two_seventy_poll_average,
            "mean_poll_derived_brier",
        ),
        "two_seventy_poll_average_mean_loss_advantage": _summary_value(
            two_seventy_poll_average,
            "mean_loss_advantage",
        ),
    }


def _h2_data(summary: pd.DataFrame, event_seed: pd.DataFrame) -> dict[str, Any]:
    primary = summary[
        summary["summary_id"].str.contains("primary_0d_to_1d", na=False)
    ].copy()
    examples = []
    for _, row in primary.iterrows():
        event = str(row["label"]).split(" | ")[0].replace("evt_2024_", "")
        examples.append({"event": event, "change_pp": float(row["value"]) * 100.0})
    return {
        "event_count": int(len(event_seed)),
        "summary_rows": int(len(summary)),
        "primary_examples": examples,
    }


def _h3_data(summary: pd.DataFrame) -> dict[str, Any]:
    tier_rows = summary[summary["summary_type"] == "wallet_tier"]
    tier_counts = {
        str(row["label"]): int(float(row["value"]))
        for _, row in tier_rows.iterrows()
    }
    corr_rows = summary[summary["summary_type"] == "lead_lag_correlation"].copy()
    corr_rows["abs_value"] = corr_rows["value"].astype(float).abs()
    top_corr = corr_rows.sort_values("abs_value", ascending=False).iloc[0]
    granger_rows = summary[summary["summary_type"] == "granger"].copy()
    granger_rows["value"] = granger_rows["value"].astype(float)
    min_granger = granger_rows.sort_values("value").iloc[0]
    return {
        "model_rows": int(_summary_value(summary, "h3_model_row_count")),
        "tier_counts": tier_counts,
        "tier_counts_text": ", ".join(f"{key}: {value}" for key, value in tier_counts.items()),
        "top_correlation_label": str(top_corr["label"]),
        "top_correlation": float(top_corr["value"]),
        "min_granger_label": str(min_granger["label"]),
        "min_granger_p": float(min_granger["value"]),
    }


def _literature_data(literature: pd.DataFrame) -> dict[str, Any]:
    """Return the bounded literature frame used for the supervisor report."""

    selected_ids = [
        "lit_emh_001",
        "lit_brier_001",
        "lit_dm_001",
        "lit_eventstudy_001",
        "lit_granger_001",
        "zotero_poly_001",
        "zotero_poly_002",
        "zotero_poly_005",
        "zotero_poly_006",
        "zotero_poly_007",
    ]
    role_by_id = {
        "lit_emh_001": "Theorie: informationelle Effizienz und EMH-Proxy-Logik.",
        "lit_brier_001": "H1-Methode: Probability-Forecast-Verifikation mit Brier-Verlust.",
        "lit_dm_001": "H1-Methode: Vergleich konkurrierender Forecast-Loss-Serien.",
        "lit_eventstudy_001": "H2-Methode: Event-Window-Design und Grenzen von Ereignisstudien.",
        "lit_granger_001": "H3-Methode: Lead-Lag-Diagnostik mit vorsichtiger Kausalitaetsabgrenzung.",
        "zotero_poly_001": "Polymarket-Kontext: Transaktionslogik, Wallet- und Volumen-Caveats.",
        "zotero_poly_002": "H1-Kontext: Prediction Markets versus Polling/Forecasting.",
        "zotero_poly_005": "H3-Kontext: Abgrenzung von Information, Informationsvorsprung- und Ethikfragen.",
        "zotero_poly_006": "Marktmikrostruktur: Maker/Taker, Bias- und Risikocaveats.",
        "zotero_poly_007": "Polymarket-Forschungskontext: Konvergenz, Volatilitaet und Biases.",
    }
    research_note_by_id = {
        "lit_emh_001": "Preise als Informationsaggregate motivieren die Proxy-Tests, beweisen aber keine Effizienz.",
        "lit_brier_001": "Begruendet H1 als Verlustvergleich von Wahrscheinlichkeitsprognosen.",
        "lit_dm_001": "Begruendet den Test auf Unterschiede in vorliegenden Forecast-Verlustreihen.",
        "lit_eventstudy_001": "Begruendet H2 als Ereignisfenster-Design statt freier News-Interpretation.",
        "lit_granger_001": "Begruendet H3 als Vorhersage-/Timingdiagnostik, nicht als starker Ursachenbeweis.",
        "zotero_poly_001": "Stuetzt Vorsicht bei on-chain Volumen, Wallet-Flows und Austausch-Equivalenten.",
        "zotero_poly_002": "Stuetzt die Vergleichsfrage Polymarket versus Polling, ersetzt aber keine lokale Transformation.",
        "zotero_poly_005": "Hilft bei der ethischen Abgrenzung von Informationsvorsprung und Marktpreisen.",
        "zotero_poly_006": "Stuetzt Mikrostruktur- und Bias-Caveats fuer spaetere Monitor-/Strategieformulierungen.",
        "zotero_poly_007": "Positioniert Polymarket-Forschung mit Konvergenz-, Volatilitaets- und Bias-Grenzen.",
    }
    sources = []
    indexed = literature.set_index("source_id")
    for source_id in selected_ids:
        if source_id not in indexed.index:
            continue
        row = indexed.loc[source_id]
        sources.append(
            {
                "source_id": source_id,
                "title": str(row["title"]),
                "authors": str(row["authors"]),
                "year": str(row["year"]),
                "venue": str(row["venue"]),
                "url": str(row["url"]),
                "topic": str(row["topic"]),
                "hypothesis": str(row["hypothesis"]),
                "status": str(row["status"]),
                "role": role_by_id[source_id],
                "research_note": research_note_by_id[source_id],
            }
        )
    status_counts = literature["status"].value_counts().sort_index()
    return {
        "source_count": int(len(literature)),
        "selected_source_count": len(sources),
        "status_counts_text": ", ".join(
            f"{status}: {int(count)}" for status, count in status_counts.items()
        ),
        "sources": sources,
        "citation_boundary": (
            "Die Quellen sind als lokaler Literaturrahmen hinterlegt. "
            "Thesis-facing Detailclaims duerfen erst nach Vollreview als "
            "reviewed oder cited verwendet werden; candidate/rejected Quellen "
            "tragen keine Ergebnisbehauptungen."
        ),
    }


def _source_review_worksheet_data(worksheet: pd.DataFrame) -> dict[str, Any]:
    """Return compact source-review worksheet counts for the report."""

    return {
        "worksheet_rows": int(len(worksheet)),
        "priority_1_rows": int(
            (worksheet["priority_band"] == "priority_1_method_foundation_review").sum()
        ),
        "blocked_rows": int((worksheet["priority_band"] == "blocked_or_future_work_only").sum()),
        "pending_rows": int((worksheet["reviewer_decision"] == "pending").sum()),
    }


def _monitor_data(
    summary: pd.DataFrame,
    anomaly_review_summary: pd.DataFrame,
) -> dict[str, Any]:
    severity = {
        str(row["label"]): int(float(row["value"]))
        for _, row in summary[summary["summary_type"] == "direct_severity_count"].iterrows()
    }
    dashboard_meta = _read_json("data/results/monitor_v2_polymarket_dashboard_metadata.json")
    graph_meta = _read_json("data/results/wallet_graph_metadata.json")
    outputs = dashboard_meta.get("outputs", {})
    graph_outputs = graph_meta.get("outputs", {})
    anomaly = anomaly_review_summary.iloc[0].to_dict()
    return {
        "snapshot_count": int(
            _summary_value(summary, "monitor_v2_snapshot_count", default=0)
        ),
        "severity_counts": severity,
        "severity_counts_text": ", ".join(f"{k}: {v}" for k, v in severity.items()),
        "live_market_count": int(outputs.get("market_count", 0)),
        "live_alert_count": int(outputs.get("alert_count", 0)),
        "live_scoring_rows": int(outputs.get("scoring_row_count", 0)),
        "wallet_graph_nodes": int(graph_outputs.get("node_count", 0)),
        "wallet_graph_edges": int(graph_outputs.get("edge_count", 0)),
        "anomaly_queue_rows": int(anomaly.get("queue_row_count", 0)),
        "anomaly_high_priority_count": int(anomaly.get("high_priority_count", 0)),
        "anomaly_medium_priority_count": int(anomaly.get("medium_priority_count", 0)),
        "anomaly_low_priority_count": int(anomaly.get("low_priority_count", 0)),
        "anomaly_review_labels": str(anomaly.get("review_label_counts", "")),
        "anomaly_review_status_counts": str(
            anomaly.get("human_review_status_counts", "")
        ),
        "anomaly_allowed_interpretation": str(
            anomaly.get("allowed_interpretation", "")
        ),
        "anomaly_limitation": (
            "Die Queue ist kein Nachweis fuer Ursachen, Regelverstoss, "
            "Handelbarkeit, Profitabilitaet oder zukuenftige Entwicklung."
        ),
    }


def _swiss_data(
    comparison: pd.DataFrame,
    latest_source: pd.DataFrame,
    information: pd.DataFrame,
    polls: pd.DataFrame,
) -> dict[str, Any]:
    metadata = _read_json("data/results/swiss_referendum_10mio_efficiency_metadata.json")
    outputs = metadata.get("outputs", {})
    latest = comparison.sort_values("collected_at_utc").iloc[-1]
    info_counts = information["information_processing_label"].value_counts().sort_index()
    latest_source_rows = []
    for _, row in latest_source.sort_values("source_name").iterrows():
        latest_source_rows.append(
            {
                "source": str(row["source_name"]),
                "poll_id": str(row["poll_id"]),
                "poll_yes": float(row["poll_yes_share"]) * 100.0,
                "raw_gap_pp": float(row["raw_yes_gap"]) * 100.0,
            }
        )
    return {
        "poll_count": int(len(polls)),
        "snapshot_count": int(outputs.get("snapshot_count", len(comparison))),
        "history_rows": int(outputs.get("history_row_count", 0)),
        "latest_poly_yes_pct": float(latest["polymarket_yes_probability"]) * 100.0,
        "latest_poll_yes_pct": float(latest["poll_yes_share"]) * 100.0,
        "latest_raw_gap_pp": float(latest["raw_yes_gap"]) * 100.0,
        "latest_decided_gap_pp": float(latest["decided_yes_gap"]) * 100.0,
        "information_response_counts_text": ", ".join(
            f"{label}: {int(count)}" for label, count in info_counts.items()
        ),
        "latest_source_rows": latest_source_rows,
    }


def _database_summary() -> dict[str, Any]:
    db_path = REPO_ROOT / "data" / "thesis.db"
    if not db_path.exists():
        return {"table_count": 0, "tables": {}}
    con = sqlite3.connect(db_path)
    try:
        tables = [
            row[0]
            for row in con.execute(
                "SELECT name FROM sqlite_master WHERE type='table' ORDER BY name"
            ).fetchall()
        ]
        counts = {
            table: int(con.execute(f'SELECT COUNT(*) FROM "{table}"').fetchone()[0])
            for table in tables
        }
    finally:
        con.close()
    return {"table_count": len(tables), "tables": counts}


def _folder_inventory() -> dict[str, int]:
    folders = [
        "data/results",
        "docs/research",
        "docs/project",
        "operations/analysis",
        "operations/collectors",
        "operations/project",
        "tests",
        "ingest",
        "legacy",
    ]
    inventory: dict[str, int] = {}
    for name in folders:
        folder = REPO_ROOT / name
        if not folder.exists():
            inventory[name] = 0
        elif name in {"data/results", "legacy"}:
            inventory[name] = sum(1 for path in folder.rglob("*") if path.is_file())
        else:
            inventory[name] = sum(1 for path in folder.iterdir() if path.is_file())
    return inventory


def _status_test_summary() -> str:
    status_path = REPO_ROOT / "STATUS.md"
    if not status_path.exists():
        return "Statusdatei fehlt"
    text = status_path.read_text(encoding="utf-8")
    for line in text.splitlines():
        if line.startswith("Pytest summary:"):
            return line.replace("Pytest summary:", "").strip().strip("`")
    return "Pytest summary nicht gefunden"


def _figure_specs() -> list[FigureSpec]:
    specs = [
        (
            "h1_forecast_quality.png",
            "H1 Forecast-Quality Vergleich",
            "Zeigt Brier Scores, Head-to-head-Tagesverluste und Forecast-Zeitreihen ohne ueberstarken Kalibrierungsanspruch.",
        ),
        (
            "h1_forecast_quality_synthesis.png",
            "H1 Forecast-Quality Synthesis",
            "Fasst alle aktuellen H1-Vergleichsquellen zusammen und trennt aggregierte Brier-Stuetze von einem breiten Viele-Faelle-Beweis.",
        ),
        (
            "h1_claim_evidence_audit.png",
            "H1 Claim-Evidence Audit",
            "Fuehrt die H1-Evidenz als Claim-Ledger zusammen: spaete Polymarket-Stuetze, widersprechendes Full-Panel und weiterhin nicht belegter breiter User-Claim.",
        ),
        (
            "h1_poll_comparison_result.png",
            "H1 Poll-Comparison Result",
            "Verdichtet die direkt pollbezogene H1-Evidenz: 262 von 285 spaeten Low/Middle-Poll-Distanz-Zeilen und 9 von 9 States stuetzen Polymarket, waehrend Vollpanel und High-Distance-Zeilen Grenzen bleiben.",
        ),
        (
            "h1_poll_claim_readiness.png",
            "H1 Poll-Claim Readiness",
            "Trennt die aktuell belegbare bounded Aussage von Gegenbeispiel-Scopes: Polymarket ist im <=90-Tage Low/Middle-Poll-Distanz-Scope stark, aber der breite Claim bleibt nicht belegt.",
        ),
        (
            "h1_poll_scope_frontier.png",
            "H1 Poll-Scope Frontier",
            "Visualisiert systematisch, wie weit sich der H1-Poll-Scope nach Horizont und quantilbasierter Poll-Distanz ausweiten laesst: groesster robuster Scope <=120 Tage Low/Middle, Vollpanel bleibt Gegenbeleg.",
        ),
        (
            "h1_poll_decision_matrix.png",
            "H1 Poll-Decision Matrix",
            "Verdichtet die H1-Poll-Evidenz in eine Claim-Matrix: robuste bounded-Yes-Zeilen, Mean-Loss-Stuetze ohne Mehrheit, Kalibrierungskontext und Vollpanel-Gegenbeleg.",
        ),
        (
            "h1_robust_poll_scope_quality.png",
            "H1 Robust Poll-Scope Quality",
            "Visualisiert Mean Brier, Fixed-Bin-ECE, Kalibrierungsbins und Lower-Loss-Zaehler fuer die robusten late Low/Middle-Poll-Distanz-Scopes.",
        ),
        (
            "h1_robust_poll_scope_unit_quality.png",
            "H1 Robust Poll-Scope Unit Quality",
            "Aggregiert die beiden robusten Poll-Scopes zu State-, State-Month-, State-Horizon- und Horizon-Tier-Einheiten, damit der bounded H1-Befund weniger von wiederholten State-Date-Zeilen abhaengt.",
        ),
        (
            "h1_poll_comparison_unit_robustness.png",
            "H1 Poll-Comparison Unit Robustness",
            "Aggregiert den primaeren H1-Poll-Scope zu State-, State-Month-, State-Horizon- und Horizon-Tier-Einheiten; Polymarket wird in allen primaeren Einheiten gestuetzt, waehrend Full-Panel und High-Distance Grenzen bleiben.",
        ),
        (
            "h1_direct_poll_loss_decomposition.png",
            "H1 Direct Poll Loss Decomposition",
            "Zerlegt direkte Poll-Transform-Vergleiche: Polymarket hat den niedrigeren mittleren Brier, obwohl poll-derived Comparatoren mehr Einzel-Faelle gewinnen.",
        ),
        (
            "h1_direct_poll_state_cluster_diagnostic.png",
            "H1 Direct Poll State-Cluster Diagnostic",
            "Prueft direkte Poll-Transform-Vergleiche mit gleichgewichteten State-Clustern: Der mittlere Verlustvorteil bleibt positiv, aber die State-Mehrheit stuetzt poll-derived Comparatoren.",
        ),
        (
            "h1_direct_poll_outlier_robustness.png",
            "H1 Direct Poll Outlier Robustness",
            "Prueft, ob der direkte Poll-State-Cluster-Vorteil nur von einzelnen Ausreissern getragen wird: alle Leave-one-state-out Means bleiben positiv, aber Top-k-Exclusions zeigen Konzentration.",
        ),
        (
            "h1_calibration_diagnostic.png",
            "H1 Calibration Diagnostic",
            "Visualisiert feste Kalibrierungsbins, Mean Brier, ECE und Pairwise-Lower-Loss-Zaehler aus geloesten H1-Fallartefakten.",
        ),
        (
            "h1_evidence_scope.png",
            "H1 Evidence-Scope Audit",
            "Trennt die 194 taeglichen Forecast-Paare von der Anzahl unabhaengiger geloester H1-Outcomes.",
        ),
        (
            "h1_expansion_readiness.png",
            "H1 Expansion-Readiness Audit",
            "Zeigt, dass zusaetzliche Polymarket-Tagespreise ohne kompatible Probability-Forecast-Vergleichsreihe noch keine weiteren H1-Brier-Paare ergeben.",
        ),
        (
            "h1_margin_threshold_readiness.png",
            "H1 Margin-Threshold Readiness",
            "Prueft sieben Trump-State-Margin-Maerkte und zeigt, dass ohne zeitliche Ueberlappung zwischen bewahrten 538-Poll-Averages und CLOB-Historie keine neuen H1-Brier-Faelle entstehen.",
        ),
        (
            "h1_final_snapshot.png",
            "H1 Final-Snapshot Extension",
            "Vergleicht acht geloeste 2024-Final-Snapshot-Outcomes gegen 538 final forecast; kleine Erweiterung, kein Viele-Faelle-Beweis.",
        ),
        (
            "h1_state_poll_snapshot.png",
            "H1 State-Poll-Snapshot Extension",
            "Vergleicht 13 geloeste State-Outcomes gegen eine dokumentiert transformierte 538 Polling-Average-Wahrscheinlichkeit; nicht Rohpoll und kein offizieller 538 State-Forecast.",
        ),
        (
            "h1_270towin_poll_average.png",
            "H1 270toWin Polling-Average Extension",
            "Vergleicht 43 gematchte State-Outcomes gegen eine dokumentiert transformierte 270toWin-Polling-Average-Wahrscheinlichkeit.",
        ),
        (
            "h1_popular_vote.png",
            "H1 Popular-Vote Extension",
            "Vergleicht 51 nationale 538-Poll-Transform-Tageszeilen mit dem Polymarket-Trump-popular-vote-Markt und zeigt einen Gegenbeleg zum starken Claim.",
        ),
        (
            "h1_state_poll_panel.png",
            "H1 State-Date Poll Panel",
            "Vergleicht 1720 gematchte State-Date-Zeilen gegen transformierte 538 Polling-Averages; das groessere Panel spricht gegen den starken Polymarket-Claim.",
        ),
        (
            "h1_state_poll_panel_temporal_diagnostic.png",
            "H1 State-Date Poll Panel Temporal Diagnostic",
            "Zeigt, dass der Vollpanel-Befund gegen Polymarket spricht, waehrend August und September 2024 als diagnostischer Teilbereich Polymarket stuetzen.",
        ),
        (
            "h1_state_poll_panel_horizon_diagnostic.png",
            "H1 State-Date Poll Panel Horizon Diagnostic",
            "Zeigt, dass das <=90-Tage-Forecast-Fenster Polymarket stuetzt, waehrend der Vollpanel-Befund weiter gegen den starken Claim spricht.",
        ),
        (
            "h1_state_poll_panel_horizon_state_support.png",
            "H1 <=90-Day State-Level Support",
            "Aggregiert das <=90-Tage-Fenster auf State-Ebene: Polymarket wird in 8 von 13 States nach mittlerem Brier und Row-Majority gestuetzt.",
        ),
        (
            "h1_state_poll_panel_near_window_quality.png",
            "H1 <=90-Day Score Quality",
            "Visualisiert Mean Brier, Fixed-Bin-ECE, Probability-Separation und lower-loss rows im <=90-Tage-Fenster des State-Date-Panels.",
        ),
        (
            "h1_state_poll_snapshot_sensitivity.png",
            "H1 Poll-Transform Sensitivity",
            "Prueft die State-Poll-Erweiterung ueber MAE-Annahmen von 2.0 bis 10.0 Prozentpunkten, ohne den Parameter auf Outcomes zu fitten.",
        ),
        (
            "h1_state_poll_snapshot_coverage.png",
            "H1 State-Poll Coverage Audit",
            "Zeigt, warum 50 States und 50 Polymarket-State-Maerkte nur 13 valide H1-Brier-Paare mit dem bewahrten 538-Polling-Average-Snapshot ergeben.",
        ),
        (
            "h1_rieke_state_forecast.png",
            "H1 Rieke 50-State Forecast Extension",
            "Vergleicht 50 Polymarket State-Winner-Maerkte mit dem pollbasierten Rieke-Forecast; Polymarket hat niedrigeren mittleren Brier, aber nur in 12 von 50 Einzelstaaten niedrigeren Verlust.",
        ),
        (
            "h1_270towin_state_forecast.png",
            "H1 270toWin/JHK 50-State Forecast Extension",
            "Vergleicht 50 Polymarket State-Winner-Maerkte mit 270toWin/JHK; Polymarket hat niedrigeren mittleren Brier, aber nur in 9 von 50 Einzelstaaten niedrigeren Verlust.",
        ),
        (
            "h1_state_source_consensus.png",
            "H1 State-Source Consensus Diagnostic",
            "Aggregiert bestehende H1-State-Artefakte ueber 156 Source-State-Vergleiche und trennt All-Source-Konsens von direktem Poll-Transform-Konsens.",
        ),
        (
            "h1_competitive_state_diagnostic.png",
            "H1 Competitive-State Diagnostic",
            "Quantilbasierte Diagnose: Polymarket ist in der niedrigsten Distanz-/kompetitivsten Terzile besser, sichere States bleiben Gegenbeleg.",
        ),
        (
            "h1_state_poll_panel_competitiveness.png",
            "H1 State-Date Competitiveness x Horizon",
            "Zeigt, dass Polymarket im <=90-Tage-Fenster bei Low/Middle-Poll-Distanz 262 von 285 State-Date-Zeilen gewinnt, waehrend spaete High-Distance-Zeilen Gegenbeleg bleiben.",
        ),
        (
            "h1_state_poll_panel_state_significance.png",
            "H1 State-Level Significance Diagnostic",
            "Zeigt den exakten State-Level-Sign-Test fuer spaete Low/Middle-Poll-Distanz-Faelle: Polymarket 9 von 9 States, einseitiger p-Wert 0.0020.",
        ),
        (
            "thesis_h2_event_window_car.png",
            "H2 Event-Window Movement",
            "Zeigt taegliche Event-Window-Bewegungen fuer die kuratierten Ereignisse.",
        ),
        (
            "thesis_h3_wallet_tier_counts.png",
            "H3 Wallet-Tier-Verteilung",
            "Zeigt, dass Wallet-Tiers aus der beobachteten Verteilung abgeleitet wurden.",
        ),
        (
            "thesis_h3_granger_pvalues.png",
            "H3 Granger-Diagnostik",
            "Fasst predictive timing diagnostics zusammen, ohne Kausalitaet zu behaupten.",
        ),
        (
            "thesis_h3_event_wallet_anomalies.png",
            "Historische Event-Wallet-Anomalien",
            "Zeigt den pausierten Monitor-Prototyp als deskriptive Review-Schicht.",
        ),
        (
            "monitor_v2_polymarket_rolling_history.png",
            "Monitor-v2 Rolling History",
            "Visualisiert die kurze read-only Polymarket-Monitor-Historie.",
        ),
        (
            "swiss_referendum_10mio_efficiency.png",
            "Swiss Referendum: Polymarket vs Polls",
            "Vergleicht die lokale Polymarket-Wahrscheinlichkeit mit kuratierten Umfragen.",
        ),
        (
            "swiss_referendum_10mio_reaction_windows.png",
            "Swiss Referendum: Reaction Windows",
            "Zeigt beschreibende 1h/6h/24h/48h-Fenster nach Poll-Releases.",
        ),
        (
            "swiss_referendum_10mio_information_response.png",
            "Swiss Referendum: Information Response",
            "Zeigt Richtungsgleichheit zwischen neuer Poll-Signalrichtung und Polymarket-Bewegungen.",
        ),
    ]
    return [
        FigureSpec(path=RESULTS_DIR / filename, caption=caption, note=note)
        for filename, caption, note in specs
        if (RESULTS_DIR / filename).exists()
    ]


def _summary_value(
    frame: pd.DataFrame,
    summary_id: str,
    *,
    default: float | None = None,
) -> float:
    rows = frame.loc[frame["summary_id"] == summary_id, "value"]
    if rows.empty:
        if default is not None:
            return default
        raise ValueError(f"summary_id not found: {summary_id}")
    return float(rows.iloc[0])


def _summary_text_value(frame: pd.DataFrame, summary_id: str) -> str:
    rows = frame.loc[frame["summary_id"] == summary_id, "value"]
    if rows.empty:
        raise ValueError(f"summary_id not found: {summary_id}")
    return str(rows.iloc[0])


def _bool_count(frame: pd.DataFrame, column: str) -> int:
    if column not in frame.columns:
        raise ValueError(f"boolean column not found: {column}")
    return int(frame[column].astype(str).str.lower().eq("true").sum())


def _read_csv(path: str) -> pd.DataFrame:
    full_path = REPO_ROOT / path
    if not full_path.exists():
        raise FileNotFoundError(f"required report source missing: {path}")
    return pd.read_csv(full_path)


def _read_json(path: str) -> dict[str, Any]:
    full_path = REPO_ROOT / path
    if not full_path.exists():
        return {}
    return json.loads(full_path.read_text(encoding="utf-8"))


def _add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    _shade_cell(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    _set_run(run, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def _add_bullets(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def _add_table(
    doc: Document,
    rows: Sequence[Sequence[Any]],
    headers: Sequence[str],
    widths: Sequence[int],
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = str(header)
        _set_cell_font(cell, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[idx].text = str(value)
            _set_cell_font(cells[idx])
    _set_table_width(table, widths)
    return table


def _shade_table_header(table) -> None:
    for cell in table.rows[0].cells:
        _shade_cell(cell, LIGHT_FILL)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_width(table, widths: Sequence[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)


def _set_cell_font(cell, *, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
        for run in paragraph.runs:
            _set_run(run, size=9.5, bold=bold)


def _set_paragraph_font(paragraph, *, size: float, color: RGBColor) -> None:
    for run in paragraph.runs:
        _set_run(run, size=size, color=color)


def _set_run(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: RGBColor | None = None,
) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--markdown-output", type=Path, default=DEFAULT_MD_OUTPUT)
    parser.add_argument("--html-output", type=Path, default=DEFAULT_HTML_OUTPUT)
    parser.add_argument("--docx-output", type=Path, default=DEFAULT_DOCX_OUTPUT)
    parser.add_argument("--asset-dir", type=Path, default=DEFAULT_ASSET_DIR)
    return parser


def main(argv: Sequence[str] | None = None) -> int:
    args = build_parser().parse_args(argv)
    result = build_report(
        markdown_output=args.markdown_output,
        html_output=args.html_output,
        docx_output=args.docx_output,
        asset_dir=args.asset_dir,
    )
    print(json.dumps(result, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
